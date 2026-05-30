from __future__ import annotations

import math
from dataclasses import dataclass
from pathlib import Path
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX_PATH = ROOT / "output" / "doc" / "BaoCao_DATN_2_14_rewritten.docx"
TARGET_DOCX_PATH = ROOT / "output" / "doc" / "BaoCao_DATN_2_14_rewritten_with_143_diagram.docx"
DRAWIO_PATH = ROOT / "output" / "diagrams" / "drawio" / "1_4_3_mo_hinh_bai_toan_tong_quat.drawio"
PNG_PATH = ROOT / "output" / "diagrams" / "png" / "1_4_3_mo_hinh_bai_toan_tong_quat.png"

FONT_REGULAR = Path(r"C:\Windows\Fonts\times.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\timesbd.ttf")


@dataclass
class Node:
    node_id: str
    x: int
    y: int
    w: int
    h: int
    text: str
    style: str


@dataclass
class Edge:
    edge_id: str
    source: str
    target: str
    label: str = ""
    dashed: bool = False
    points: list[tuple[int, int]] | None = None


PAGE_WIDTH = 1760
PAGE_HEIGHT = 560

GROUP_STYLE = (
    "rounded=1;arcSize=18;whiteSpace=wrap;html=1;"
    "fillColor=#ffffff;strokeColor=#000000;fontColor=#000000;"
    "fontSize=16;fontStyle=1;fontFamily=Times New Roman;align=center;verticalAlign=middle;spacing=8;"
)


NODES = [
    Node(
        "g1",
        60,
        120,
        360,
        260,
        "1. Định danh và chuẩn bị vào trận\n\n"
        "PlayFabLeaderboardManager\n"
        "NameInputPanel\n"
        "ChallengePostNPC / ChallengePanel",
        GROUP_STYLE,
    ),
    Node(
        "g2",
        480,
        120,
        360,
        260,
        "2. Gameplay runtime\n\n"
        "PlayerController / PlayerAttack\n"
        "Enemy / WaveSpawner\n"
        "MapThemeManager / ObjectPool",
        GROUP_STYLE,
    ),
    Node(
        "g3",
        900,
        120,
        360,
        260,
        "3. Progression / Buff\n\n"
        "ExpDropper / PlayerLevelSystem\n"
        "BuffCardManager\n"
        "CardSelectionPanel",
        GROUP_STYLE,
    ),
    Node(
        "g4",
        1320,
        120,
        360,
        260,
        "4. UI / Backend\n\n"
        "GameUI / PauseMenuPanel\n"
        "LeaderboardPanel / HighScore\n"
        "PlayFab Service / Result Sync",
        GROUP_STYLE,
    ),
]

EDGES = [
    Edge("e1", "g1", "g2", "vào trận"),
    Edge("e2", "g2", "g3", "EXP / Level"),
    Edge("e3", "g3", "g4", "HUD / kết quả"),
]

SECTION_143_PARAGRAPHS = [
    (
        "Mô hình tổng quát",
        "Mô hình bài toán tổng quát của hệ thống được chia thành bốn khối chính. "
        "Khối định danh và chuẩn bị vào trận phụ trách đăng nhập PlayFab, kiểm tra hồ sơ và mở điều kiện bắt đầu trận. "
        "Khối gameplay runtime chịu trách nhiệm điều khiển player, tổ chức enemy, wave, object pool và thay đổi theme bản đồ. "
        "Khối progression và buff tiếp nhận dữ liệu EXP từ chiến đấu, xử lý lên cấp và sinh lựa chọn tăng sức mạnh. "
        "Khối UI và backend đảm nhiệm việc phản ánh trạng thái trận đấu, tạm dừng, tổng kết và đồng bộ dữ liệu online."
    ),
    (
        "Quan hệ giữa các khối",
        "Dữ liệu hệ thống đi theo chuỗi: PlayFab / NameInput / Challenge tạo điều kiện vào trận; "
        "sau đó Player, Enemy, Wave và Theme tạo ra trạng thái runtime; "
        "trạng thái này sinh EXP, level và buff; cuối cùng GameUI, PauseMenuPanel, LeaderboardPanel và PlayFab tiếp nhận dữ liệu đã xử lý để hiển thị và lưu kết quả. "
        "Mỗi khối vừa nhận dữ liệu từ khối trước, vừa trả tín hiệu điều phối cho khối sau."
    ),
    (
        "Ý nghĩa mô hình",
        "Phần 1.4.3 sử dụng sơ đồ khối để mô tả cấu trúc bài toán và quan hệ dữ liệu giữa các nhóm chức năng, "
        "không đi vào thứ tự thời gian chi tiết của từng bước xử lý. "
        "Vì vậy, đây là mô hình bài toán tổng quát của hệ thống, còn phần 1.4.4 tiếp tục đảm nhiệm vai trò mô tả luồng hoạt động theo trình tự xảy ra trong một lượt chơi."
    ),
    (
        "Liên hệ sơ đồ chi tiết",
        "Sơ đồ khối tổng quan này là lớp nhìn mức cao; các nhóm thành phần bên trong nó được triển khai cụ thể bởi bộ sơ đồ chi tiết từ 1_4_1_1 đến 1_4_1_12."
    ),
]


def load_font(size: int, bold: bool = False):
    candidate = FONT_BOLD if bold else FONT_REGULAR
    if candidate.exists():
        return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def set_run_font(run, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if bold is not None:
        run.bold = bold


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_paragraph(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def add_labeled_paragraph(anchor: Paragraph, label: str, text: str) -> Paragraph:
    para = insert_paragraph_after(anchor, style="Normal")
    run1 = para.add_run(f"{label}: ")
    set_run_font(run1, bold=True)
    run2 = para.add_run(text)
    set_run_font(run2)
    return para


def escape_drawio_text(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "&lt;br&gt;")
    )


def node_center(node: Node) -> tuple[float, float]:
    return (node.x + node.w / 2, node.y + node.h / 2)


def write_drawio() -> None:
    lines = [
        '<mxfile host="Electron" agent="Codex">',
        '  <diagram id="143-model" name="1.4.3 Mô hình bài toán tổng quát">',
        f'    <mxGraphModel dx="{PAGE_WIDTH}" dy="{PAGE_HEIGHT}" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="{PAGE_WIDTH}" pageHeight="{PAGE_HEIGHT}" math="0" shadow="0">',
        "      <root>",
        '        <mxCell id="0" />',
        '        <mxCell id="1" parent="0" />',
    ]
    for node in NODES:
        lines.extend(
            [
                f'        <mxCell id="{node.node_id}" parent="1" style="{node.style}" value="{escape_drawio_text(node.text)}" vertex="1">',
                f'          <mxGeometry x="{node.x}" y="{node.y}" width="{node.w}" height="{node.h}" as="geometry" />',
                "        </mxCell>",
            ]
        )
    for edge in EDGES:
        dashed = ";dashed=1" if edge.dashed else ""
        lines.append(
            f'        <mxCell id="{edge.edge_id}" edge="1" parent="1" source="{edge.source}" target="{edge.target}" value="{escape_drawio_text(edge.label)}" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;endFill=1;strokeColor=#000000;fontColor=#000000;fontSize=12;fontFamily=Times New Roman{dashed}">'
        )
        lines.append('          <mxGeometry relative="1" as="geometry">')
        if edge.points:
            lines.append('            <Array as="points">')
            for x, y in edge.points:
                lines.append(f'              <mxPoint x="{x}" y="{y}" />')
            lines.append("            </Array>")
        lines.append("          </mxGeometry>")
        lines.append("        </mxCell>")
    lines.extend(["      </root>", "    </mxGraphModel>", "  </diagram>", "</mxfile>"])
    DRAWIO_PATH.parent.mkdir(parents=True, exist_ok=True)
    DRAWIO_PATH.write_text("\n".join(lines), encoding="utf-8")
    ET.parse(DRAWIO_PATH)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw_line in text.splitlines():
        words = raw_line.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            trial = f"{current} {word}"
            if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
                current = trial
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def draw_multiline_centered(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, max(40, x2 - x1 - 18))
    heights = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line or "Ag", font=font)
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + max(0, len(lines) - 1) * 6
    y = y1 + ((y2 - y1 - total_h) / 2)
    for line, height in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line or "Ag", font=font)
        width = bbox[2] - bbox[0]
        x = x1 + ((x2 - x1 - width) / 2)
        draw.text((x, y), line, font=font, fill="#000000")
        y += height + 6


def draw_multiline_top_centered(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font,
    top_padding: int = 18,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, max(40, x2 - x1 - 18))
    y = y1 + top_padding
    for line in lines:
        bbox = draw.textbbox((0, 0), line or "Ag", font=font)
        width = bbox[2] - bbox[0]
        height = bbox[3] - bbox[1]
        x = x1 + ((x2 - x1 - width) / 2)
        draw.text((x, y), line, font=font, fill="#000000")
        y += height + 6


def edge_points(source: Node, target: Node) -> list[tuple[float, float]]:
    sx, sy = node_center(source)
    tx, ty = node_center(target)
    dx = tx - sx
    dy = ty - sy
    if abs(dx) >= abs(dy):
        if dx >= 0:
            start = (source.x + source.w, sy)
            end = (target.x, ty)
        else:
            start = (source.x, sy)
            end = (target.x + target.w, ty)
        if abs(start[1] - end[1]) < 6:
            return [start, end]
        mid_x = (start[0] + end[0]) / 2
        return [start, (mid_x, start[1]), (mid_x, end[1]), end]
    if dy >= 0:
        start = (sx, source.y + source.h)
        end = (tx, target.y)
    else:
        start = (sx, source.y)
        end = (tx, target.y + target.h)
    if abs(start[0] - end[0]) < 6:
        return [start, end]
    mid_y = (start[1] + end[1]) / 2
    return [start, (start[0], mid_y), (end[0], mid_y), end]


def anchor_towards(node: Node, point: tuple[float, float]) -> tuple[float, float]:
    cx, cy = node_center(node)
    dx = point[0] - cx
    dy = point[1] - cy
    if abs(dx) >= abs(dy):
        return (node.x + node.w, cy) if dx >= 0 else (node.x, cy)
    return (cx, node.y + node.h) if dy >= 0 else (cx, node.y)


def full_edge_points(edge: Edge, nodes_by_id: dict[str, Node]) -> list[tuple[float, float]]:
    source = nodes_by_id[edge.source]
    target = nodes_by_id[edge.target]
    if not edge.points:
        return edge_points(source, target)
    start = anchor_towards(source, edge.points[0])
    end = anchor_towards(target, edge.points[-1])
    return [start, *edge.points, end]


def draw_polyline(draw: ImageDraw.ImageDraw, points: list[tuple[float, float]], dashed: bool) -> None:
    segments = list(zip(points, points[1:]))
    for start, end in segments:
        if dashed:
            segment_len = max(abs(end[0] - start[0]), abs(end[1] - start[1]))
            dash = 12
            gap = 8
            count = max(1, int(segment_len // (dash + gap)))
            for i in range(count + 1):
                t0 = min(1.0, ((dash + gap) * i) / max(1.0, segment_len))
                t1 = min(1.0, (((dash + gap) * i) + dash) / max(1.0, segment_len))
                p0 = (start[0] + (end[0] - start[0]) * t0, start[1] + (end[1] - start[1]) * t0)
                p1 = (start[0] + (end[0] - start[0]) * t1, start[1] + (end[1] - start[1]) * t1)
                draw.line([p0, p1], fill="#000000", width=3)
        else:
            draw.line([start, end], fill="#000000", width=3)


def label_position(points: list[tuple[float, float]]) -> tuple[float, float]:
    segments = list(zip(points, points[1:]))
    best = max(segments, key=lambda pair: abs(pair[1][0] - pair[0][0]) + abs(pair[1][1] - pair[0][1]))
    return ((best[0][0] + best[1][0]) / 2, (best[0][1] + best[1][1]) / 2)


def draw_arrow_head(draw: ImageDraw.ImageDraw, start: tuple[float, float], end: tuple[float, float]) -> None:
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_len = 12
    left = (
        end[0] - arrow_len * math.cos(angle - math.pi / 6),
        end[1] - arrow_len * math.sin(angle - math.pi / 6),
    )
    right = (
        end[0] - arrow_len * math.cos(angle + math.pi / 6),
        end[1] - arrow_len * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill="#000000")


def render_png() -> None:
    scale = 1.2
    margin = 30
    canvas_w = int((PAGE_WIDTH + margin * 2) * scale)
    canvas_h = int((PAGE_HEIGHT + margin * 2) * scale)
    image = Image.new("RGB", (canvas_w, canvas_h), "#ffffff")
    draw = ImageDraw.Draw(image)

    title_font = load_font(18, bold=True)
    detail_font = load_font(16, bold=False)
    label_font = load_font(12, bold=False)
    nodes_by_id = {node.node_id: node for node in NODES}

    def scaled_point(point: tuple[float, float]) -> tuple[float, float]:
        return ((point[0] + margin) * scale, (point[1] + margin) * scale)

    for edge in EDGES:
        raw_points = full_edge_points(edge, nodes_by_id)
        points = [scaled_point(point) for point in raw_points]
        draw_polyline(draw, points, edge.dashed)
        draw_arrow_head(draw, points[-2], points[-1])
        if edge.label:
            mx, my = scaled_point(label_position(raw_points))
            bbox = draw.textbbox((0, 0), edge.label, font=label_font)
            label_box = (
                int(mx - ((bbox[2] - bbox[0]) / 2) - 8),
                int(my - ((bbox[3] - bbox[1]) / 2) - 3),
                int(mx + ((bbox[2] - bbox[0]) / 2) + 8),
                int(my + ((bbox[3] - bbox[1]) / 2) + 3),
            )
            draw.rounded_rectangle(label_box, radius=8, fill="#ffffff", outline="#d0d0d0")
            draw.text((label_box[0] + 8, label_box[1] + 1), edge.label, font=label_font, fill="#000000")

    for node in NODES:
        x1 = int((node.x + margin) * scale)
        y1 = int((node.y + margin) * scale)
        x2 = int((node.x + node.w + margin) * scale)
        y2 = int((node.y + node.h + margin) * scale)
        radius = 18 if node.node_id.startswith("n") else 24
        draw.rounded_rectangle((x1, y1, x2, y2), radius=radius, outline="#000000", fill="#ffffff", width=4)
        if node.node_id.startswith("g"):
            draw_multiline_top_centered(draw, (x1, y1, x2, y2), node.text, title_font, top_padding=18)
        else:
            draw_multiline_centered(draw, (x1, y1, x2, y2), node.text, detail_font)

    PNG_PATH.parent.mkdir(parents=True, exist_ok=True)
    image.save(PNG_PATH, format="PNG")


def replace_section_143_in_docx() -> None:
    doc = Document(SOURCE_DOCX_PATH)
    start_para = find_paragraph(doc, "1.4.3. Mô hình bài toán tổng quát")
    end_para = find_paragraph(doc, "1.4.4. Mô tả hoạt động của hệ thống")

    paragraphs = doc.paragraphs
    start_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == start_para.text.strip())
    end_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == end_para.text.strip() and i > start_idx)
    for paragraph in list(doc.paragraphs[start_idx + 1:end_idx]):
        remove_paragraph(paragraph)

    anchor = start_para
    ref_para = insert_paragraph_after(anchor, style="Normal")
    run1 = ref_para.add_run("Tệp sơ đồ nguồn: ")
    set_run_font(run1, bold=True)
    run2 = ref_para.add_run("1_4_3_mo_hinh_bai_toan_tong_quat.drawio - Sơ đồ mô hình bài toán tổng quát.")
    set_run_font(run2)
    anchor = ref_para

    image_para = insert_paragraph_after(anchor, style="Normal")
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_run = image_para.add_run()
    image_run.add_picture(str(PNG_PATH), width=Inches(6.3))
    anchor = image_para

    caption_para = insert_paragraph_after(anchor, style="Normal")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = caption_para.add_run("Hình 1.4.3: Sơ đồ mô hình bài toán tổng quát của hệ thống.")
    set_run_font(cap_run)
    cap_run.italic = True
    anchor = caption_para

    for label, text in SECTION_143_PARAGRAPHS:
        anchor = add_labeled_paragraph(anchor, label, text)

    doc.save(TARGET_DOCX_PATH)


def main() -> None:
    write_drawio()
    render_png()
    replace_section_143_in_docx()
    print(DRAWIO_PATH)
    print(PNG_PATH)
    print(TARGET_DOCX_PATH)


if __name__ == "__main__":
    main()
