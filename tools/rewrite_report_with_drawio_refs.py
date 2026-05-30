from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_restructured.docx"
TARGET_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_drawio_aligned.docx"


SECTION_14_FILES = [
    ("1.4.1.1. Chức năng định danh người chơi", "1_4_1_1_dinh_danh_nguoi_choi.drawio", "Sơ đồ định danh người chơi qua PlayFab"),
    ("1.4.1.2. Chức năng bắt đầu trận đấu", "1_4_1_2_bat_dau_tran_dau.drawio", "Sơ đồ bắt đầu trận đấu"),
    ("1.4.1.3. Chức năng điều khiển nhân vật", "1_4_1_3_player_controller.drawio", "Sơ đồ điều khiển nhân vật"),
    ("1.4.1.4. Chức năng chiến đấu của người chơi", "1_4_1_4_chien_dau_player.drawio", "Sơ đồ chiến đấu của người chơi"),
    ("1.4.1.5. Chức năng quản lý enemy", "1_4_1_5_enemy_system.drawio", "Sơ đồ nhóm enemy"),
    ("1.4.1.6. Chức năng quản lý wave và độ khó", "1_4_1_6_wave_va_do_kho.drawio", "Sơ đồ quản lý wave và độ khó"),
    ("1.4.1.7. Chức năng đổi theme bản đồ", "1_4_1_7_theme_ban_do.drawio", "Sơ đồ đổi theme bản đồ"),
    ("1.4.1.8. Chức năng kinh nghiệm và lên cấp", "1_4_1_8_exp_va_len_cap.drawio", "Sơ đồ EXP và lên cấp"),
    ("1.4.1.9. Chức năng chọn buff và tăng sức mạnh", "1_4_1_9_buff_va_tang_suc_manh.drawio", "Sơ đồ chọn buff và tăng sức mạnh"),
    ("1.4.1.10. Chức năng giao diện trong trận", "1_4_1_10_ui_trong_tran.drawio", "Sơ đồ UI trong trận"),
    ("1.4.1.11. Chức năng tạm dừng và kết thúc trận", "1_4_1_11_pause_va_ket_thuc_tran.drawio", "Sơ đồ pause và kết thúc trận"),
    ("1.4.1.12. Chức năng bảng xếp hạng", "1_4_1_12_leaderboard.drawio", "Sơ đồ leaderboard"),
]


def set_run_font(run, size: int = 13, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph: Paragraph, size: int = 13) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_paragraph(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def remove_range(document: Document, start_text: str, end_text: str, include_start: bool = True) -> None:
    removing = False
    to_remove: list[Paragraph] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == start_text:
            removing = True
            if include_start:
                to_remove.append(paragraph)
            continue
        if removing and text == end_text:
            break
        if removing:
            to_remove.append(paragraph)
    for paragraph in to_remove:
        remove_paragraph(paragraph)


def add_labeled_paragraph(anchor: Paragraph, label: str, text: str) -> Paragraph:
    para = insert_paragraph_after(anchor, style="Normal")
    run1 = para.add_run(f"{label}: ")
    set_run_font(run1, bold=True)
    run2 = para.add_run(text)
    set_run_font(run2)
    return para


def add_file_reference(anchor: Paragraph, filename: str, caption: str) -> Paragraph:
    para = insert_paragraph_after(anchor, style="Normal")
    run1 = para.add_run("Tệp sơ đồ nguồn: ")
    set_run_font(run1, bold=True)
    run2 = para.add_run(f"{filename} - {caption}.")
    set_run_font(run2)
    return para


def insert_section_14_file_refs(document: Document) -> None:
    for heading, filename, caption in reversed(SECTION_14_FILES):
        anchor = find_paragraph(document, heading)
        add_file_reference(anchor, filename, caption)


def rebuild_section_143_144(document: Document) -> None:
    remove_range(document, "1.4.3. Mô hình bài toán tổng quát", "CHƯƠNG 2 \nMỘT SỐ KIẾN THỨC CƠ BẢN THỰC HIỆN ĐỀ TÀI")
    anchor = find_paragraph(document, "Ổn định luồng chơi: Các trạng thái nhập tên, pause, chọn buff, leaderboard và đổi theme không được để lại input cũ hoặc phá vỡ trạng thái wave hiện hành.")

    anchor = insert_paragraph_after(anchor, "1.4.3. Mô hình bài toán tổng quát", style="Heading 3")
    anchor = add_labeled_paragraph(
        anchor,
        "Định hướng mô hình",
        "Theo bộ sơ đồ mới, mô hình bài toán không còn trình bày bằng một hình lớn duy nhất mà được tách thành các cụm chức năng độc lập. Cách tiếp cận này giúp người đọc theo dõi từng dòng dữ liệu và từng khối xử lý của hệ thống mà không phải giải mã một sơ đồ tổng hợp quá dày đặc."
    )
    anchor = add_labeled_paragraph(
        anchor,
        "Các nhóm chức năng chính",
        "Mô hình tổng quát của game được chia thành 12 nhóm: định danh người chơi, bắt đầu trận, điều khiển player, chiến đấu của player, enemy, wave, đổi theme bản đồ, EXP và level, buff, UI trong trận, pause/kết thúc trận và leaderboard. Mỗi nhóm được ánh xạ sang một file draw.io riêng để có thể chỉnh sửa trực tiếp mà không ảnh hưởng các phần còn lại."
    )
    anchor = add_labeled_paragraph(
        anchor,
        "Ý nghĩa trình bày",
        "Việc tách sơ đồ theo nhóm tính năng giúp báo cáo nhất quán với cấu trúc triển khai thực tế của project. Khi cần điều chỉnh logic hoặc bổ sung ghi chú, người viết chỉ cần chỉnh đúng file sơ đồ của nhóm liên quan thay vì sửa toàn bộ hình tổng hợp."
    )
    for _, filename, caption in SECTION_14_FILES:
        anchor = add_labeled_paragraph(anchor, "Sơ đồ chi tiết", f"{filename} - {caption}")

    anchor = insert_paragraph_after(anchor, "1.4.4. Mô tả hoạt động của hệ thống", style="Heading 3")
    anchor = add_labeled_paragraph(
        anchor,
        "Nguyên tắc đọc luồng",
        "Luồng hệ thống giờ được mô tả theo các sơ đồ nguồn chuyên biệt. Thay vì gom toàn bộ trạng thái vào một hình duy nhất, báo cáo trình bày quá trình vận hành bằng cách đi qua từng nhóm chức năng theo thứ tự xảy ra trong trận."
    )
    anchor = add_labeled_paragraph(
        anchor,
        "Chuỗi hoạt động chính",
        "Hệ thống bắt đầu từ định danh PlayFab, chuyển sang pha chờ trong scene, mở trận qua ChallengePostNPC, sau đó vào vòng lặp spawn wave - chiến đấu - cộng EXP - lên cấp - chọn buff - tiếp tục wave. Khi người chơi bị hạ, PlayerHealth tổng hợp totalExpGained để gửi HighScore và mở LeaderboardPanel."
    )
    anchor = add_labeled_paragraph(
        anchor,
        "Ánh xạ tới sơ đồ nguồn",
        "Phần định danh và chuẩn bị vào trận tương ứng với 1_4_1_1 và 1_4_1_2. Phần điều khiển và chiến đấu của player tương ứng với 1_4_1_3 và 1_4_1_4. Phần enemy, wave, theme, progression và buff tương ứng với 1_4_1_5 đến 1_4_1_9. Lớp UI, pause và leaderboard tương ứng với 1_4_1_10 đến 1_4_1_12."
    )
    anchor = add_labeled_paragraph(
        anchor,
        "Lợi ích khi viết báo cáo",
        "Cách viết theo bộ sơ đồ mới làm rõ dữ liệu vào/ra, điểm kích hoạt và điểm kết thúc của từng khối. Điều này phù hợp với yêu cầu phân tích hệ thống vì mỗi chức năng giờ có thể được mô tả, sửa sơ đồ và kiểm tra chéo độc lập."
    )


CH3_BLOCKS = [
    ("3.1. Use Case Diagram", [
        ("3.1.1. Use Case người chơi trước trận", "3_1_1_use_case_nguoi_choi_truoc_tran.drawio", [
            "Nhóm use case trước trận tập trung vào các tương tác xảy ra trước khi wave đầu tiên được kích hoạt. Trọng tâm ở đây là việc người chơi đăng nhập bằng Custom ID, kiểm tra hồ sơ PlayFab, nhập Display Name nếu cần và khởi động phiên chơi thông qua ChallengePostNPC.",
            "Việc tách riêng sơ đồ này giúp báo cáo nhấn mạnh rằng phần định danh và phần bắt đầu trận là một cụm nghiệp vụ độc lập, có actor riêng, dữ liệu riêng và có phụ thuộc trực tiếp vào dịch vụ PlayFab. Cách trình bày này rõ ràng hơn so với việc nhồi toàn bộ use case vào một hình lớn."
        ]),
        ("3.1.2. Use Case gameplay core loop", "3_1_2_use_case_gameplay_core_loop.drawio", [
            "Nhóm use case gameplay core loop mô tả toàn bộ vòng lặp sinh tồn mà người chơi trực tiếp tham gia: điều khiển nhân vật, chiến đấu tự động, tiêu diệt enemy, nhận EXP, lên cấp, chọn buff, hoàn thành wave và đối mặt với boss wave hoặc đổi theme bản đồ.",
            "Sơ đồ này được tách riêng vì đây là phần tạo ra giá trị chơi lại và cũng là nơi hội tụ các hệ thống Player, Enemy, Wave và Buff. Khi viết theo sơ đồ mới, báo cáo có thể mô tả rõ từng use case phụ mà không làm mờ ranh giới giữa gameplay runtime với backend."
        ]),
        ("3.1.3. Use Case kết thúc trận và leaderboard", "3_1_3_use_case_ket_thuc_tran_leaderboard.drawio", [
            "Nhóm use case cuối tập trung vào các hành vi tạm dừng, kết thúc trận và đồng bộ kết quả. Sau khi người chơi chết hoặc chủ động rời lượt chơi, hệ thống tổng hợp điểm từ totalExpGained, gửi lên PlayFab và hiển thị lại bảng xếp hạng trong LeaderboardPanel.",
            "Việc dành một sơ đồ riêng cho giai đoạn này giúp báo cáo làm rõ phần hậu xử lý trận đấu không chỉ là một bước UI, mà là một cụm nghiệp vụ có liên quan đến lưu trữ trực tuyến, phản hồi kết quả và so sánh thành tích giữa nhiều người chơi."
        ]),
    ]),
    ("3.2. Sequence Diagram", [
        ("3.2.1. Luồng đăng nhập và nhập tên hiển thị", "3_2_1_sequence_dang_nhap_nhap_ten.drawio", [
            "Sequence đầu tiên mô tả chuỗi xác thực PlayFab và thiết lập Display Name. Player kích hoạt quá trình đăng nhập, PlayFabLeaderboardManager gửi LoginWithCustomID, đọc hồ sơ người chơi rồi quyết định có cần mở NameInputPanel hay không.",
            "Việc tách sequence này thành file riêng giúp thể hiện rõ vai trò trung gian của PlayFabLeaderboardManager, đồng thời làm nổi bật dữ liệu quan trọng của luồng là Custom ID, CurrentPlayFabId và CurrentDisplayName."
        ]),
        ("3.2.2. Luồng chiến đấu, nhận EXP và chọn buff", "3_2_2_sequence_chien_dau_exp_buff.drawio", [
            "Sequence thứ hai tập trung vào mạch chiến đấu cốt lõi: PlayerAttack sinh projectile qua ObjectPool, Enemy bị hạ, ExpDropper cộng expValue cho PlayerLevelSystem và OnLevelUp kích hoạt BuffCardManager cùng CardSelectionPanel.",
            "Theo sơ đồ mới, luồng này được tách riêng để giữ cho báo cáo bám sát quan hệ giữa combat runtime và progression. Cách trình bày như vậy cho phép mô tả rõ attackCooldown, expToNextLevel và cardsPerSelection=3 như các tham số quyết định nhịp tăng trưởng của trận."
        ]),
        ("3.2.3. Luồng game over, gửi điểm và leaderboard", "3_2_3_sequence_game_over_gui_diem_tai_leaderboard.drawio", [
            "Sequence cuối cùng bắt đầu tại thời điểm PlayerHealth xác nhận nhân vật đã chết. Hệ thống lấy totalExpGained từ PlayerLevelSystem, quy đổi thành điểm cuối trận, gọi SubmitScore tại PlayFabLeaderboardManager rồi tải lại leaderboard tổng và leaderboard quanh người chơi hiện tại.",
            "Nhờ có một sequence riêng, báo cáo làm rõ rằng kết thúc trận là giao điểm giữa gameplay cục bộ và backend. Điểm số ở đây bám đúng thống kê HighScore, không dùng thời gian sống hay số quái tiêu diệt."
        ]),
    ]),
    ("3.3. Activity Diagram", [
        ("3.3.1. Activity gameplay core loop", "3_3_activity_gameplay_core_loop.drawio", [
            "Activity Diagram được giữ thành một file riêng nhưng được vẽ lại theo các lane và các nhánh điều kiện chính thay vì một trục dài. Luồng bắt đầu từ đăng nhập PlayFab, kiểm tra Display Name, tương tác với NPC để vào trận, sau đó lặp qua các bước spawn wave, chiến đấu, enemy chết, cộng EXP và chọn buff.",
            "Ở nửa sau của activity, hệ thống rẽ nhánh sang hoàn tất wave, boss wave, đổi theme hoặc kết thúc trận tùy điều kiện. Viết lại báo cáo theo activity mới giúp người đọc hiểu được đâu là điểm lặp, đâu là điều kiện rẽ nhánh và đâu là điều kiện dừng hoàn toàn của một phiên chơi."
        ]),
    ]),
    ("3.4. Component Diagram", [
        ("3.4.1. Component gameplay runtime", "3_4_1_component_gameplay_runtime.drawio", [
            "Component gameplay runtime gom các thành phần trực tiếp vận hành trong lúc trận đấu đang diễn ra, bao gồm Input, Player, Enemy, Wave, ObjectPool, MapThemeManager và các shared services. Sơ đồ này nhấn mạnh quan hệ giữa spawn quái, chiến đấu và tái sử dụng đối tượng runtime.",
            "Việc tách runtime thành một component riêng giúp báo cáo tập trung vào vòng lặp chiến đấu, tránh làm rối phần trình bày bởi các dịch vụ giao diện hoặc backend."
        ]),
        ("3.4.2. Component progression và UI", "3_4_2_component_progression_ui.drawio", [
            "Component progression và UI mô tả PlayerLevelSystem, BuffCardManager, CardSelectionPanel và GameUI cùng các panel phụ trợ. Đây là cụm chịu trách nhiệm phản hồi trạng thái trận đấu, hiển thị HUD và cho phép người chơi đưa ra quyết định nâng cấp.",
            "Sơ đồ riêng cho progression/UI giúp báo cáo làm rõ quan hệ giữa event runtime với màn hình hiển thị, đặc biệt là cách OnLevelUp đi từ PlayerLevelSystem sang BuffCardManager rồi tới CardSelectionPanel."
        ]),
        ("3.4.3. Component backend và services", "3_4_3_component_backend_services.drawio", [
            "Component backend và services tập trung vào các thành phần giao tiếp với dịch vụ ngoài và các tiện ích nền như PlayFabLeaderboardManager, PlayFab, NameInputPanel, LeaderboardPanel, LoadingUIManager và AudioManager. Cụm này không tham gia trực tiếp vào tính toán combat nhưng quyết định tính hoàn thiện của trải nghiệm.",
            "Việc tách phần backend/services thành sơ đồ riêng giúp người đọc dễ phân biệt dịch vụ cốt lõi của gameplay với các dịch vụ hỗ trợ xác thực, chuyển cảnh, âm thanh và hiển thị kết quả."
        ]),
    ]),
    ("3.5. Class Diagram", [
        ("3.5.1. Nhóm Player và Progression", "3_5_1_class_player_progression.drawio", [
            "Sơ đồ lớp Player và Progression phản ánh mối quan hệ giữa PlayerController, PlayerAttack, PlayerHealth, PlayerData, PlayerLevelSystem, BuffCardManager, PlayerStatsPanel và CardSelectionPanel. Đây là cụm lớp quan trọng nhất vì nó nối liền điều khiển, chiến đấu, phát triển sức mạnh và hiển thị trạng thái.",
            "Khi viết theo sơ đồ mới, báo cáo có thể mô tả rõ PlayerData là nơi giữ tham số runtime, PlayerLevelSystem là nơi phát sự kiện tăng cấp, còn BuffCardManager và CardSelectionPanel là lớp quyết định cách sức mạnh của người chơi được điều chỉnh trong từng lượt chơi."
        ]),
        ("3.5.2. Nhóm Enemy và Projectile", "3_5_2_class_enemy_projectile.drawio", [
            "Sơ đồ lớp Enemy và Projectile mô tả tuyến kế thừa từ Enemy tới MeleeEnemy, RangedEnemy, FlyEnemy và BossEnemy, đồng thời giữ mối liên hệ giữa EnemyData, EnemyConfig và nhánh Projectile. IDamageable đóng vai trò hợp nhất hành vi nhận sát thương cho các đối tượng có thể bị tấn công.",
            "Cách tách mới giúp báo cáo trình bày tốt hơn hai tầng của hệ thống đối kháng: tầng dữ liệu chỉ số và tầng hành vi runtime. Điều này đặc biệt hữu ích khi phân tích vì project hiện không dùng NavMeshAgent mà điều khiển AI bằng CharacterController và logic lớp Enemy."
        ]),
        ("3.5.3. Nhóm UI, NPC và Backend", "3_5_3_class_ui_backend.drawio", [
            "Sơ đồ lớp UI, NPC và Backend tập trung vào PanelBase, các panel con, GameUI, lớp cha NPC và hai lớp con ChallengePostNPC cùng ChestBuffBox, kết hợp với PlayFabLeaderboardManager ở tầng backend. Nhóm này cho thấy project không chỉ có gameplay mà còn có lớp điều hướng, nhập liệu và đồng bộ trực tuyến tương đối hoàn chỉnh.",
            "Viết lại theo sơ đồ mới cho phép nhấn mạnh GameUI là hub giao diện, NPC là điểm khởi tạo các luồng tương tác trong scene, còn PlayFabLeaderboardManager là điểm hội tụ của xác thực, gửi điểm và tải bảng xếp hạng."
        ]),
    ]),
]


def rebuild_chapter_3(document: Document) -> None:
    remove_range(document, "3.1. Use Case Diagram", "CHƯƠNG 4 \nXÂY DỰNG CHƯƠNG TRÌNH", include_start=True)
    chapter_anchor = find_paragraph(document, "CHƯƠNG 3 \nPHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")

    anchor = chapter_anchor
    intro = insert_paragraph_after(anchor, style="Normal")
    intro.add_run(
        "Chương 3 được viết lại theo bộ sơ đồ draw.io mới, trong đó mỗi nhóm tính năng hoặc mỗi luồng nghiệp vụ được tách thành một file nguồn riêng. Cách tổ chức này giúp báo cáo bám sát cấu trúc hệ thống thực tế và thuận lợi hơn cho việc chỉnh sửa sơ đồ sau này."
    )
    set_paragraph_font(intro)
    anchor = intro

    for section_heading, items in CH3_BLOCKS:
        anchor = insert_paragraph_after(anchor, section_heading, style="Heading 2")
        for heading, filename, paragraphs in items:
            anchor = insert_paragraph_after(anchor, heading, style="Heading 3")
            anchor = add_file_reference(anchor, filename, heading)
            for text in paragraphs:
                para = insert_paragraph_after(anchor, style="Normal")
                para.add_run(text)
                set_paragraph_font(para)
                anchor = para


def main() -> None:
    shutil.copyfile(SOURCE_DOC, TARGET_DOC)
    doc = Document(TARGET_DOC)
    insert_section_14_file_refs(doc)
    rebuild_section_143_144(doc)
    rebuild_chapter_3(doc)
    doc.save(TARGET_DOC)
    print(TARGET_DOC)


if __name__ == "__main__":
    main()
