from __future__ import annotations

import math
import os
from pathlib import Path
from typing import Iterable, List, Sequence, Tuple

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIAGRAM_DIR = ROOT / "output" / "diagrams" / "ch3"
OUTPUT_DOC_DIR = ROOT / "output" / "doc"

FONT_REGULAR = r"C:\Windows\Fonts\times.ttf"
FONT_BOLD = r"C:\Windows\Fonts\timesbd.ttf"
FONT_ITALIC = r"C:\Windows\Fonts\timesi.ttf"


def ensure_dirs() -> None:
    OUTPUT_DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DOC_DIR.mkdir(parents=True, exist_ok=True)


def wrap_text(text: str, font: ImageFont.FreeTypeFont, max_width: int) -> List[str]:
    if not text:
        return [""]

    lines: List[str] = []
    for raw_line in text.split("\n"):
        words = raw_line.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}"
            if font.getlength(candidate) <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def text_block_size(lines: Sequence[str], font: ImageFont.FreeTypeFont, line_spacing: int = 6) -> Tuple[int, int]:
    widths = []
    total_height = 0
    for line in lines:
        bbox = font.getbbox(line or "A")
        widths.append(int(math.ceil(font.getlength(line))) if line else bbox[2] - bbox[0])
        total_height += bbox[3] - bbox[1]
    if lines:
        total_height += line_spacing * (len(lines) - 1)
    return max(widths) if widths else 0, total_height


class DiagramPNG:
    def __init__(self, width: int, height: int) -> None:
        self.width = width
        self.height = height
        self.image = Image.new("RGB", (width, height), "white")
        self.draw = ImageDraw.Draw(self.image)

    def font(self, size: int, bold: bool = False, italic: bool = False) -> ImageFont.FreeTypeFont:
        path = FONT_REGULAR
        if bold:
            path = FONT_BOLD
        elif italic:
            path = FONT_ITALIC
        return ImageFont.truetype(path, size=size)

    def title(self, text: str) -> None:
        font = self.font(42, bold=True)
        bbox = self.draw.textbbox((0, 0), text, font=font)
        x = (self.width - (bbox[2] - bbox[0])) / 2
        self.draw.text((x, 28), text, font=font, fill="black")

    def text(
        self,
        x: int,
        y: int,
        text: str,
        font_size: int = 28,
        max_width: int | None = None,
        align: str = "center",
        bold: bool = False,
        italic: bool = False,
        fill: str = "black",
        line_spacing: int = 6,
    ) -> Tuple[int, int]:
        font = self.font(font_size, bold=bold, italic=italic)
        lines = wrap_text(text, font, max_width) if max_width else text.split("\n")
        width, height = text_block_size(lines, font, line_spacing)
        start_y = y
        for line in lines:
            line_width = int(math.ceil(font.getlength(line))) if line else 0
            line_bbox = font.getbbox(line or "A")
            line_height = line_bbox[3] - line_bbox[1]
            if align == "center":
                tx = x - line_width / 2
            elif align == "right":
                tx = x - line_width
            else:
                tx = x
            self.draw.text((tx, start_y), line, font=font, fill=fill)
            start_y += line_height + line_spacing
        return width, height

    def box(
        self,
        x: int,
        y: int,
        w: int,
        h: int,
        text: str,
        *,
        fill: str = "#f6f6f6",
        outline: str = "black",
        radius: int = 22,
        font_size: int = 28,
        bold: bool = False,
    ) -> None:
        self.draw.rounded_rectangle((x, y, x + w, y + h), radius=radius, outline=outline, width=4, fill=fill)
        self.text(x + w // 2, y + h // 2 - 14, text, font_size=font_size, max_width=w - 24, align="center", bold=bold)

    def plain_box(
        self,
        x: int,
        y: int,
        w: int,
        h: int,
        text: str,
        *,
        fill: str = "#ffffff",
        outline: str = "black",
        font_size: int = 28,
        bold: bool = False,
    ) -> None:
        self.draw.rectangle((x, y, x + w, y + h), outline=outline, width=4, fill=fill)
        self.text(x + w // 2, y + h // 2 - 14, text, font_size=font_size, max_width=w - 24, align="center", bold=bold)

    def ellipse(
        self,
        x: int,
        y: int,
        w: int,
        h: int,
        text: str,
        *,
        fill: str = "#ffffff",
        outline: str = "black",
        font_size: int = 26,
    ) -> None:
        self.draw.ellipse((x, y, x + w, y + h), outline=outline, width=4, fill=fill)
        self.text(x + w // 2, y + h // 2 - 14, text, font_size=font_size, max_width=w - 28, align="center")

    def diamond(
        self,
        x: int,
        y: int,
        w: int,
        h: int,
        text: str,
        *,
        fill: str = "#ffffff",
        outline: str = "black",
        font_size: int = 26,
    ) -> None:
        points = [(x + w // 2, y), (x + w, y + h // 2), (x + w // 2, y + h), (x, y + h // 2)]
        self.draw.polygon(points, outline=outline, fill=fill)
        self.draw.line(points + [points[0]], fill=outline, width=4)
        self.text(x + w // 2, y + h // 2 - 14, text, font_size=font_size, max_width=w - 34, align="center")

    def actor(self, x: int, y: int, label: str) -> None:
        self.draw.ellipse((x + 25, y, x + 75, y + 50), outline="black", width=4, fill="white")
        self.draw.line((x + 50, y + 50, x + 50, y + 122), fill="black", width=4)
        self.draw.line((x + 12, y + 76, x + 88, y + 76), fill="black", width=4)
        self.draw.line((x + 50, y + 122, x + 18, y + 166), fill="black", width=4)
        self.draw.line((x + 50, y + 122, x + 82, y + 166), fill="black", width=4)
        self.text(x + 50, y + 182, label, font_size=28, max_width=160, align="center", bold=True)

    def line(self, x1: int, y1: int, x2: int, y2: int, *, dashed: bool = False, width: int = 4, fill: str = "black") -> None:
        if dashed:
            dash = 14
            gap = 10
            length = math.hypot(x2 - x1, y2 - y1)
            if length == 0:
                return
            dx = (x2 - x1) / length
            dy = (y2 - y1) / length
            pos = 0.0
            while pos < length:
                start = pos
                end = min(pos + dash, length)
                sx = x1 + dx * start
                sy = y1 + dy * start
                ex = x1 + dx * end
                ey = y1 + dy * end
                self.draw.line((sx, sy, ex, ey), fill=fill, width=width)
                pos += dash + gap
        else:
            self.draw.line((x1, y1, x2, y2), fill=fill, width=width)

    def arrow(
        self,
        x1: int,
        y1: int,
        x2: int,
        y2: int,
        *,
        dashed: bool = False,
        label: str | None = None,
        label_dx: int = 0,
        label_dy: int = -32,
    ) -> None:
        self.line(x1, y1, x2, y2, dashed=dashed)
        angle = math.atan2(y2 - y1, x2 - x1)
        size = 16
        p1 = (x2, y2)
        p2 = (x2 - size * math.cos(angle - math.pi / 6), y2 - size * math.sin(angle - math.pi / 6))
        p3 = (x2 - size * math.cos(angle + math.pi / 6), y2 - size * math.sin(angle + math.pi / 6))
        self.draw.polygon([p1, p2, p3], fill="black")
        if label:
            mx = (x1 + x2) // 2 + label_dx
            my = (y1 + y2) // 2 + label_dy
            self.text(mx, my, label, font_size=24, max_width=300, align="center")

    def save(self, path: Path) -> None:
        self.image.save(path)


class DiagramSVG:
    def __init__(self, width: int, height: int) -> None:
        self.width = width
        self.height = height
        self.elements: List[str] = []

    def add(self, raw: str) -> None:
        self.elements.append(raw)

    def _text(
        self,
        x: int,
        y: int,
        text: str,
        *,
        font_size: int = 28,
        max_width: int | None = None,
        anchor: str = "middle",
        bold: bool = False,
        italic: bool = False,
        line_spacing: int = 6,
        font_measure: ImageFont.FreeTypeFont | None = None,
    ) -> None:
        if font_measure is None:
            path = FONT_BOLD if bold else FONT_ITALIC if italic else FONT_REGULAR
            font_measure = ImageFont.truetype(path, size=font_size)
        lines = wrap_text(text, font_measure, max_width) if max_width else text.split("\n")
        weight = "bold" if bold else "normal"
        style = "italic" if italic else "normal"
        bbox = font_measure.getbbox("Ag")
        line_h = bbox[3] - bbox[1]
        total_h = line_h * len(lines) + line_spacing * max(0, len(lines) - 1)
        start_y = y - total_h / 2 + line_h
        self.add(
            f"<text x='{x}' y='{start_y}' text-anchor='{anchor}' font-family='Times New Roman' "
            f"font-size='{font_size}' font-weight='{weight}' font-style='{style}'>"
        )
        current_y = start_y
        for idx, line in enumerate(lines):
            dy = 0 if idx == 0 else line_h + line_spacing
            self.add(f"<tspan x='{x}' dy='{dy}'>{escape_xml(line)}</tspan>")
        self.add("</text>")

    def title(self, text: str) -> None:
        font = ImageFont.truetype(FONT_BOLD, size=42)
        self._text(self.width // 2, 58, text, font_size=42, bold=True, font_measure=font)

    def box(self, x: int, y: int, w: int, h: int, text: str, *, fill: str = "#f6f6f6", radius: int = 22, font_size: int = 28, bold: bool = False) -> None:
        self.add(f"<rect x='{x}' y='{y}' width='{w}' height='{h}' rx='{radius}' ry='{radius}' fill='{fill}' stroke='black' stroke-width='4'/>")
        font = ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size=font_size)
        self._text(x + w // 2, y + h // 2 + 4, text, font_size=font_size, max_width=w - 24, bold=bold, font_measure=font)

    def plain_box(self, x: int, y: int, w: int, h: int, text: str, *, fill: str = "#ffffff", font_size: int = 28, bold: bool = False) -> None:
        self.add(f"<rect x='{x}' y='{y}' width='{w}' height='{h}' fill='{fill}' stroke='black' stroke-width='4'/>")
        font = ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size=font_size)
        self._text(x + w // 2, y + h // 2 + 4, text, font_size=font_size, max_width=w - 24, bold=bold, font_measure=font)

    def ellipse(self, x: int, y: int, w: int, h: int, text: str, *, fill: str = "#ffffff", font_size: int = 26) -> None:
        self.add(f"<ellipse cx='{x + w / 2}' cy='{y + h / 2}' rx='{w / 2}' ry='{h / 2}' fill='{fill}' stroke='black' stroke-width='4'/>")
        font = ImageFont.truetype(FONT_REGULAR, size=font_size)
        self._text(x + w // 2, y + h // 2 + 4, text, font_size=font_size, max_width=w - 28, font_measure=font)

    def diamond(self, x: int, y: int, w: int, h: int, text: str, *, fill: str = "#ffffff", font_size: int = 26) -> None:
        points = f"{x + w // 2},{y} {x + w},{y + h // 2} {x + w // 2},{y + h} {x},{y + h // 2}"
        self.add(f"<polygon points='{points}' fill='{fill}' stroke='black' stroke-width='4'/>")
        font = ImageFont.truetype(FONT_REGULAR, size=font_size)
        self._text(x + w // 2, y + h // 2 + 4, text, font_size=font_size, max_width=w - 34, font_measure=font)

    def actor(self, x: int, y: int, label: str) -> None:
        self.add(f"<circle cx='{x + 50}' cy='{y + 25}' r='25' fill='white' stroke='black' stroke-width='4'/>")
        self.add(f"<line x1='{x + 50}' y1='{y + 50}' x2='{x + 50}' y2='{y + 122}' stroke='black' stroke-width='4'/>")
        self.add(f"<line x1='{x + 12}' y1='{y + 76}' x2='{x + 88}' y2='{y + 76}' stroke='black' stroke-width='4'/>")
        self.add(f"<line x1='{x + 50}' y1='{y + 122}' x2='{x + 18}' y2='{y + 166}' stroke='black' stroke-width='4'/>")
        self.add(f"<line x1='{x + 50}' y1='{y + 122}' x2='{x + 82}' y2='{y + 166}' stroke='black' stroke-width='4'/>")
        font = ImageFont.truetype(FONT_BOLD, size=28)
        self._text(x + 50, y + 204, label, font_size=28, max_width=160, bold=True, font_measure=font)

    def line(self, x1: int, y1: int, x2: int, y2: int, *, dashed: bool = False, width: int = 4) -> None:
        dash = " stroke-dasharray='14 10'" if dashed else ""
        self.add(f"<line x1='{x1}' y1='{y1}' x2='{x2}' y2='{y2}' stroke='black' stroke-width='{width}'{dash}/>")

    def arrow(self, x1: int, y1: int, x2: int, y2: int, *, dashed: bool = False, label: str | None = None, label_dx: int = 0, label_dy: int = -32) -> None:
        self.line(x1, y1, x2, y2, dashed=dashed)
        angle = math.atan2(y2 - y1, x2 - x1)
        size = 16
        p2 = (x2 - size * math.cos(angle - math.pi / 6), y2 - size * math.sin(angle - math.pi / 6))
        p3 = (x2 - size * math.cos(angle + math.pi / 6), y2 - size * math.sin(angle + math.pi / 6))
        self.add(f"<polygon points='{x2},{y2} {p2[0]},{p2[1]} {p3[0]},{p3[1]}' fill='black'/>")
        if label:
            font = ImageFont.truetype(FONT_REGULAR, size=24)
            self._text((x1 + x2) // 2 + label_dx, (y1 + y2) // 2 + label_dy, label, font_size=24, max_width=300, font_measure=font)

    def save(self, path: Path) -> None:
        content = [
            "<?xml version='1.0' encoding='UTF-8'?>",
            f"<svg xmlns='http://www.w3.org/2000/svg' width='{self.width}' height='{self.height}' viewBox='0 0 {self.width} {self.height}'>",
            "<rect width='100%' height='100%' fill='white'/>",
            *self.elements,
            "</svg>",
        ]
        path.write_text("\n".join(content), encoding="utf-8")


def escape_xml(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&apos;")
    )


class DualDiagram:
    def __init__(self, width: int, height: int, title: str) -> None:
        self.png = DiagramPNG(width, height)
        self.svg = DiagramSVG(width, height)
        self.png.title(title)
        self.svg.title(title)

    def box(self, *args, **kwargs) -> None:
        self.png.box(*args, **kwargs)
        self.svg.box(*args, **kwargs)

    def plain_box(self, *args, **kwargs) -> None:
        self.png.plain_box(*args, **kwargs)
        self.svg.plain_box(*args, **kwargs)

    def ellipse(self, *args, **kwargs) -> None:
        self.png.ellipse(*args, **kwargs)
        self.svg.ellipse(*args, **kwargs)

    def diamond(self, *args, **kwargs) -> None:
        self.png.diamond(*args, **kwargs)
        self.svg.diamond(*args, **kwargs)

    def actor(self, *args, **kwargs) -> None:
        self.png.actor(*args, **kwargs)
        self.svg.actor(*args, **kwargs)

    def arrow(self, *args, **kwargs) -> None:
        self.png.arrow(*args, **kwargs)
        self.svg.arrow(*args, **kwargs)

    def line(self, *args, **kwargs) -> None:
        self.png.line(*args, **kwargs)
        self.svg.line(*args, **kwargs)

    def text(self, x: int, y: int, text: str, **kwargs) -> None:
        self.png.text(x, y, text, **kwargs)
        font_size = kwargs.get("font_size", 28)
        bold = kwargs.get("bold", False)
        italic = kwargs.get("italic", False)
        max_width = kwargs.get("max_width")
        align = kwargs.get("align", "center")
        anchor = {"center": "middle", "left": "start", "right": "end"}[align]
        path = FONT_BOLD if bold else FONT_ITALIC if italic else FONT_REGULAR
        font = ImageFont.truetype(path, size=font_size)
        self.svg._text(x, y, text, font_size=font_size, max_width=max_width, anchor=anchor, bold=bold, italic=italic, font_measure=font)

    def save(self, stem: str) -> Tuple[Path, Path]:
        png_path = OUTPUT_DIAGRAM_DIR / f"{stem}.png"
        svg_path = OUTPUT_DIAGRAM_DIR / f"{stem}.svg"
        self.png.save(png_path)
        self.svg.save(svg_path)
        return png_path, svg_path


def add_field(paragraph, field: str, text: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field} "
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(OxmlElement("w:r"))
    paragraph.add_run(text)
    run._r.append(fld_char_end)


def set_run_font(run, size: int = 13, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 15, 3: 14, 4: 13}.get(level, 13), bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=13)


def add_diagram_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=12, italic=True)
    p.paragraph_format.space_after = Pt(8)


def create_use_case_diagram() -> Tuple[Path, Path]:
    d = DualDiagram(2200, 1600, "Hình 3.1. Use Case Diagram của hệ thống")
    d.actor(90, 190, "Người chơi")
    d.actor(1940, 190, "PlayFab")
    d.plain_box(350, 140, 1500, 1290, "", fill="#ffffff", font_size=30, bold=True)
    d.text(430, 165, "Biên hệ thống", font_size=24, bold=True, align="left")

    use_cases = [
        (480, 250, "Đăng nhập PlayFab"),
        (480, 430, "Nhập tên hiển thị"),
        (480, 610, "Bắt đầu trận đấu"),
        (480, 790, "Di chuyển và Dash"),
        (900, 250, "Tấn công tự động"),
        (900, 430, "Nhận sát thương"),
        (900, 610, "Nhận EXP"),
        (900, 790, "Lên cấp và chọn buff"),
        (1320, 250, "Tạm dừng trò chơi"),
        (1320, 430, "Xem leaderboard"),
        (1320, 610, "Gửi điểm số"),
        (1320, 790, "Đổi theme theo wave"),
    ]
    centers = []
    for x, y, label in use_cases:
        d.ellipse(x, y, 280, 110, label)
        centers.append((x + 140, y + 55, label))

    player_targets = [2, 3, 4, 5, 6, 7, 8, 9]
    for idx in player_targets:
        cx, cy, _ = centers[idx]
        d.line(190, 345, cx - 140, cy)

    for idx in (0, 1, 9, 10):
        cx, cy, _ = centers[idx]
        d.line(cx + 140, cy, 1940, 345)

    d.arrow(620, 305, 620, 430, dashed=True, label="Nếu chưa có\nDisplay Name", label_dx=120)
    d.arrow(1040, 665, 1040, 790, dashed=True, label="Khi đủ EXP")
    d.arrow(1460, 665, 1460, 790, dashed=True, label="Theo tiến trình wave")
    return d.save("3_1_use_case")


def create_sequence_diagram() -> Tuple[Path, Path]:
    d = DualDiagram(2500, 2400, "Hình 3.2. Sequence Diagram của các luồng chính")
    d.text(280, 125, "Luồng 1. Đăng nhập và nhập tên hiển thị", font_size=30, bold=True, align="left")
    lifelines_1 = [("PlayFabLeaderboardManager", 240), ("PlayFab", 760), ("NameInputPanel", 1280), ("Người chơi", 1800)]
    for label, x in lifelines_1:
        d.plain_box(x - 120, 180, 240, 70, label, font_size=24, bold=True)
        d.line(x, 250, x, 690, dashed=True, width=3)
    seq1 = [
        (240, 760, 320, "LoginWithCustomID()"),
        (760, 240, 395, "Login success"),
        (240, 760, 470, "GetPlayerProfile()"),
        (760, 240, 545, "Display Name / hồ sơ"),
        (240, 1280, 620, "Show() nếu chưa có tên"),
        (1800, 1280, 680, "Nhập tên và xác nhận"),
        (1280, 240, 760, "SubmitName()"),
        (240, 760, 835, "UpdateUserTitleDisplayName()"),
        (760, 240, 910, "Tên hợp lệ"),
        (240, 1280, 985, "Hide()"),
    ]
    for x1, x2, y, label in seq1:
        d.arrow(x1, y, x2, y, label=label)

    d.text(280, 1135, "Luồng 2. Tiêu diệt enemy, cộng EXP và chọn buff", font_size=30, bold=True, align="left")
    lifelines_2 = [("PlayerAttack", 200), ("Enemy", 620), ("ExpDropper", 1040), ("PlayerLevelSystem", 1460), ("CardSelectionPanel", 1880), ("BuffCardManager", 2300)]
    for label, x in lifelines_2:
        d.plain_box(x - 120, 1185, 240, 70, label, font_size=22, bold=True)
        d.line(x, 1255, x, 1770, dashed=True, width=3)
    seq2 = [
        (200, 620, 1325, "TryAttack() / SpawnProjectiles()"),
        (620, 620, 1400, "TakeDamage()"),
        (620, 1040, 1475, "OnDeath"),
        (1040, 1460, 1550, "AddExp(expValue)"),
        (1460, 1460, 1625, "Kiểm tra ngưỡng lên cấp"),
        (1460, 1880, 1700, "OnLevelUp"),
        (1880, 2300, 1775, "GetRandomCards(3)"),
        (2300, 1880, 1850, "Danh sách buff hợp lệ"),
    ]
    for x1, x2, y, label in seq2:
        d.arrow(x1, y, x2, y, label=label)
    d.arrow(1880, 1880, 1880, 1925, label="ShowCards() + Time.timeScale = 0", label_dx=180)
    d.arrow(1880, 2000, 2300, 2000, label="OnCardSelected()")
    d.arrow(2300, 2075, 1460, 2075, label="ApplyCard()")
    d.arrow(1880, 2075, 1880, 2150, label="HideCards() + tiếp tục game", label_dx=170)

    d.text(280, 1885, "Luồng 3. Người chơi chết và gửi điểm", font_size=30, bold=True, align="left")
    lifelines_3 = [("Enemy / Projectile", 240), ("PlayerHealth", 720), ("PlayerLevelSystem", 1200), ("PlayFabLeaderboardManager", 1680), ("PlayFab", 2160)]
    for label, x in lifelines_3:
        d.plain_box(x - 120, 1935, 240, 70, label, font_size=22, bold=True)
        d.line(x, 2005, x, 2320, dashed=True, width=3)
    seq3 = [
        (240, 720, 2075, "TakeDamage()"),
        (720, 720, 2145, "Die()"),
        (720, 1200, 2215, "GetTotalExpGained()"),
        (720, 1680, 2285, "SubmitScore(finalScore)"),
        (1680, 2160, 2355, "UpdatePlayerStatistics()"),
    ]
    for x1, x2, y, label in seq3:
        d.arrow(x1, y, x2, y, label=label)
    return d.save("3_2_sequence")


def create_activity_diagram() -> Tuple[Path, Path]:
    d = DualDiagram(1700, 2600, "Hình 3.3. Activity Diagram của vòng lặp gameplay")
    d.ellipse(710, 110, 280, 90, "Bắt đầu")
    d.box(620, 250, 460, 100, "Khởi động game và đăng nhập PlayFab", font_size=30)
    d.diamond(645, 410, 410, 150, "Đã có\nDisplay Name?")
    d.box(190, 660, 430, 110, "Hiển thị NameInputPanel và cập nhật tên", font_size=28)
    d.box(900, 660, 430, 110, "Vào scene chính và chờ người chơi tương tác", font_size=28)
    d.box(620, 860, 460, 110, "Tương tác ChallengePostNPC và StartGame", font_size=28)
    d.box(620, 1040, 460, 110, "WaveSpawner chuẩn bị và bắt đầu wave", font_size=28)
    d.box(620, 1220, 460, 110, "Chiến đấu: di chuyển, dash, tự động tấn công", font_size=28)
    d.diamond(645, 1395, 410, 150, "Enemy chết?")
    d.box(620, 1615, 460, 110, "ExpDropper cộng EXP vào PlayerLevelSystem", font_size=28)
    d.diamond(645, 1790, 410, 150, "Đủ EXP\nlên cấp?")
    d.box(620, 2010, 460, 110, "Hiển thị CardSelectionPanel và áp dụng buff", font_size=28)
    d.diamond(645, 2185, 410, 150, "Hoàn thành wave?")
    d.box(140, 2380, 420, 120, "Boss wave hoặc đổi theme map rồi tiếp tục", font_size=28)
    d.box(620, 2380, 460, 120, "Quay lại vòng chiến đấu wave mới", font_size=28)
    d.box(1140, 2380, 420, 120, "Người chơi chết, gửi điểm và tải leaderboard", font_size=28)

    d.arrow(850, 200, 850, 250)
    d.arrow(850, 350, 850, 410)
    d.arrow(645, 485, 405, 660, label="Không")
    d.arrow(1055, 485, 1115, 660, label="Có", label_dx=36)
    d.arrow(405, 770, 850, 860)
    d.arrow(1115, 770, 850, 860)
    d.arrow(850, 970, 850, 1040)
    d.arrow(850, 1150, 850, 1220)
    d.arrow(850, 1330, 850, 1395)
    d.arrow(645, 1470, 420, 1220, label="Không", label_dx=-24, label_dy=-18)
    d.arrow(850, 1545, 850, 1615, label="Có", label_dx=42)
    d.arrow(850, 1725, 850, 1790)
    d.arrow(645, 1865, 420, 1220, label="Không", label_dx=-24, label_dy=-18)
    d.arrow(850, 1940, 850, 2010, label="Có", label_dx=42)
    d.arrow(850, 2120, 850, 2185)
    d.arrow(645, 2260, 350, 2380, label="Có:\nđến mốc boss\nhoặc đổi theme", label_dx=-18, label_dy=-6)
    d.arrow(850, 2335, 850, 2380, label="Có: wave mới", label_dx=82)
    d.arrow(1055, 2260, 1350, 2380, label="Không nhưng\nngười chơi chết", label_dx=30, label_dy=-6)
    d.arrow(350, 2500, 850, 1220)
    d.arrow(850, 2500, 850, 1220)
    d.ellipse(1210, 2550, 280, 90, "Kết thúc")
    d.arrow(1350, 2500, 1350, 2550)
    return d.save("3_3_activity")


def create_component_diagram() -> Tuple[Path, Path]:
    d = DualDiagram(2200, 1650, "Hình 3.4. Component Diagram của hệ thống")
    d.box(90, 220, 360, 120, "Input\nInputSystem_Actions", font_size=28, bold=True)
    d.box(520, 120, 500, 220, "Player\nPlayerController\nPlayerAttack\nPlayerHealth\nPlayerLevelSystem", font_size=24, bold=True)
    d.box(1220, 120, 560, 280, "UI\nGameUI\nChallengePanel\nNameInputPanel\nPlayerStatsPanel\nCardSelectionPanel\nPauseMenuPanel\nLeaderboardPanel", font_size=22, bold=True)
    d.box(520, 430, 500, 220, "Enemy\nEnemy\nMeleeEnemy\nRangedEnemy\nFlyEnemy\nBossEnemy", font_size=24, bold=True)
    d.box(90, 760, 500, 190, "Wave\nWaveSpawner\nWaveConfig", font_size=26, bold=True)
    d.box(680, 760, 500, 240, "Buff\nBuffCardManager\nBuffCardConfig\nSpiritManager\nOrbitingBallManager", font_size=24, bold=True)
    d.box(1320, 500, 460, 180, "Backend\nPlayFabLeaderboardManager\nPlayFab API", font_size=26, bold=True)
    d.box(1320, 820, 700, 320, "Shared Services\nObjectPool\nAudioManager\nMapThemeManager\nLoadingUIManager\nUtils", font_size=24, bold=True, fill="#f0f5ff")

    d.arrow(450, 280, 520, 230, label="Điều khiển")
    d.arrow(450, 300, 1220, 260, label="Tương tác UI")
    d.arrow(770, 340, 770, 430, label="Chiến đấu")
    d.arrow(770, 650, 330, 760, label="Sự kiện wave")
    d.arrow(340, 950, 340, 1040, label="Wave mới")
    d.arrow(340, 1040, 770, 1040, label="OnLevelUp / buff", label_dy=-14)
    d.arrow(1220, 520, 1110, 520, label="HUD / panel")
    d.arrow(1780, 400, 1780, 500, label="Display Name")
    d.arrow(1550, 680, 1670, 820, label="Dịch vụ dùng chung")
    d.arrow(1180, 900, 1320, 900, label="Spirit / pool")
    d.arrow(590, 860, 1320, 980, label="Spawn / transition")
    d.arrow(1020, 540, 1320, 560, label="Boss / leaderboard")
    return d.save("3_4_component")


def create_class_diagram() -> Tuple[Path, Path]:
    d = DualDiagram(2500, 1800, "Hình 3.5. Class Diagram của các lớp chính")
    d.plain_box(1030, 120, 320, 90, "Singleton<T>", font_size=34, bold=True, fill="#f6f6f6")
    managers = [
        ("PlayerController", 300, 300),
        ("PlayerLevelSystem", 700, 300),
        ("WaveSpawner", 1100, 300),
        ("BuffCardManager", 1500, 300),
        ("GameUI", 1900, 300),
        ("PlayFabLeaderboardManager", 1030, 470),
    ]
    for label, x, y in managers:
        d.plain_box(x, y, 320, 90, label, font_size=26, fill="#ffffff")
        d.arrow(x + 160, y, 1190, 210, label="kế thừa", label_dy=-16)

    d.plain_box(170, 700, 260, 90, "<<interface>>\nIDamageable", font_size=24, bold=True, fill="#eef9ff")
    d.plain_box(560, 700, 280, 90, "PlayerHealth", font_size=28, fill="#ffffff")
    d.plain_box(950, 700, 280, 90, "Enemy", font_size=30, fill="#ffffff")
    d.arrow(700, 700, 300, 790, dashed=True, label="triển khai", label_dx=-20, label_dy=-18)
    d.arrow(1090, 700, 300, 790, dashed=True, label="triển khai", label_dx=20, label_dy=-18)

    enemy_children = [
        ("MeleeEnemy", 760, 900),
        ("RangedEnemy", 1040, 900),
        ("FlyEnemy", 1320, 900),
        ("BossEnemy", 1600, 900),
    ]
    for label, x, y in enemy_children:
        d.plain_box(x, y, 220, 80, label, font_size=24)
        d.arrow(x + 110, y, 1090, 790, label="kế thừa", label_dy=-16)

    bosses = [
        ("LawaChurl_Geo", 1480, 1085),
        ("LawaChurl_Pyro", 1700, 1085),
        ("LawaChurl_Electro", 1920, 1085),
    ]
    for label, x, y in bosses:
        d.plain_box(x, y, 210, 80, label, font_size=22)
        d.arrow(x + 105, y, 1710, 980, label="kế thừa", label_dy=-14)

    d.plain_box(240, 1100, 260, 80, "Projectile", font_size=28)
    projectiles = [
        ("PlayerProjectile", 80, 1280),
        ("EnemyProjectile", 320, 1280),
        ("SpiritProjectileScript", 560, 1280),
    ]
    for label, x, y in projectiles:
        d.plain_box(x, y, 220, 80, label, font_size=21)
        d.arrow(x + 110, y, 370, 1180, label="kế thừa", label_dy=-14)

    d.plain_box(980, 1280, 240, 80, "NPC", font_size=28)
    npcs = [("ChallengePostNPC", 860, 1460), ("ChestBuffBox", 1120, 1460)]
    for label, x, y in npcs:
        d.plain_box(x, y, 220, 80, label, font_size=23)
        d.arrow(x + 110, y, 1100, 1360, label="kế thừa", label_dy=-14)

    d.plain_box(1500, 1280, 260, 80, "PanelBase", font_size=28)
    panels = [
        ("ChallengePanel", 1360, 1460),
        ("NameInputPanel", 1590, 1460),
        ("LeaderboardPanel", 1820, 1460),
        ("CardSelectionPanel", 2050, 1460),
    ]
    for label, x, y in panels:
        d.plain_box(x, y, 210, 80, label, font_size=20)
        d.arrow(x + 105, y, 1630, 1360, label="kế thừa", label_dy=-12)

    d.plain_box(1990, 1280, 230, 80, "PauseMenuPanel", font_size=22)
    d.arrow(2105, 1280, 1630, 1360, label="kế thừa", label_dy=-12)

    data_nodes = [
        ("PlayerData", 370, 1550),
        ("PlayerConfig", 370, 1660),
        ("EnemyData", 700, 1550),
        ("EnemyConfig", 700, 1660),
        ("WaveConfig", 1030, 1660),
        ("BuffCardConfig", 1360, 1660),
    ]
    for label, x, y in data_nodes:
        d.plain_box(x, y, 220, 70, label, font_size=22, fill="#fffdf2")
    d.arrow(480, 1550, 700, 790, label="sở hữu dữ liệu", label_dx=60, label_dy=-28)
    d.arrow(480, 1620, 480, 1660, label="đọc cấu hình", label_dx=66, label_dy=-14)
    d.arrow(810, 1550, 1090, 790, label="sở hữu dữ liệu", label_dx=40, label_dy=-20)
    d.arrow(810, 1620, 810, 1660, label="đọc cấu hình", label_dx=66, label_dy=-14)
    d.arrow(1140, 1660, 1260, 560, label="dùng để spawn", label_dx=80, label_dy=-20)
    d.arrow(1470, 1660, 1660, 560, label="dùng để chọn buff", label_dx=80, label_dy=-20)
    return d.save("3_5_class")


CHAPTER_TEXT = {
    "title": "PHẦN NỘI DUNG THAY THẾ CHƯƠNG 1 VÀ CHƯƠNG 3",
    "sections": [
        (
            "CHƯƠNG 1. TỔNG QUAN",
            [
                ("1.1. Tổng quan", [
                    "Đề tài tập trung xây dựng một trò chơi sinh tồn 3D trên nền tảng Unity, trong đó người chơi điều khiển một nhân vật trong không gian ba chiều, thực hiện di chuyển, né tránh, sử dụng kỹ năng lướt và tự động tấn công kẻ địch trong phạm vi. Trò chơi được thiết kế theo hướng dễ tiếp cận ở thao tác điều khiển nhưng vẫn có chiều sâu về chiến thuật nhờ hệ thống wave tăng dần, boss wave, lên cấp chọn buff và bảng xếp hạng trực tuyến. Bên cạnh việc tạo ra một sản phẩm có thể chơi được, đề tài còn hướng tới tổ chức mã nguồn thành các module rõ ràng, tách dữ liệu cấu hình khỏi logic xử lý và áp dụng các giải pháp tối ưu phù hợp với game có nhiều đối tượng runtime hoạt động đồng thời.",
                ]),
                ("1.1.1. Lý do chọn đề tài", [
                    "Các trò chơi sinh tồn mang yếu tố Roguelike hiện nay thu hút người chơi nhờ vòng lặp gameplay nhanh, giá trị chơi lại cao và khả năng tạo nhiều tổ hợp phát triển sức mạnh khác nhau qua mỗi lần chơi. Tuy nhiên, phần lớn các sản phẩm phổ biến vẫn nghiêng về 2D, trong khi các trò chơi 3D cùng thể loại phải giải quyết đồng thời nhiều bài toán hơn về camera, chuyển động trong không gian, số lượng enemy lớn, projectile dày đặc và ổn định hiệu năng. Việc lựa chọn đề tài phát triển trò chơi sinh tồn 3D vì vậy vừa có ý nghĩa thực tiễn, vừa tạo điều kiện áp dụng kiến thức về lập trình hướng đối tượng, tổ chức dữ liệu, xử lý va chạm, thiết kế giao diện và tích hợp dịch vụ trực tuyến trong cùng một sản phẩm hoàn chỉnh.",
                ]),
                ("1.1.2. Mục tiêu nghiên cứu", [
                    "Mục tiêu của đề tài là xây dựng một trò chơi sinh tồn 3D có thể vận hành ổn định trên PC, trong đó người chơi được điều khiển nhân vật với các cơ chế cốt lõi gồm di chuyển, dash và tự động tấn công. Hệ thống gameplay phải hỗ trợ nhiều loại enemy như cận chiến, đánh xa, bay và boss nguyên tố; đồng thời cho phép người chơi tăng tiến sức mạnh thông qua kinh nghiệm, lên cấp và lựa chọn buff ngẫu nhiên. Ngoài ra, hệ thống wave phải có khả năng mở rộng theo hướng vô hạn, tăng độ khó theo tiến trình, hỗ trợ boss wave và chuyển đổi theme bản đồ theo chu kỳ. Cuối cùng, đề tài cần tích hợp PlayFab để quản lý tên hiển thị và lưu điểm số trực tuyến, từ đó hoàn thiện yếu tố cạnh tranh của sản phẩm.",
                ]),
                ("1.1.3. Phạm vi nghiên cứu", [
                    "Đề tài được triển khai trên nền tảng PC, ưu tiên môi trường Windows, sử dụng Unity 6 kết hợp ngôn ngữ C# để xây dựng toàn bộ logic gameplay, giao diện và luồng điều khiển tổng thể. Dịch vụ trực tuyến được sử dụng là PlayFab, với phạm vi tập trung vào đăng nhập bằng Custom ID, lưu Display Name và quản lý leaderboard. Đề tài không bao gồm việc xây dựng máy chủ riêng, không triển khai multiplayer thời gian thực và không phát triển hệ quản trị cơ sở dữ liệu độc lập. Về phạm vi chức năng, sản phẩm tập trung vào vòng lặp gameplay chính gồm bắt đầu trận đấu, spawn enemy, chiến đấu, cộng EXP, lên cấp chọn buff, tiếp tục wave, đối đầu boss, đổi theme map và gửi điểm số khi người chơi thất bại.",
                ]),
                ("1.2. Khảo sát thực trạng", [
                    "Qua khảo sát các trò chơi sinh tồn mang phong cách Roguelike trên thị trường, có thể thấy hướng tiếp cận phổ biến là giảm số lượng thao tác điều khiển trực tiếp để người chơi tập trung vào di chuyển, giữ vị trí và lựa chọn phương án nâng cấp. Các trò chơi thành công thường có nhịp độ trận đấu tăng nhanh, hệ thống sức mạnh phát triển rõ ràng và nhiều tổ hợp kỹ năng tạo ra giá trị chơi lại. Tuy nhiên, nhiều sản phẩm chỉ dừng ở không gian 2D hoặc 2.5D, trong khi các trò chơi 3D cùng thể loại phải đối mặt với yêu cầu cao hơn về tổ chức camera, ổn định điều khiển và tối ưu hiệu năng khi số lượng enemy, projectile và hiệu ứng xuất hiện lớn.",
                ]),
                ("1.2.4. Kết luận khảo sát và hướng đi của đề tài", [
                    "Từ kết quả khảo sát, hướng đi phù hợp của đề tài là xây dựng một trò chơi sinh tồn 3D giữ được sự đơn giản trong điều khiển của dòng Survivor, đồng thời tận dụng không gian ba chiều để tăng tính trực quan và mức độ hoàn chỉnh của sản phẩm. Dự án lựa chọn cách tiếp cận trong đó người chơi chủ yếu di chuyển, dash, né tránh và chọn buff, còn thao tác tấn công được tự động hóa nhằm duy trì nhịp độ nhanh. Hệ thống được định hình quanh các khối chính gồm Player, Enemy, Wave, Buff, UI, PlayFab, ObjectPool và MapThemeManager, bám sát những gì đã được triển khai trong project hiện tại.",
                ]),
                ("1.3. Đề xuất phương pháp giải quyết", [
                    "Về công nghệ, đề tài sử dụng Unity 6 với Universal Render Pipeline để xây dựng môi trường 3D, sử dụng C# để hiện thực gameplay và Unity Input System để tách riêng lớp xử lý input khỏi phần logic phía sau. Các dữ liệu quan trọng như cấu hình player, enemy, buff card và wave được tách thành ScriptableObject để thuận tiện cho việc cân bằng và mở rộng. Nhân vật người chơi và enemy được điều khiển bằng CharacterController, kết hợp với xử lý theo trạng thái chiến đấu và sự kiện gameplay. Logic AI của enemy được xây dựng trực tiếp trong các lớp gameplay, không phụ thuộc NavMeshAgent. Về hiệu năng, project áp dụng Object Pooling cho projectile, enemy, effect, damage text và một số đối tượng runtime khác. Về dịch vụ trực tuyến, PlayFab được dùng để đăng nhập, quản lý tên hiển thị và đồng bộ leaderboard mà không cần xây dựng backend riêng.",
                ]),
                ("1.4. Mô tả yêu cầu và mô hình bài toán", []),
                ("1.4.1. Yêu cầu chức năng", []),
                ("1.4.1.1. Chức năng định danh người chơi", [
                    "Mục đích của chức năng này là bảo đảm mỗi người chơi có một danh tính hợp lệ trước khi tham gia vào hệ thống xếp hạng trực tuyến. Điều kiện kích hoạt bắt đầu ngay khi trò chơi khởi động. Thành phần chịu trách nhiệm chính là PlayFabLeaderboardManager, phối hợp với NameInputPanel và LoadingUIManager.",
                    "Dữ liệu đầu vào của chức năng gồm Custom ID cục bộ được lưu trong PlayerPrefs và dữ liệu hồ sơ người chơi trả về từ PlayFab. Hệ thống trước hết thực hiện đăng nhập bằng LoginWithCustomID, sau đó gọi API tải hồ sơ để kiểm tra Display Name hiện có. Nếu hồ sơ chưa có tên hiển thị, NameInputPanel phải được hiển thị để yêu cầu người chơi nhập tên mới. Tên nhập vào cần được cắt khoảng trắng dư thừa và kiểm tra độ dài hợp lệ trước khi gửi lên PlayFab.",
                    "Xử lý chính của hệ thống gồm ba bước: đăng nhập, tải hồ sơ và cập nhật tên khi cần thiết. Nếu tên không hợp lệ hoặc bị từ chối, hệ thống phải hiển thị thông báo lỗi rõ ràng và giữ người chơi ở trạng thái nhập lại. Nếu cập nhật thành công, CurrentDisplayName được đồng bộ về client, panel nhập tên đóng lại và trò chơi chuyển sang trạng thái sẵn sàng.",
                    "Kết quả đầu ra của chức năng là người chơi có một Display Name hợp lệ, có thể được dùng để hiển thị trong leaderboard và trong các thông báo chào mừng. Ràng buộc triển khai là toàn bộ luồng này phải hoàn thành trước khi người chơi bước vào vòng gameplay chính, đồng thời cần tránh khóa cứng giao diện nếu API trả lỗi hoặc mạng chậm.",
                ]),
                ("1.4.1.2. Chức năng bắt đầu trận đấu", [
                    "Mục đích của chức năng này là cho phép người chơi chủ động chuyển từ trạng thái chờ sang trạng thái chiến đấu. Điều kiện kích hoạt là người chơi di chuyển đến vùng tương tác của ChallengePostNPC và nhấn phím tương tác. Các thành phần chính tham gia gồm NPC, InteractPanel, ChallengePanel, PlayerController, WaveSpawner và PlayerStatsPanel.",
                    "Dữ liệu đầu vào của chức năng là trạng thái vị trí của player trong vùng trigger, input tương tác và sự tồn tại của ChallengePanel trong scene. Khi người chơi vào phạm vi hợp lệ, hệ thống hiển thị InteractPanel để gợi ý thao tác. Khi người chơi xác nhận tương tác, ChallengePanel xuất hiện, khóa input di chuyển và chờ người chơi nhấn StartGame.",
                    "Xử lý chính gồm mở panel thử thách, phát tín hiệu onGameStart, gọi StartNextWave của WaveSpawner và hiển thị HUD thống kê trận đấu. Sau khi StartGame được kích hoạt, ChallengePostNPC bị vô hiệu hóa để không tiếp tục can thiệp vào vòng chơi hiện tại.",
                    "Kết quả đầu ra là trận đấu chính thức bắt đầu, wave đầu tiên được chuẩn bị hoặc spawn, HUD bắt đầu cập nhật máu, EXP, level và wave. Ràng buộc là quá trình chuyển trạng thái phải đồng bộ giữa giao diện, thời gian trò chơi và input người chơi để không tạo ra tình trạng nhân vật vẫn di chuyển khi panel đang mở.",
                ]),
                ("1.4.1.3. Chức năng điều khiển nhân vật", [
                    "Mục đích của chức năng này là cung cấp khả năng điều khiển cơ bản cho nhân vật trong môi trường 3D. Điều kiện kích hoạt là người chơi đã vào trận và input được mở. Thành phần chịu trách nhiệm chính là PlayerController kết hợp InputSystem_Actions, CharacterController và CameraFollow.",
                    "Dữ liệu đầu vào của chức năng gồm vector di chuyển từ Input System, trạng thái mặt đất, hướng camera hiện tại, thông số dash và trạng thái khóa input. Hệ thống phải quy đổi input từ mặt phẳng camera sang hướng di chuyển thực tế trong thế giới 3D, đồng thời xoay model nhân vật theo hướng chuyển động để tạo cảm giác điều khiển trực quan.",
                    "Xử lý chính gồm di chuyển thường, áp dụng gravity, kiểm tra grounded và thực thi dash khi người chơi nhấn lệnh trong lúc dash cooldown đã hết. Khi dash đang diễn ra, tốc độ di chuyển được thay bằng dashSpeed, thời gian dash được đếm lùi và sau đó khôi phục trạng thái di chuyển bình thường. Hệ thống cũng phải hỗ trợ khóa hoặc mở input khi người chơi đang pause, nhập tên, chọn buff hoặc trong thời gian chuyển theme.",
                    "Kết quả đầu ra là nhân vật có thể di chuyển mượt, đổi hướng hợp lý và thực hiện dash đúng nhịp. Ràng buộc triển khai là không được để input tồn đọng khi khóa điều khiển, tránh trường hợp người chơi thả phím nhưng nhân vật vẫn di chuyển sau khi mở lại input.",
                ]),
                ("1.4.1.4. Chức năng chiến đấu của người chơi", [
                    "Mục đích của chức năng này là tự động hóa thao tác tấn công nhưng vẫn duy trì tính chiến thuật thông qua việc chọn mục tiêu, chu kỳ bắn và các buff tăng cường. Điều kiện kích hoạt là player còn sống, có PlayerData hợp lệ và attack cooldown đã về 0. Thành phần chính gồm PlayerAttack, PlayerProjectile, ObjectPool và PlayerData.",
                    "Dữ liệu đầu vào của chức năng gồm attack range, attack cooldown, projectile speed, projectile lifetime, chỉ số damage hiện tại, trạng thái multishot, AoE và các hiệu ứng hỗ trợ từ spirit hoặc orbiting ball. Hệ thống phải quét vùng xung quanh player để tìm enemy gần nhất còn sống, sau đó tính hướng bắn từ attackPoint tới mục tiêu đó.",
                    "Xử lý chính của hệ thống là gọi TryAttack theo chu kỳ, xác định target, sinh projectile từ ObjectPool và khởi tạo projectile bằng các thông số tấn công hiện tại. Nếu người chơi có multishot, hệ thống phải tạo nhiều hướng bắn theo góc lệch xác định; nếu có AoE, projectile phải truyền thông tin bán kính và hệ số sát thương nổ; nếu có các buff spirit hoặc orbiting ball, những đối tượng phụ trợ này tiếp tục hoạt động độc lập quanh player nhưng vẫn lấy damage source từ PlayerData.",
                    "Kết quả đầu ra là các projectile hợp lệ được sinh ra, bay đúng hướng, gây sát thương lên enemy trong phạm vi va chạm và phản ánh đúng sức mạnh hiện tại của player. Ràng buộc triển khai là không được instantiate projectile liên tục ngoài pool trong vòng lặp chiến đấu bình thường, trừ trường hợp dự phòng khi pool trống và hệ thống có cơ chế tạo bổ sung.",
                ]),
                ("1.4.1.5. Chức năng quản lý enemy", [
                    "Mục đích của chức năng này là tạo ra các đối tượng đối kháng đa dạng và duy trì áp lực chiến đấu cho người chơi. Điều kiện kích hoạt là WaveSpawner phát lệnh spawn từ cấu hình wave hoặc boss wave. Thành phần chính gồm Enemy, EnemyData, các lớp kế thừa như MeleeEnemy, RangedEnemy, FlyEnemy, BossEnemy và projectile của enemy.",
                    "Dữ liệu đầu vào của chức năng gồm loại pool cần spawn, cấu hình chỉ số enemy từ EnemyConfig, vị trí spawn, loại enemy, attack range, damage, cooldown và random variation của chỉ số. Sau khi được spawn từ ObjectPool, mỗi enemy phải lấy dữ liệu từ EnemyData, tìm Player theo tag và bắt đầu vòng update AI riêng.",
                    "Xử lý chính là phân nhánh hành vi theo EnemyType. Enemy cận chiến di chuyển áp sát và gây sát thương tiếp xúc; enemy đánh xa giữ khoảng cách và bắn projectile về phía player; enemy bay có biến thể nhịp bắn hoặc chuyển động riêng; boss kế thừa từ BossEnemy để hỗ trợ phase, kỹ năng đặc biệt và logic tử vong riêng. Toàn bộ chuyển động của enemy trong project hiện tại được thực hiện bằng CharacterController và logic trạng thái trong lớp Enemy, không sử dụng NavMeshAgent.",
                    "Kết quả đầu ra là các enemy có thể truy đuổi, tấn công, nhận sát thương, chết và phát sự kiện OnDeath phục vụ cho hệ thống wave và kinh nghiệm. Ràng buộc triển khai là enemy sau khi bị hạ gục phải được dọn trạng thái đúng cách trước khi quay về pool, tránh mang theo dữ liệu máu hoặc listener cũ sang lần spawn tiếp theo.",
                ]),
                ("1.4.1.6. Chức năng quản lý wave và độ khó", [
                    "Mục đích của chức năng này là điều phối nhịp độ trận đấu theo từng đợt enemy, kiểm soát boss wave và tăng độ khó theo thời gian. Điều kiện kích hoạt là người chơi đã bắt đầu trận thông qua ChallengePanel. Thành phần điều phối trung tâm là WaveSpawner, sử dụng dữ liệu từ WaveConfig và danh sách active enemy trong runtime.",
                    "Dữ liệu đầu vào gồm danh sách wave cấu hình sẵn, các EnemyGroup trong từng wave, preparation time, spawn position, spread radius, bossPoolTypes, autoScale và scalePerWave. Khi bắt đầu mỗi wave, hệ thống phải chọn dữ liệu wave phù hợp theo currentWave. Nếu currentWave vượt ra ngoài số wave cấu hình, hệ thống cần sinh endless wave bằng cách tái sử dụng mẫu cũ và tăng số lượng enemy cùng bán kính spawn theo vòng lặp.",
                    "Xử lý chính của hệ thống là tăng currentWave, chuẩn bị wave, bắt đầu spawn enemy theo nhóm hoặc theo hiệu ứng circle spawn, theo dõi activeEnemies, cập nhật số lượng enemy đã spawn và số lượng còn lại. Với boss wave, hệ thống phải chọn ngẫu nhiên một boss hợp lệ từ bossPoolTypes, phát sự kiện OnBossWaveStart và spawn boss ở vị trí quy định. Trong suốt tiến trình chơi, nếu autoScale được bật, WaveSpawner phải tăng chỉ số máu, sát thương tiếp xúc và sát thương projectile của enemy theo hệ số scalePerWave.",
                    "Kết quả đầu ra là nhịp độ wave được duy trì liên tục, độ khó tăng dần và boss wave xuất hiện theo đúng cấu hình. Ràng buộc triển khai là phải tách rõ trạng thái giữa wave hiện tại và session mới, tránh trường hợp coroutine cũ tiếp tục spawn khi người chơi đã nhảy wave hoặc chuyển trạng thái map.",
                ]),
                ("1.4.1.7. Chức năng đổi theme bản đồ", [
                    "Mục đích của chức năng này là thay đổi hình thức trình bày của môi trường chơi theo tiến trình wave mà không phá vỡ vòng lặp gameplay. Điều kiện kích hoạt là khi WaveSpawner hoàn thành một wave và phát hiện wave tiếp theo thuộc theme khác. Thành phần chính gồm WaveSpawner, MapThemeManager và LoadingUIManager.",
                    "Dữ liệu đầu vào của chức năng gồm currentWave, upcomingWave, danh sách theme trong scene và các vật liệu hoặc effectRoot tương ứng với từng theme. Hệ thống phải xác định theme mục tiêu bằng ResolveThemeIndexForWave, kiểm tra xem việc chuyển theme có thực sự cần thiết hay không, rồi mới tiến hành transition.",
                    "Xử lý chính bao gồm khóa gameplay tạm thời bằng cách dừng timeScale, ẩn các giao diện tương tác không cần thiết, vô hiệu hóa input của player và yêu cầu MapThemeManager thực hiện chuyển theme. Sau khi loading transition hoàn tất, theme mới được áp dụng lên ground, wall và các hiệu ứng môi trường, sau đó gameplay được mở lại và wave tiếp theo bắt đầu.",
                    "Kết quả đầu ra là bản đồ được thay theme đúng mốc wave, người chơi quay lại vòng chiến đấu với trạng thái điều khiển ổn định. Ràng buộc là quá trình chuyển theme không được làm mất dữ liệu wave hiện tại, không được để người chơi thao tác trong lúc scene đang bị che bởi loading transition.",
                ]),
                ("1.4.1.8. Chức năng kinh nghiệm và lên cấp", [
                    "Mục đích của chức năng này là tạo cơ chế tăng tiến sức mạnh và điều chỉnh nhịp lựa chọn chiến thuật của người chơi. Điều kiện kích hoạt bắt đầu khi một enemy bị tiêu diệt. Thành phần chính gồm Enemy, ExpDropper, PlayerLevelSystem và PlayerStatsPanel.",
                    "Dữ liệu đầu vào là expValue từ EnemyData, currentExp, expToNextLevel, currentLevel và expScalingFactor. Khi enemy chết, ExpDropper phải nhận sự kiện OnDeath và cộng lượng EXP tương ứng vào PlayerLevelSystem. Sau mỗi lần cộng EXP, hệ thống phải cập nhật progress hiện tại và kiểm tra xem người chơi đã vượt ngưỡng lên cấp hay chưa.",
                    "Xử lý chính bao gồm cộng currentExp, cộng tổng totalExpGained, phát sự kiện cập nhật HUD và lặp kiểm tra điều kiện lên cấp nếu lượng EXP nhận được vượt qua nhiều mốc cùng lúc. Khi lên cấp, currentLevel tăng thêm một đơn vị, expToNextLevel được scale lên cho mốc tiếp theo, hệ thống phát OnLevelUp và tạm dừng trò chơi để chờ chọn buff.",
                    "Kết quả đầu ra là cấp độ của người chơi tăng theo tiến trình chiến đấu, thanh EXP và level trên HUD luôn được cập nhật đúng. Ràng buộc triển khai là thời gian trò chơi phải dừng đúng lúc hiển thị lựa chọn buff, nhưng không được làm mất trạng thái wave hoặc khiến sự kiện gameplay quan trọng bị bỏ qua.",
                ]),
                ("1.4.1.9. Chức năng chọn buff và tăng sức mạnh", [
                    "Mục đích của chức năng này là cho phép người chơi cá nhân hóa hướng phát triển sức mạnh ngay trong trận đấu. Điều kiện kích hoạt là PlayerLevelSystem phát sự kiện OnLevelUp hoặc một luồng gameplay khác chủ động gọi mở chọn buff. Thành phần chính gồm BuffCardManager, BuffCardConfig, CardSelectionPanel, PlayerData, PlayerHealth, SpiritManager và OrbitingBallManager.",
                    "Dữ liệu đầu vào là toàn bộ danh sách buff card, rarity, maxLevel, loại buff, chỉ số hiện tại của player và lịch sử card đã chọn. Trước khi hiển thị, BuffCardManager phải lọc các thẻ đã đạt maxLevel, tính trọng số ngẫu nhiên theo rarity và luckBonus, sau đó rút ra đúng số lượng card cho mỗi lần hiển thị. CardSelectionPanel chịu trách nhiệm sinh UI cho từng thẻ, hiển thị mô tả và level hiện tại để người chơi lựa chọn.",
                    "Xử lý chính sau khi chọn card là áp dụng buff vào hệ thống phù hợp. Các buff cơ bản sẽ tăng damage, move speed, attack speed, max health hoặc EXP bonus trực tiếp trên PlayerData và PlayerHealth. Các buff dạng tính năng như multishot, AoE, spirit hoặc orbiting ball sẽ cập nhật tham số chiến đấu hoặc thêm thực thể hỗ trợ vào player thông qua SpiritManager và OrbitingBallManager.",
                    "Kết quả đầu ra là player có bộ chỉ số hoặc kỹ năng mới ngay trong runtime, đồng thời panel lựa chọn đóng lại và trò chơi tiếp tục. Ràng buộc triển khai là chỉ được tiếp tục khi người chơi đã chọn xong một card, và buff áp dụng phải đồng bộ với hệ thống chiến đấu hiện thời thay vì chỉ thay đổi dữ liệu lưu trữ.",
                ]),
                ("1.4.1.10. Chức năng giao diện trong trận", [
                    "Mục đích của chức năng này là cung cấp phản hồi trực quan liên tục về trạng thái trận đấu và trạng thái hệ thống. Điều kiện kích hoạt là game đang trong scene chính và các panel liên quan đã được khởi tạo trong GameUI. Thành phần chính gồm PlayerStatsPanel, InteractPanel, CardSelectionPanel, NotiPanel, LeaderboardPanel, PauseMenuPanel, NameInputPanel và các UI entry thành phần.",
                    "Dữ liệu đầu vào của chức năng gồm các sự kiện từ PlayerHealth, PlayerLevelSystem, WaveSpawner và PlayFabLeaderboardManager. PlayerStatsPanel phải cập nhật máu, thanh EXP, level và wave theo thời gian thực. InteractPanel hiển thị hướng dẫn tương tác khi người chơi đứng gần NPC. NotiPanel hoặc các panel thông báo khác dùng để phản hồi một số trạng thái quan trọng như chào mừng người chơi hoặc lỗi nhập tên.",
                    "Xử lý chính là đăng ký listener vào các event phù hợp, hiển thị hoặc ẩn panel đúng trạng thái game và ngăn chồng chéo giữa các màn hình như nhập tên, leaderboard, pause, chọn buff. Khi người chơi xem leaderboard, hệ thống phải tải top xếp hạng và vị trí hiện tại của người chơi rồi hiển thị vào LeaderboardPanel. Khi chọn buff hoặc pause, UI phải đồng bộ với trạng thái dừng thời gian và khóa input.",
                    "Kết quả đầu ra là người chơi luôn nhận được thông tin cần thiết để tiếp tục trận đấu mà không phải suy đoán trạng thái hệ thống. Ràng buộc là UI không được che lấp nhau sai logic, đồng thời mỗi panel cần quản lý rõ quyền sở hữu input và trạng thái đóng mở của chính nó.",
                ]),
                ("1.4.1.11. Chức năng tạm dừng và kết thúc trận", [
                    "Mục đích của chức năng này là cho phép người chơi dừng tạm thời trận đấu, điều chỉnh một số cài đặt giao diện và chủ động kết thúc lượt chơi. Điều kiện kích hoạt là người chơi nhấn phím ESC trong lúc các panel chặn input khác không chiếm quyền điều khiển. Thành phần chính là PauseMenuPanel, kết hợp với LeaderboardPanel, PlayerController và PlayerHealth.",
                    "Dữ liệu đầu vào của chức năng là trạng thái mở của các panel hiện tại, currentView của PauseMenuPanel và trạng thái sống/chết của player. Khi được kích hoạt, hệ thống phải dừng timeScale, khóa input người chơi, mở lớp phủ pause và đưa người chơi vào menu chính của màn hình pause. Từ đây, người chơi có thể quay lại trận, mở phần cài đặt, mở leaderboard hoặc chọn kết thúc game.",
                    "Nếu người chơi chọn Resume, hệ thống phải khôi phục timeScale và mở lại input. Nếu người chơi mở leaderboard từ pause, LeaderboardPanel phải xuất hiện mà không cướp trạng thái pause tổng thể. Nếu người chơi chọn kết thúc trận, hệ thống cần đưa player về trạng thái chết logic, từ đó tái sử dụng luôn luồng game over và gửi điểm số hiện tại lên PlayFab.",
                    "Kết quả đầu ra là người chơi có thể kiểm soát phiên chơi tốt hơn, tạm dừng và kết thúc đúng cách. Ràng buộc là pause không được hoạt động khi các panel như NameInputPanel hoặc CardSelectionPanel đang giữ quyền thao tác độc lập, tránh xung đột trạng thái.",
                ]),
                ("1.4.1.12. Chức năng bảng xếp hạng", [
                    "Mục đích của chức năng này là lưu trữ và hiển thị điểm số trực tuyến của người chơi. Điều kiện kích hoạt chính là khi PlayerHealth xử lý cái chết của người chơi hoặc khi LeaderboardPanel cần tải dữ liệu hiển thị. Thành phần cốt lõi gồm PlayerHealth, PlayerLevelSystem, PlayFabLeaderboardManager và LeaderboardPanel.",
                    "Dữ liệu đầu vào của chức năng là tổng lượng EXP người chơi đã tích lũy trong suốt trận đấu, được lưu tại totalExpGained của PlayerLevelSystem. Khi người chơi chết, PlayerHealth phải lấy giá trị này, chuyển sang số nguyên điểm cuối cùng và gửi đến PlayFabLeaderboardManager để cập nhật statistic tương ứng trên PlayFab. Ngoài top leaderboard, hệ thống còn cần tải vị trí hiện tại của người chơi để làm nổi bật dòng dữ liệu của chính họ.",
                    "Xử lý chính bao gồm SubmitScore, GetLeaderboardData và GetPlayerLeaderboardData. Sau khi dữ liệu được trả về, LeaderboardPanel phải sinh danh sách entry, hiển thị tên người chơi, hạng và điểm số. Những dòng tương ứng với tài khoản hiện tại cần được đánh dấu khác biệt để tăng khả năng nhận biết.",
                    "Kết quả đầu ra là dữ liệu xếp hạng được đồng bộ lên dịch vụ trực tuyến và hiển thị lại trong giao diện của game. Ràng buộc triển khai là điểm số phải bám đúng totalExpGained, không thay bằng thời gian sống, số quái tiêu diệt hoặc các chỉ số ngoài phạm vi hệ thống hiện tại.",
                ]),
                ("1.4.2. Yêu cầu phi chức năng", [
                    "Về hiệu năng, hệ thống phải duy trì khả năng hoạt động ổn định khi trên màn hình đồng thời xuất hiện nhiều enemy, projectile, effect và damage text. Điều này đòi hỏi phải sử dụng ObjectPool cho các đối tượng runtime có tần suất tạo hủy cao, đồng thời hạn chế instantiate và destroy trực tiếp trong vòng lặp gameplay thông thường.",
                    "Về tổ chức mã nguồn, project cần được chia theo các module rõ ràng như Player, Enemy, Wave, UI, Buff, Backend và Utils. Việc phân tách này giúp giảm phụ thuộc chéo, hỗ trợ kiểm thử cục bộ và thuận lợi khi bổ sung tính năng mới mà không phá vỡ các thành phần đang ổn định.",
                    "Về khả năng mở rộng, hệ thống phải cho phép thêm enemy mới, buff mới hoặc wave mới chủ yếu thông qua prefab và ScriptableObject, thay vì buộc phải sửa sâu vào logic lõi. Điều này đặc biệt quan trọng với đồ án game, nơi yêu cầu cân bằng hoặc điều chỉnh nội dung thường thay đổi trong giai đoạn hoàn thiện.",
                    "Về tính ổn định luồng chơi, các trạng thái đặc biệt như pause, chọn buff, chuyển theme map và nhập tên không được làm sai lệch wave hiện tại, không được để player giữ input cũ và không được phá vỡ sự đồng bộ giữa gameplay với UI. Về khả năng sử dụng, các panel và HUD phải đủ rõ để người chơi biết trạng thái máu, EXP, level, wave và leaderboard. Về dữ liệu online, các thao tác đăng nhập, cập nhật tên, gửi điểm và tải xếp hạng cần có phản hồi rõ ràng để người chơi nhận biết tình trạng kết nối và kết quả xử lý.",
                ]),
                ("1.4.3. Mô hình bài toán tổng quát", [
                    "Mô hình bài toán của hệ thống được tổ chức quanh tám khối chính là Player, Enemy, Wave, BuffCard, UI, ObjectPool, PlayFab và MapTheme. Trong đó Player là khối trung tâm tiếp nhận input, điều khiển di chuyển, dash, tấn công và nhận sát thương. Enemy là khối tiêu thụ đầu ra chiến đấu từ Player và tạo áp lực ngược lại thông qua va chạm hoặc projectile. Wave đóng vai trò điều phối tiến trình trận đấu, sinh enemy, theo dõi wave hiện tại và xác định khi nào xuất hiện boss hoặc cần chuyển theme.",
                    "BuffCard là khối thay đổi trạng thái chiến đấu của Player bằng cách áp dụng các hiệu ứng tăng chỉ số hoặc mở rộng kỹ năng. UI là khối trình bày, lấy dữ liệu từ Player, Wave và Backend để phản ánh đúng trạng thái của toàn hệ thống. ObjectPool là khối hạ tầng dùng chung, chịu trách nhiệm cấp phát và thu hồi lại các đối tượng runtime như projectile, enemy, effect hoặc spirit. PlayFab là khối dịch vụ trực tuyến, tiêu thụ thông tin định danh và điểm số từ game client rồi cung cấp lại dữ liệu xếp hạng. MapTheme là khối hỗ trợ trình bày môi trường, được Wave điều phối để đổi vật liệu và hiệu ứng theo mốc tiến trình trận đấu.",
                ]),
                ("1.4.4. Mô tả hoạt động của hệ thống", [
                    "Khi trò chơi khởi động, PlayFabLeaderboardManager thực hiện đăng nhập bằng Custom ID cục bộ và gọi API lấy hồ sơ người chơi. Nếu hồ sơ chưa có Display Name, NameInputPanel được mở để người chơi nhập tên hợp lệ. Sau khi định danh thành công, game chuyển sang trạng thái chờ trong scene chính. Người chơi tiếp cận ChallengePostNPC, hệ thống hiện InteractPanel để gợi ý thao tác, sau đó ChallengePanel mở ra và cho phép nhấn StartGame để bắt đầu trận đấu.",
                    "Ngay khi trận đấu bắt đầu, WaveSpawner tăng currentWave, lấy dữ liệu từ WaveConfig, chuẩn bị wave và bắt đầu sinh enemy theo từng nhóm hoặc theo cơ chế circle spawn tùy cấu hình. Trong suốt quá trình chơi, PlayerController tiếp nhận input di chuyển và dash, còn PlayerAttack tự động tìm enemy gần nhất trong tầm rồi sinh projectile từ ObjectPool để gây sát thương. Enemy được điều khiển bằng CharacterController và logic AI trong lớp Enemy hoặc các lớp dẫn xuất, sau đó phản công lại người chơi bằng tiếp xúc hoặc projectile.",
                    "Khi một enemy bị hạ gục, sự kiện OnDeath của nó kích hoạt ExpDropper cộng EXP trực tiếp vào PlayerLevelSystem. Nếu đạt ngưỡng lên cấp, PlayerLevelSystem tăng cấp, tính lại mốc EXP mới, phát OnLevelUp và tạm dừng Time.timeScale. CardSelectionPanel lắng nghe sự kiện này, yêu cầu BuffCardManager chọn ngẫu nhiên ba buff hợp lệ rồi hiển thị để người chơi chọn. Sau khi buff được áp dụng lên PlayerData, PlayerHealth hoặc các manager kỹ năng phụ trợ, game tiếp tục vận hành.",
                    "Khi wave hiện tại kết thúc, WaveSpawner kiểm tra xem có cần chuyển theme map cho wave tiếp theo hay không. Nếu cần, MapThemeManager phối hợp LoadingUIManager tạm khóa gameplay, đổi vật liệu và hiệu ứng môi trường rồi mở lại input để bắt đầu wave tiếp theo. Ở các mốc boss wave, WaveSpawner chọn ngẫu nhiên boss từ danh sách bossPoolTypes và phát sự kiện thông báo tương ứng.",
                    "Khi máu người chơi giảm về 0, PlayerHealth xử lý trạng thái chết, lấy totalExpGained từ PlayerLevelSystem làm điểm cuối cùng và gọi PlayFabLeaderboardManager.SubmitScore để gửi dữ liệu lên PlayFab. Sau khi cập nhật thành công, hệ thống có thể tải leaderboard tổng và leaderboard quanh người chơi hiện tại để hiển thị trong LeaderboardPanel. Toàn bộ vòng lặp này tạo thành một chu trình khép kín từ định danh người chơi đến gameplay, tăng tiến sức mạnh và đồng bộ kết quả trực tuyến.",
                ]),
            ],
        ),
        (
            "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG",
            [
                ("3.1. Use Case Diagram", [
                    "Use Case Diagram của hệ thống được xây dựng dựa trực tiếp trên nhóm yêu cầu chức năng ở mục 1.4.1. Hai tác nhân chính của hệ thống là Người chơi và PlayFab. Người chơi tương tác với trò chơi thông qua các hành vi định danh, bắt đầu trận đấu, điều khiển nhân vật, chiến đấu, lên cấp, chọn buff, tạm dừng và xem bảng xếp hạng. PlayFab đóng vai trò là tác nhân ngoài hệ thống, chịu trách nhiệm xử lý đăng nhập, cập nhật Display Name, nhận điểm số và trả về dữ liệu leaderboard. Việc tách hai actor như vậy giúp làm rõ ranh giới giữa gameplay nội bộ và các thao tác phụ thuộc dịch vụ trực tuyến.",
                    "Gợi ý vẽ sơ đồ: đặt actor Người chơi ở bên trái và PlayFab ở bên phải. Bên trong khung hệ thống cần có các use case: Đăng nhập PlayFab, Nhập tên hiển thị, Bắt đầu trận đấu, Di chuyển và Dash, Tấn công tự động, Nhận sát thương, Nhận EXP, Lên cấp và chọn buff, Tạm dừng trò chơi, Xem leaderboard, Gửi điểm số, Đổi theme theo wave.",
                ]),
                ("3.2. Sequence Diagram", [
                    "Sequence Diagram của hệ thống nên thể hiện ít nhất ba luồng xử lý trọng tâm. Luồng thứ nhất mô tả quá trình đăng nhập PlayFab và nhập tên hiển thị, bắt đầu từ PlayFabLeaderboardManager, đi qua PlayFab và kết thúc ở NameInputPanel khi người chơi chưa có Display Name. Luồng thứ hai mô tả chuỗi chiến đấu chính gồm PlayerAttack tấn công Enemy, Enemy chết, ExpDropper cộng EXP, PlayerLevelSystem kiểm tra ngưỡng lên cấp rồi CardSelectionPanel phối hợp với BuffCardManager hiển thị và áp dụng buff. Luồng thứ ba mô tả game over và leaderboard, trong đó PlayerHealth lấy totalExpGained từ PlayerLevelSystem rồi gửi điểm thông qua PlayFabLeaderboardManager trước khi tải lại dữ liệu xếp hạng.",
                    "Gợi ý vẽ sơ đồ: dùng một hình sequence lớn và chia thành ba cụm dọc. Cụm 1 gồm các lifeline PlayFabLeaderboardManager, PlayFab, NameInputPanel, Người chơi. Cụm 2 gồm PlayerAttack, Enemy, ExpDropper, PlayerLevelSystem, CardSelectionPanel, BuffCardManager. Cụm 3 gồm Enemy hoặc Projectile, PlayerHealth, PlayerLevelSystem, PlayFabLeaderboardManager, PlayFab, LeaderboardPanel.",
                ]),
                ("3.3. Activity Diagram", [
                    "Activity Diagram của hệ thống cần mô tả vòng lặp gameplay chính từ lúc vào game tới khi kết thúc trận đấu. Luồng bắt đầu từ đăng nhập PlayFab, kiểm tra Display Name và nhập tên nếu cần. Sau đó người chơi tương tác với NPC để bắt đầu trận đấu, hệ thống spawn wave, chuyển sang chiến đấu, xử lý việc enemy chết, cộng EXP, kiểm tra lên cấp và mở màn hình chọn buff. Sau khi wave hoàn thành, hệ thống lặp lại quá trình với wave tiếp theo, có thể chèn boss wave hoặc chuyển theme map ở các mốc phù hợp. Luồng chỉ dừng khi người chơi chết, lúc này game gửi điểm lên PlayFab rồi chuyển sang trạng thái hiển thị leaderboard.",
                    "Gợi ý vẽ sơ đồ: dùng một activity flow theo trục từ trên xuống với các nút quyết định Đã có Display Name?, Enemy chết?, Đủ EXP lên cấp?, Hoàn thành wave? và Người chơi chết?. Các activity trung tâm là Khởi động game và đăng nhập PlayFab, Hiển thị NameInputPanel, Tương tác ChallengePostNPC và StartGame, WaveSpawner bắt đầu wave, Chiến đấu, ExpDropper cộng EXP, CardSelectionPanel chọn buff, Boss wave hoặc đổi theme map, Gửi điểm và tải leaderboard.",
                ]),
                ("3.4. Component Diagram", [
                    "Component Diagram cần phản ánh cách project được chia module trong triển khai hiện tại. Các component chính gồm Input, Player, Enemy, Wave, Buff, UI, Backend và Shared Services. Input tiếp nhận dữ liệu điều khiển và chuyển cho các hệ thống gameplay hoặc UI. Player bao gồm điều khiển, tấn công, máu, dữ liệu và cấp độ. Enemy quản lý đối tượng đối kháng và logic AI. Wave chịu trách nhiệm điều phối tiến trình spawn và nhịp độ trận đấu. Buff quản lý việc chọn và áp dụng thẻ tăng sức mạnh. UI cung cấp tất cả panel và HUD. Backend là lớp trung gian làm việc với PlayFab. Shared Services cung cấp các dịch vụ dùng chung như ObjectPool, AudioManager, MapThemeManager và LoadingUIManager.",
                    "Gợi ý vẽ sơ đồ: đặt 8 component thành 8 khối riêng. Input nối với Player và UI. Player tương tác hai chiều với Enemy. Wave điều phối Enemy và kích hoạt MapThemeManager khi cần. Buff nối với Player và CardSelectionPanel trong UI. Backend nối với PlayFabLeaderboardManager, NameInputPanel và LeaderboardPanel. Shared Services phục vụ Player, Enemy, Wave, Buff và UI.",
                ]),
                ("3.5. Class Diagram", [
                    "Class Diagram của hệ thống phải phản ánh rõ mô hình hướng đối tượng đang có trong project. Ở tầng nền, Singleton<T> là lớp cơ sở của các manager chính như PlayerController, PlayerLevelSystem, WaveSpawner, BuffCardManager, GameUI và PlayFabLeaderboardManager. Giao diện IDamageable được dùng để thống nhất xử lý sát thương cho PlayerHealth và Enemy. Trong nhóm gameplay, Enemy là lớp cha của MeleeEnemy, RangedEnemy, FlyEnemy và BossEnemy; riêng BossEnemy tiếp tục được mở rộng thành các boss nguyên tố. Projectile là lớp cha của PlayerProjectile, EnemyProjectile và SpiritProjectileScript. Về tương tác môi trường, NPC là lớp cha của ChallengePostNPC và ChestBuffBox. Về giao diện, PanelBase là lớp cơ sở của ChallengePanel, NameInputPanel, LeaderboardPanel, CardSelectionPanel và PauseMenuPanel.",
                    "Ngoài cấu trúc kế thừa, sơ đồ cũng phải thể hiện mối liên hệ giữa các lớp dữ liệu và lớp xử lý. PlayerData đọc thông số từ PlayerConfig, EnemyData đọc từ EnemyConfig, WaveSpawner phụ thuộc vào WaveConfig và BuffCardManager sử dụng BuffCardConfig để sinh và áp dụng buff. Cách tổ chức này cho thấy gameplay được triển khai theo hướng kết hợp giữa component-based của Unity và tư duy hướng đối tượng trong lập trình C#.",
                    "Gợi ý vẽ sơ đồ: đặt Singleton<T> ở trên cùng, các manager kế thừa nằm phía dưới. Đặt IDamageable ở một nhánh riêng và nối nét đứt tới PlayerHealth và Enemy. Vẽ cụm Enemy với các lớp con và cụm BossEnemy với các boss nguyên tố. Vẽ cụm Projectile với ba lớp dẫn xuất. Vẽ cụm UI quanh PanelBase và cụm dữ liệu quanh PlayerData, EnemyData, WaveConfig, BuffCardConfig cùng các quan hệ cấu hình tương ứng.",
                ]),
            ],
        ),
    ],
}


def build_text_file() -> Path:
    path = OUTPUT_DOC_DIR / "Chuong1_Chuong3_Roguelike_Revised.txt"
    parts: List[str] = [CHAPTER_TEXT["title"], ""]
    for chapter_title, entries in CHAPTER_TEXT["sections"]:
        parts.append(chapter_title)
        parts.append("")
        for heading, paragraphs in entries:
            if heading != chapter_title:
                parts.append(heading)
                parts.append("")
            for para in paragraphs:
                parts.append(para)
                parts.append("")
    path.write_text("\n".join(parts).strip() + "\n", encoding="utf-8")
    return path


def configure_document(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.98)
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(13)


def build_full_doc(diagram_paths: List[Tuple[str, Path]]) -> Path:
    doc = Document()
    configure_document(doc)
    add_heading(doc, CHAPTER_TEXT["title"], level=1)

    diagram_map = {name: path for name, path in diagram_paths}
    for chapter_title, entries in CHAPTER_TEXT["sections"]:
        add_heading(doc, chapter_title, level=1)
        for heading, paragraphs in entries:
            add_heading(doc, heading, level=2 if heading.startswith("1.") or heading.startswith("3.") else 3)
            for para in paragraphs:
                add_body(doc, para)
            if heading == "3.1. Use Case Diagram":
                doc.add_picture(str(diagram_map["3_1_use_case"]), width=Inches(6.5))
                add_diagram_caption(doc, "Hình 3.1. Use Case Diagram của hệ thống")
            elif heading == "3.2. Sequence Diagram":
                doc.add_picture(str(diagram_map["3_2_sequence"]), width=Inches(6.5))
                add_diagram_caption(doc, "Hình 3.2. Sequence Diagram của các luồng chính")
            elif heading == "3.3. Activity Diagram":
                doc.add_picture(str(diagram_map["3_3_activity"]), width=Inches(6.2))
                add_diagram_caption(doc, "Hình 3.3. Activity Diagram của vòng lặp gameplay")
            elif heading == "3.4. Component Diagram":
                doc.add_picture(str(diagram_map["3_4_component"]), width=Inches(6.5))
                add_diagram_caption(doc, "Hình 3.4. Component Diagram của hệ thống")
            elif heading == "3.5. Class Diagram":
                doc.add_picture(str(diagram_map["3_5_class"]), width=Inches(6.5))
                add_diagram_caption(doc, "Hình 3.5. Class Diagram của các lớp chính")

    output = OUTPUT_DOC_DIR / "Chuong1_Chuong3_Roguelike_Revised.docx"
    doc.save(output)
    return output


def build_diagram_doc(diagram_paths: List[Tuple[str, Path]]) -> Path:
    captions = {
        "3_1_use_case": "Hình 3.1. Use Case Diagram của hệ thống",
        "3_2_sequence": "Hình 3.2. Sequence Diagram của các luồng chính",
        "3_3_activity": "Hình 3.3. Activity Diagram của vòng lặp gameplay",
        "3_4_component": "Hình 3.4. Component Diagram của hệ thống",
        "3_5_class": "Hình 3.5. Class Diagram của các lớp chính",
    }
    doc = Document()
    configure_document(doc)
    add_heading(doc, "BỘ SƠ ĐỒ CHƯƠNG 3", level=1)
    for idx, (name, path) in enumerate(diagram_paths):
        if idx > 0:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
        add_heading(doc, captions[name], level=2)
        doc.add_picture(str(path), width=Inches(6.5))
        add_diagram_caption(doc, captions[name])
    output = OUTPUT_DOC_DIR / "SoDo_Chuong3_Roguelike.docx"
    doc.save(output)
    return output


def main() -> None:
    ensure_dirs()
    diagrams = [
        ("3_1_use_case", create_use_case_diagram()[0]),
        ("3_2_sequence", create_sequence_diagram()[0]),
        ("3_3_activity", create_activity_diagram()[0]),
        ("3_4_component", create_component_diagram()[0]),
        ("3_5_class", create_class_diagram()[0]),
    ]
    build_text_file()
    build_full_doc(diagrams)
    build_diagram_doc(diagrams)


if __name__ == "__main__":
    main()
