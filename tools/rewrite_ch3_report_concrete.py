from __future__ import annotations

import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_ch3_restructured_full.docx"
BACKUP_DOC = ROOT / "tmp" / "docs" / "BaoCao_DATN_2_ch3_before_concrete_rewrite.docx"

BLACK = RGBColor(0, 0, 0)


@dataclass
class TableSpec:
    caption: str
    headers: list[str]
    rows: list[list[str]]
    widths_cm: list[float] | None = None
    font_size: int = 11


@dataclass
class SectionSpec:
    heading: str
    next_heading: str
    figure_caption: str | None
    body: list[str]
    table: TableSpec | None = None


def set_run_font(run, size: int = 13, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_font(paragraph, size: int = 13, bold: bool | None = None, italic: bool | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, italic=italic)


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def replace_paragraph_text(
    paragraph,
    text: str,
    *,
    size: int = 13,
    bold: bool | None = None,
    italic: bool | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    if alignment is not None:
        paragraph.alignment = alignment


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def insert_table_after(
    paragraph,
    rows: int,
    cols: int,
    style: str = "Table Grid",
    width_cm: float = 16.0,
) -> Table:
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Cm(width_cm))
    paragraph._p.addnext(table._tbl)
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    return table


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_table(table: Table) -> None:
    table._tbl.getparent().remove(table._tbl)


def has_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//*[local-name()='drawing']"))


def find_paragraph(document: Document, text: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def paragraphs_between(document: Document, start_text: str, end_text: str):
    collecting = False
    items = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == start_text:
            collecting = True
            continue
        if collecting and text == end_text:
            break
        if collecting:
            items.append(paragraph)
    return items


def cleanup_section(document: Document, heading: str, next_heading: str):
    start = find_paragraph(document, heading)
    between = paragraphs_between(document, heading, next_heading)
    keep = []
    caption = None
    for paragraph in between:
        text = paragraph.text.strip()
        if text.startswith("Tệp sơ đồ nguồn:") or has_drawing(paragraph) or text.startswith("Hình "):
            keep.append(paragraph)
            if text.startswith("Hình ") and caption is None:
                caption = paragraph
            continue
        remove_paragraph(paragraph)
    anchor = keep[-1] if keep else start
    return anchor, caption


def remove_existing_chapter3_tables(document: Document) -> None:
    target_headers = {
        ("Thông số", "Giá trị", "Nguồn áp dụng runtime"),
        ("Hạng mục", "Giá trị hiện tại", "Ý nghĩa"),
        ("Loại enemy", "HP", "Tốc độ", "Tầm đánh", "Sát thương chính", "EXP"),
        ("Thành phần", "Giá trị cụ thể", "Ghi chú runtime"),
        ("Component", "Lớp/Service chính", "Trách nhiệm runtime"),
    }

    for paragraph in list(document.paragraphs):
        if paragraph.text.strip().startswith("Bảng 3."):
            remove_paragraph(paragraph)

    for table in list(document.tables):
        if not table.rows:
            continue
        header = tuple(cell.text.strip() for cell in table.rows[0].cells)
        if header in target_headers:
            remove_table(table)


def add_body_blocks(anchor, blocks: list[str]):
    current = anchor
    for text in blocks:
        current = insert_paragraph_after(current, text, style="Normal")
        set_paragraph_font(current, size=13)
        current.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return current


def add_table(anchor, spec: TableSpec):
    caption = insert_paragraph_after(anchor, spec.caption, style="Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(caption, size=12, bold=True)

    total_width = sum(spec.widths_cm) if spec.widths_cm else 16.0
    table = insert_table_after(
        caption,
        rows=1 + len(spec.rows),
        cols=len(spec.headers),
        width_cm=total_width,
    )
    header_cells = table.rows[0].cells
    for idx, text in enumerate(spec.headers):
        header_cells[idx].text = text

    for row_idx, row_values in enumerate(spec.rows, start=1):
        for col_idx, value in enumerate(row_values):
            table.rows[row_idx].cells[col_idx].text = value

    if spec.widths_cm:
        for row in table.rows:
            for idx, width in enumerate(spec.widths_cm):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=spec.font_size, bold=(row_idx == 0))

    return table


def remove_between(document: Document, start_text: str, end_text: str) -> None:
    for paragraph in list(paragraphs_between(document, start_text, end_text)):
        remove_paragraph(paragraph)


def normalize_chapter3_text_black(document: Document) -> None:
    in_chapter3 = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == "3.1. Use Case Diagram":
            in_chapter3 = True
        if text == "CHƯƠNG 4 \nXÂY DỰNG CHƯƠNG TRÌNH":
            in_chapter3 = False
        if not in_chapter3:
            continue
        for run in paragraph.runs:
            run.font.color.rgb = BLACK

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = BLACK


PLAYER_TABLE = TableSpec(
    caption="Bảng 3.1: Thông số player nền và combat cơ bản.",
    headers=["Thông số", "Giá trị", "Nguồn áp dụng runtime"],
    rows=[
        ["moveSpeed", "7", "PlayerConfig.asset"],
        ["dashSpeed", "20", "PlayerConfig.asset"],
        ["dashDuration", "0.5 s", "PlayerConfig.asset"],
        ["dashCooldown", "1 s", "PlayerConfig.asset"],
        ["maxHealth", "1000", "PlayerConfig.asset"],
        ["attackDamage", "50", "PlayerConfig.asset"],
        ["attackCooldown", "1 s", "PlayerConfig.asset"],
        ["projectileSpeed", "30", "PlayerConfig.asset"],
        ["projectileLifetime", "10 s", "PlayerConfig.asset"],
    ],
    widths_cm=[4.0, 3.0, 8.5],
)

PROGRESSION_TABLE = TableSpec(
    caption="Bảng 3.2: Thông số level-up và buff tiêu biểu.",
    headers=["Hạng mục", "Giá trị hiện tại", "Ý nghĩa"],
    rows=[
        ["currentLevel ban đầu", "0", "Level runtime trước khi vào trận"],
        ["expToNextLevel", "100", "Ngưỡng EXP cho level đầu tiên"],
        ["expScalingFactor", "1.1", "Hệ số tăng ngưỡng EXP sau mỗi level"],
        ["cardsPerSelection", "3", "Số thẻ buff hiển thị mỗi lần chọn"],
        ["BuffType trong code", "15", "Số loại buff được định nghĩa trong enum"],
        ["Asset buff hiện có", "13", "Số ScriptableObject thẻ hiện đang cấu hình"],
        ["RarityType", "4", "Common, Rare, Epic, Legendary"],
        ["MultiShot maxLevel", "3", "Có thể tăng lặp lại 3 mốc"],
        ["AoEExplosion maxLevel", "1", "Buff một lần rồi giữ trạng thái"],
    ],
    widths_cm=[4.4, 3.2, 7.9],
)

ENEMY_TABLE = TableSpec(
    caption="Bảng 3.3: So sánh chỉ số enemy đại diện.",
    headers=["Loại enemy", "HP", "Tốc độ", "Tầm đánh", "Sát thương chính", "EXP"],
    rows=[
        ["Melee", "100", "6", "2", "contactDamage = 10", "20"],
        ["Ranged", "100", "8", "20", "projectileDamage = 15", "50"],
        ["Fly", "100", "8", "50", "projectileDamage = 10", "50"],
        ["Boss (Geo đại diện)", "1000", "5", "50", "bossProjectileDamage = 25", "200"],
    ],
    widths_cm=[4.5, 1.6, 1.8, 1.9, 4.8, 1.6],
)

WAVE_TABLE = TableSpec(
    caption="Bảng 3.4: Cấu hình wave và spawn đại diện.",
    headers=["Thành phần", "Giá trị cụ thể", "Ghi chú runtime"],
    rows=[
        ["autoScale", "true", "Bật tăng độ khó theo wave"],
        ["scalePerWave", "1.1", "Nhân theo currentWave khi autoScale hoạt động"],
        ["spawnRandomRadius", "2", "Bán kính ngẫu nhiên quanh điểm spawn"],
        ["maxEnemySpawnsPerFrame", "8", "Giới hạn spawn mỗi frame"],
        ["maxSpawnEffectsPerFrame", "12", "Giới hạn effect spawn mỗi frame"],
        ["Wave 1 enemyGroups", "4 nhóm", "Cả 4 nhóm đều dùng pool type 3"],
        ["Wave 1 enemyCount", "3 enemy / nhóm", "Tổng 12 enemy ở wave đầu"],
        ["Wave 1 spawnDelay", "0, 1, 2, 3", "Spawn theo 4 nhịp liên tiếp"],
        ["Wave 1 preparationTime", "0", "Vào trận và spawn ngay"],
    ],
    widths_cm=[4.8, 3.4, 7.0],
)

COMPONENT_TABLE = TableSpec(
    caption="Bảng 3.5: Mapping component và service chính trong kiến trúc triển khai.",
    headers=["Component", "Lớp/Service chính", "Trách nhiệm runtime"],
    rows=[
        ["Input", "InputSystem_Actions, PlayerController", "Nhận input di chuyển, dash, tương tác và phân phối cho gameplay/UI"],
        ["Player", "PlayerData, PlayerAttack, PlayerHealth", "Giữ trạng thái chiến đấu và vòng đời của người chơi"],
        ["Enemy", "Enemy, MeleeEnemy, RangedEnemy, FlyEnemy, BossEnemy", "AI, nhận damage, tấn công và phát EXP khi chết"],
        ["Wave", "WaveSpawner, WaveConfig", "Khởi tạo wave, spawn enemy và chuyển wave"],
        ["Buff", "PlayerLevelSystem, BuffCardManager, CardSelectionPanel", "Tính EXP, level-up và áp buff vào runtime"],
        ["UI", "GameUI, PanelBase, PlayerStatsPanel", "HUD, panel nghiệp vụ và khóa/mở input đúng ngữ cảnh"],
        ["Backend", "PlayFabLeaderboardManager", "Login, cập nhật Display Name, gửi điểm và tải leaderboard"],
        ["Shared Services", "ObjectPool, AudioManager, MapThemeManager, LoadingUIManager", "Dịch vụ dùng chung cho hiệu năng, âm thanh và chuyển cảnh"],
    ],
    widths_cm=[3.2, 5.0, 7.0],
    font_size=10,
)


SECTION_SPECS = [
    SectionSpec(
        heading="3.1.1. Use Case tổng quan",
        next_heading="3.1.2. Use Case chi tiết",
        figure_caption="Use Case tổng quan của hệ thống",
        body=[
            "Mô tả: Sơ đồ use case tổng quan gom hai tác nhân chính của hệ thống là Người chơi và PlayFab. Người chơi là actor trung tâm của toàn bộ vòng lặp sinh tồn, còn PlayFab là actor ngoài hệ thống dùng để xác thực, cập nhật tên hiển thị, nhận điểm và trả về dữ liệu bảng xếp hạng.",
            "Luồng thực hiện: 1) Khi vào game, actor Người chơi khởi phát chuỗi đăng nhập. 2) Nếu chưa có Display Name, hệ thống mở luồng nhập tên. 3) Sau khi bắt đầu trận, người chơi đi qua nhánh gameplay gồm di chuyển, chiến đấu, nhận EXP, lên cấp và chọn buff. 4) Khi lượt chơi kết thúc, actor PlayFab tiếp nhận điểm số và trả dữ liệu leaderboard để hiển thị lại trên UI.",
            "Thông số cụ thể: Mức tổng quan hiện có 2 actor chính, 1 statistic backend là HighScore và 4 nhóm nghiệp vụ lớn gồm định danh, gameplay runtime, progression/buff và leaderboard. Tên hiển thị hợp lệ được kiểm tra trong khoảng 3 đến 25 ký tự trước khi đồng bộ lên PlayFab.",
            "Dữ liệu vào/ra: Đầu vào của actor Người chơi là input điều khiển, thao tác với NPC và quyết định chọn buff; đầu vào của actor PlayFab là Custom ID, Display Name và finalScore. Đầu ra của toàn hệ thống là currentWave, currentLevel, CurrentDisplayName, CurrentPlayFabId và vị trí leaderboard tương ứng.",
            "Kết quả đầu ra: Use case tổng quan cho thấy ranh giới rõ ràng giữa gameplay cục bộ trong scene và nhánh dữ liệu trực tuyến, từ đó làm nền cho các sequence và class diagram phía sau.",
        ],
    ),
    SectionSpec(
        heading="3.1.2. Use Case chi tiết",
        next_heading="3.2. Sequence Diagram",
        figure_caption="Use Case chi tiết của vòng lặp gameplay",
        body=[
            "Mô tả: Sơ đồ use case chi tiết đi sâu vào vòng lặp gameplay của một lượt chơi roguelike, nơi người chơi liên tục xen giữa chiến đấu, tích EXP, chọn buff, hoàn thành wave và xử lý trạng thái thất bại.",
            "Luồng thực hiện: 1) Người chơi tương tác ChallengePostNPC để khởi động trận. 2) WaveSpawner bắt đầu wave và spawn enemy theo cấu hình. 3) PlayerAttack tự động tìm mục tiêu, tạo projectile và gây damage. 4) Enemy chết sẽ trả EXP cho PlayerLevelSystem. 5) Khi đủ ngưỡng, CardSelectionPanel mở ra để người chơi chọn một trong 3 buff. 6) Chuỗi này lặp lại cho đến khi player bị hạ và hệ thống gửi điểm lên leaderboard.",
            "Thông số cụ thể: Vòng lặp chi tiết gắn trực tiếp với các biến runtime như currentWave, activeEnemies, currentExp, expToNextLevel, totalExpGained và currentRunScore. Điểm cuối trận hiện được tính bằng floor(totalExpGained), vì vậy tiến trình chiến đấu và tiến trình xếp hạng dùng chung một nguồn dữ liệu.",
            "Dữ liệu vào/ra: Đầu vào của use case chi tiết gồm trạng thái wave, enemy config, buff card pool và PlayFab profile hiện tại. Đầu ra là trạng thái HUD theo thời gian thực, bộ chỉ số player sau buff và kết quả gửi HighScore ở cuối lượt chơi.",
            "Kết quả đầu ra: Sơ đồ chi tiết làm rõ gameplay không chỉ là chuỗi chiến đấu, mà là một pipeline dữ liệu liên tục từ spawn quái, cộng EXP, chọn buff đến tổng kết điểm và đồng bộ leaderboard.",
        ],
    ),
    SectionSpec(
        heading="3.2.1. Luồng đăng nhập và nhập tên hiển thị",
        next_heading="3.2.2. Luồng bắt đầu trận và khởi tạo wave",
        figure_caption="Luồng đăng nhập và nhập tên hiển thị",
        body=[
            "Mô tả: Sequence này biểu diễn nhánh khởi tạo danh tính người chơi trước khi gameplay bắt đầu. Mục tiêu của nhánh là bảo đảm mỗi phiên chơi đều có CurrentPlayFabId và CurrentDisplayName hợp lệ để dùng lại cho leaderboard.",
            "Luồng thực hiện: 1) PlayFabLeaderboardManager gọi Login() ngay khi singleton khởi động. 2) Hệ thống lấy hoặc tạo Custom ID bằng GetOrCreateCustomId(), sau đó gửi LoginWithCustomID lên PlayFab. 3) Khi đăng nhập thành công, manager tiếp tục gọi GetPlayerProfile() để đọc Display Name hiện có. 4) Nếu tên trống, NameInputPanel được mở và SubmitName() sẽ cập nhật lại hồ sơ người chơi.",
            "Thông số cụ thể: Custom ID được lưu trong PlayerPrefs và mặc định lấy từ SystemInfo.deviceUniqueIdentifier; Display Name chỉ hợp lệ khi độ dài sau khi trim nằm trong khoảng 3 đến 25 ký tự; hai trường đồng bộ cốt lõi là CurrentDisplayName và CurrentPlayFabId.",
            "Dữ liệu vào/ra: Đầu vào là Custom ID cục bộ, phản hồi LoginResult/GetPlayerProfile và chuỗi tên người dùng nhập từ UI. Đầu ra là một danh tính PlayFab hợp lệ, sẵn sàng cho các thao tác SubmitScore() và GetLeaderboardData() sau này.",
            "Kết quả đầu ra: Sau sequence này, phần backend của client đã sẵn sàng, còn gameplay có thể đọc trực tiếp CurrentDisplayName để hiển thị tên người chơi trên các panel hoặc leaderboard.",
        ],
    ),
    SectionSpec(
        heading="3.2.2. Luồng bắt đầu trận và khởi tạo wave",
        next_heading="3.2.3. Luồng hệ thống spawn quái",
        figure_caption="Luồng bắt đầu trận và khởi tạo wave",
        body=[
            "Mô tả: Sequence bắt đầu trận chuyển game từ trạng thái chờ trong scene sang trạng thái gameplay runtime. Đây là điểm nối giữa nhóm NPC/UI và nhóm WaveSpawner.",
            "Luồng thực hiện: 1) Người chơi đi vào trigger của ChallengePostNPC và nhấn phím tương tác. 2) ChallengePanel mở ra, tạm khóa input gameplay và chờ người chơi xác nhận. 3) Khi nút Start Game được chọn, ChallengePanel phát sự kiện onGameStart, đồng thời ẩn panel và trả quyền hiển thị về HUD. 4) WaveSpawner.Instance.StartNextWave() được gọi để tăng currentWave và khởi tạo wave đầu tiên.",
            "Thông số cụ thể: Ở dữ liệu hiện tại của WaveConfig.asset, wave đầu có preparationTime = 0 nên sau khi onGameStart phát ra, quá trình spawn có thể diễn ra ngay. Ngoài ra ChallengePanel còn đóng vai trò nguồn event cho PlayerLevelSystem và PlayerStatsPanel để đồng bộ trạng thái đầu trận.",
            "Dữ liệu vào/ra: Đầu vào là trạng thái playerInRange của NPC, thao tác xác nhận trên ChallengePanel và instance của WaveSpawner. Đầu ra là currentWave mới, HUD hoạt động, NPC tương tác bị ẩn và vòng gameplay chính được kích hoạt.",
            "Kết quả đầu ra: Sau sequence này, trò chơi rời trạng thái chuẩn bị và đi vào nhịp chiến đấu theo wave mà không cần tải lại scene.",
        ],
    ),
    SectionSpec(
        heading="3.2.3. Luồng hệ thống spawn quái",
        next_heading="3.2.4. Luồng chiến đấu, nhận EXP và chọn buff",
        figure_caption="Luồng hệ thống spawn quái",
        body=[
            "Mô tả: Sequence spawn quái mô tả phần lõi của gameplay theo wave, nơi WaveSpawner đóng vai trò bộ điều phối trung tâm giữa ChallengePanel, WaveConfig và ObjectPool.",
            "Luồng thực hiện: 1) StartNextWave() tăng currentWave rồi lấy dữ liệu bằng GetWave(currentWave); nếu vượt quá tập wave cấu hình, hệ thống rẽ sang GenerateEndlessWave(currentWave). 2) RunWave() nhận SimpleWaveData, tính totalEnemiesToSpawn và chờ preparationTime trước khi spawn. 3) WaveSpawner duyệt từng EnemyGroup, gọi ObjectPool để sinh đúng pool type tại spawnPosition, sau đó cập nhật activeEnemies và phát OnEnemyCountChanged. 4) Khi wave bắt đầu chính thức, OnWaveStart(currentWave) được phát để HUD cập nhật trạng thái.",
            "Thông số cụ thể: WaveConfig.asset đang bật autoScale = true với scalePerWave = 1.1; WaveSpawner dùng spawnRandomRadius = 2, maxEnemySpawnsPerFrame = 8 và maxSpawnEffectsPerFrame = 12. Ở wave 1, dữ liệu hiện tại gồm 4 enemyGroups melee, mỗi nhóm 3 enemy, spawnDelay lần lượt là 0, 1, 2 và 3 giây, tổng cộng 12 enemy.",
            "Dữ liệu vào/ra: Đầu vào của sequence là waveConfig, currentWave, bossPoolTypes, spawnPosition và enemyGroups. Đầu ra là danh sách activeEnemies, totalEnemiesToSpawn, totalEnemiesSpawned và các event OnWaveStart/OnEnemyCountChanged phục vụ HUD và nhạc nền.",
            "Kết quả đầu ra: Sequence này chính là biểu diễn động của cấu trúc lớp được trình bày lại ở mục 3.5.7, nơi WaveConfig cung cấp dữ liệu, WaveSpawner điều phối, còn ObjectPool thực thi việc tạo đối tượng.",
        ],
    ),
    SectionSpec(
        heading="3.2.4. Luồng chiến đấu, nhận EXP và chọn buff",
        next_heading="3.2.5. Luồng game over, gửi điểm và leaderboard",
        figure_caption="Luồng chiến đấu, nhận EXP và chọn buff",
        body=[
            "Mô tả: Sequence này mô tả mạch gameplay lặp nhiều nhất trong game: player tự động tấn công, enemy nhận damage, EXP được cộng và hệ thống mở lựa chọn buff khi đủ điều kiện lên cấp.",
            "Luồng thực hiện: 1) PlayerAttack kiểm tra cooldown rồi tìm enemy gần nhất trong attackRange. 2) Khi có mục tiêu, ObjectPool cấp PlayerProjectile và hàm Initialize() nạp damage, speed, lifetime cùng các cờ AoE hoặc multishot. 3) Khi enemy chết, expValue được cộng vào PlayerLevelSystem qua AddExp(). 4) Nếu currentExp vượt expToNextLevel, hệ thống thực hiện LevelUp(), phát OnLevelUp và CardSelectionPanel hiển thị 3 thẻ để người chơi chọn.",
            "Thông số cụ thể: Bộ thông số combat hiện tại của player là attackDamage = 50, attackRange = 50, projectileSpeed = 30 và projectileLifetime = 10. Phần progression dùng expToNextLevel = 100, expScalingFactor = 1.1, cardsPerSelection = 3; dữ liệu mở rộng cho combat gồm multiShotCount mặc định 1, multiShotAngle = 12 và aoeRadius = 3.",
            "Dữ liệu vào/ra: Đầu vào của sequence là PlayerData, enemy trong phạm vi, expValue của từng enemy và card pool của BuffCardManager. Đầu ra là projectile đang hoạt động, currentExp/currentLevel mới, totalExpGained tăng lên và trạng thái buff được áp trực tiếp vào runtime.",
            "Kết quả đầu ra: Combat và progression dùng chung chuỗi dữ liệu, nên mỗi lần tiêu diệt enemy không chỉ giảm activeEnemies của wave mà còn có thể làm thay đổi hẳn bộ chỉ số chiến đấu của player ở những nhịp tiếp theo.",
        ],
    ),
    SectionSpec(
        heading="3.2.5. Luồng game over, gửi điểm và leaderboard",
        next_heading="3.3. Activity Diagram",
        figure_caption="Luồng game over, gửi điểm và leaderboard",
        body=[
            "Mô tả: Sequence cuối cùng mô tả nhánh kết thúc một lượt chơi, nơi gameplay cục bộ nhường quyền cho quá trình dọn scene, tính điểm và đồng bộ dữ liệu với PlayFab.",
            "Luồng thực hiện: 1) PlayerHealth phát hiện currentHealth <= 0 và gọi Die(). 2) Hệ thống khóa input, dọn object gameplay còn lại thông qua StopAndClearCurrentWave() và chuẩn bị leaderboard sau khi chết. 3) currentRunScore được tính bằng floor(PlayerLevelSystem.totalExpGained). 4) PlayFabLeaderboardManager.SubmitScore() gửi giá trị này lên statistic HighScore, sau đó GetLeaderboardData() và GetPlayerLeaderboardData() nạp lại bảng xếp hạng để hiển thị.",
            "Thông số cụ thể: Statistic dùng để gửi điểm là HighScore; GetLeaderboardData() lấy tối đa 100 kết quả đầu bảng; điểm gửi đi không lấy từ thời gian sống mà lấy trực tiếp từ totalExpGained của lượt chơi. Trong PlayerHealth, chuỗi dọn dẹp còn có các mốc thời gian nội bộ như DeathCleanupDelaySeconds = 1 và DeathLeaderboardDelaySeconds = 0.5.",
            "Dữ liệu vào/ra: Đầu vào là currentHealth, totalExpGained, CurrentPlayFabId và CurrentDisplayName. Đầu ra là điểm số đã lưu trên PlayFab, top leaderboard, thứ hạng gần người chơi và trạng thái kết thúc trận trên UI.",
            "Kết quả đầu ra: Sequence này khép vòng gameplay theo cùng một chuẩn dữ liệu backend, bảo đảm điểm hiển thị trên leaderboard đúng với tiến trình chiến đấu mà người chơi vừa hoàn thành.",
        ],
    ),
    SectionSpec(
        heading="3.3. Activity Diagram",
        next_heading="3.4. Component Diagram",
        figure_caption="Activity Diagram của vòng lặp gameplay",
        body=[
            "Mô tả: Activity Diagram mô tả toàn bộ vòng đời của một phiên chơi từ bước đăng nhập, chuẩn bị trận, chiến đấu theo wave cho đến lúc gửi điểm và quay về leaderboard.",
            "Luồng thực hiện: 1) Game đăng nhập PlayFab và kiểm tra Display Name. 2) Nếu hợp lệ, người chơi tương tác với ChallengePostNPC để phát onGameStart. 3) WaveSpawner bắt đầu chuỗi spawn và chiến đấu lặp. 4) Mỗi lần tiêu diệt enemy, luồng có thể rẽ sang nhánh AddExp() và CardSelectionPanel khi đủ điều kiện lên cấp. 5) Sau khi người chơi bị hạ, luồng chuyển sang dọn scene, SubmitScore() và tải lại leaderboard.",
            "Thông số cụ thể: Nhánh level-up đầu tiên dùng expToNextLevel = 100 và các level sau tăng theo expScalingFactor = 1.1. Nhánh spawn dùng autoScale = true với scalePerWave = 1.1; điểm số cuối trận dùng currentRunScore = floor(totalExpGained). Điều này cho thấy ba trục gameplay, progression và leaderboard đang dùng chung dữ liệu runtime thay vì tách rời.",
            "Dữ liệu vào/ra: Đầu vào của activity là input người chơi, profile PlayFab, cấu hình wave và bộ chỉ số player/enemy. Đầu ra là trạng thái wave, trạng thái HUD, bộ buff đang hoạt động và kết quả HighScore của lượt chơi hiện tại.",
            "Kết quả đầu ra: Sơ đồ hoạt động làm rõ game được tổ chức như một vòng lặp trạng thái khép kín, trong đó mỗi nút quyết định đều gắn với dữ liệu cụ thể trong code thay vì chỉ là mô tả khái quát.",
        ],
    ),
    SectionSpec(
        heading="3.4. Component Diagram",
        next_heading="3.5. Class Diagram",
        figure_caption="Component Diagram của hệ thống",
        body=[
            "Mô tả: Component Diagram trình bày cách project được tách thành các mô-đun triển khai rõ ràng thay vì để toàn bộ logic dồn vào một cụm script duy nhất. Trên build hiện tại có thể quy về 8 component chính: Input, Player, Enemy, Wave, Buff, UI, Backend và Shared Services.",
            "Luồng thực hiện: Input thu thao tác và chuyển vào Player hoặc UI; Player phối hợp với Enemy và Projectile để tạo combat loop; Wave điều phối việc spawn và hoàn tất từng wave; Buff nhận dữ liệu từ PlayerLevelSystem để thay đổi PlayerData; cuối cùng Backend tiếp nhận currentRunScore và đồng bộ với PlayFab. Shared Services như ObjectPool, AudioManager hay MapThemeManager được nhiều component dùng chung ở mọi pha runtime.",
            "Thông số cụ thể: Các singleton đang đóng vai trò hub gồm GameUI, WaveSpawner, PlayerLevelSystem và PlayFabLeaderboardManager. Dòng dữ liệu giữa component được ràng qua những event hoặc API chính như onGameStart, OnWaveStart, OnEnemyCountChanged, OnLevelUp và SubmitScore().",
            "Dữ liệu vào/ra: Đầu vào của kiến trúc component là input người chơi, dữ liệu ScriptableObject, profile PlayFab và các event runtime. Đầu ra là scene gameplay đang chạy, HUD đang phản ánh đúng trạng thái và leaderboard đã được đồng bộ khi lượt chơi kết thúc.",
            "Kết quả đầu ra: Từ góc nhìn triển khai, sơ đồ component cho thấy project đã được mô-đun hóa đủ rõ để có thể mở rộng enemy, wave, buff hoặc backend mà không cần viết lại toàn bộ vòng lặp gameplay.",
        ],
        table=COMPONENT_TABLE,
    ),
    SectionSpec(
        heading="3.5.1. Nhóm Player Core",
        next_heading="3.5.2. Nhóm Progression và Buff",
        figure_caption="Class Diagram - Nhóm Player Core",
        body=[
            "Mô tả: Nhóm Player Core gom bốn lớp điều khiển trực tiếp vòng đời người chơi là PlayerController, PlayerAttack, PlayerHealth và PlayerData. Đây là cụm lớp trung tâm của gameplay runtime vì toàn bộ trạng thái di chuyển, tấn công và máu đều hội tụ tại đây.",
            "Cách vận hành: PlayerController nhận input và đọc PlayerData để xử lý di chuyển, gravity, xoay theo camera và dash. PlayerAttack sử dụng cùng PlayerData để xác định attackRange, attackDamage, projectileSpeed rồi tạo projectile thông qua ObjectPool. PlayerHealth theo dõi currentHealth, phát OnHealthChanged/OnDeath và gọi chuỗi game over khi HP về 0.",
            "Thông số cụ thể: Dữ liệu hiện hành trong PlayerConfig.asset là moveSpeed = 7, dashSpeed = 20, dashDuration = 0.5, dashCooldown = 1, maxHealth = 1000, attackDamage = 50, projectileSpeed = 30 và projectileLifetime = 10. Đây là bộ tham số thật đang nạp vào runtime qua PlayerData.LoadFromConfig().",
            "Dữ liệu vào/ra: Đầu vào của cụm này là moveInput, trạng thái dash, enemy target và mọi buff đã áp vào PlayerData. Đầu ra là vị trí nhân vật, projectile đang bay, currentHealth/maxHealth trên HUD và trạng thái sống/chết để các hệ khác tiếp nhận.",
            "Kết quả đầu ra: Việc tách PlayerData khỏi ba lớp xử lý còn lại giúp tất cả thay đổi chỉ số do buff hoặc config đều được áp tập trung, tránh việc một giá trị combat bị nhân bản ở nhiều nơi.",
        ],
        table=PLAYER_TABLE,
    ),
    SectionSpec(
        heading="3.5.2. Nhóm Progression và Buff",
        next_heading="3.5.3. Nhóm Enemy Core",
        figure_caption="Class Diagram - Nhóm Progression và Buff",
        body=[
            "Mô tả: Nhóm này chịu trách nhiệm cho tiến trình tăng sức mạnh của người chơi qua EXP, level và buff card. Cụm lớp chính gồm PlayerLevelSystem, BuffCardManager, CardSelectionPanel và PlayerStatsPanel.",
            "Cách vận hành: PlayerLevelSystem cộng EXP bằng AddExp(), tự kiểm tra currentExp >= expToNextLevel rồi thực hiện LevelUp(). Sau khi level tăng, hệ thống phát OnExpChanged, OnLevelChanged và OnLevelUp; CardSelectionPanel đăng ký OnLevelUp để gọi BuffCardManager.SelectCards(), còn PlayerStatsPanel cập nhật thanh HUD theo dữ liệu mới. Ở code hiện tại, ChallengePanel.onGameStart còn được gắn với LevelUp() để mở mốc tiến trình đầu trận.",
            "Thông số cụ thể: Hệ thống hiện bắt đầu ở currentLevel = 0, expToNextLevel = 100 và expScalingFactor = 1.1. BuffCardManager hiển thị 3 thẻ mỗi lần chọn; enum BuffType trong code có 15 loại, thư mục asset hiện cấu hình 13 thẻ, và rarity được chia thành 4 mức Common, Rare, Epic, Legendary. Một số mốc điển hình là MultiShot maxLevel = 3 và AoEExplosion maxLevel = 1.",
            "Dữ liệu vào/ra: Đầu vào là expValue từ enemy, luckBonus trong PlayerData và danh sách buff card chưa vượt maxLevel. Đầu ra là currentLevel, currentExp, totalExpGained và các modifier mới áp thẳng vào PlayerData hoặc PlayerHealth.",
            "Kết quả đầu ra: Nhờ cơ chế event, nhóm Progression và Buff không cần gọi cứng sang gameplay mà vẫn thay đổi tức thời được cách người chơi tấn công, di chuyển hoặc hồi phục ở các wave tiếp theo.",
        ],
        table=PROGRESSION_TABLE,
    ),
    SectionSpec(
        heading="3.5.3. Nhóm Enemy Core",
        next_heading="3.5.4. Nhóm Projectile và Damage Flow",
        figure_caption="Class Diagram - Nhóm Enemy Core",
        body=[
            "Mô tả: Enemy Core mô tả trục kế thừa của đối tượng địch từ lớp Enemy cơ sở sang MeleeEnemy, RangedEnemy, FlyEnemy và BossEnemy, đồng thời chỉ ra vai trò của EnemyConfig, EnemyData và giao diện IDamageable.",
            "Cách vận hành: EnemyConfig giữ số liệu thiết kế; EnemyData nhận cấu hình này rồi cung cấp dữ liệu runtime cho từng đối tượng. Enemy xử lý di chuyển, nhận damage, phát OnDeath và tương tác với player hoặc projectile. Các lớp con bổ sung chiến thuật riêng như tiếp cận cận chiến, giữ khoảng cách bắn, bay theo quỹ đạo hoặc chuyển phase ở boss.",
            "Thông số cụ thể: Bộ dữ liệu hiện tại cho thấy MeleeEnemy có 100 HP, moveSpeed = 6, attackRange = 2, contactDamage = 10 và expValue = 20; RangedEnemy có 100 HP, moveSpeed = 8, attackRange = 20, projectileDamage = 15 và expValue = 50; FlyEnemy có 100 HP, moveSpeed = 8, attackRange = 50, projectileDamage = 10 và expValue = 50. Các boss đại diện dùng maxHealth = 1000 với ngưỡng chuyển phase 0.6 và 0.3.",
            "Dữ liệu vào/ra: Đầu vào của cụm này là EnemyConfig tương ứng, spawn position từ WaveSpawner và damage tới từ IDamageable.TakeDamage(). Đầu ra là enemy đã hoạt động đúng AI, expValue trả về khi chết và event OnDeath phục vụ progression cũng như wave runtime.",
            "Kết quả đầu ra: Sơ đồ cho thấy boss không phải hệ riêng tách biệt mà là lớp kế thừa mở rộng, nhờ đó toàn bộ pipeline spawn, damage và despawn vẫn đi theo cùng một chuẩn đối tượng.",
        ],
        table=ENEMY_TABLE,
    ),
    SectionSpec(
        heading="3.5.4. Nhóm Projectile và Damage Flow",
        next_heading="3.5.5. Nhóm UI và Scene Interaction",
        figure_caption="Class Diagram - Nhóm Projectile và Damage Flow",
        body=[
            "Mô tả: Nhóm Projectile và Damage Flow tách riêng đường đi của sát thương ra khỏi AI enemy hoặc input player. Lõi của nhóm là lớp Projectile cùng ba nhánh PlayerProjectile, EnemyProjectile và SpiritProjectileScript.",
            "Cách vận hành: Projectile.Initialize(damage, speed, lifetime, direction, targetLayer, owner) nạp mọi tham số cần dùng rồi tự di chuyển theo Time.deltaTime. Khi va chạm, OnHit() gọi IDamageable.TakeDamage() và DispawnProjectile() để trả đối tượng về pool. PlayerProjectile bổ sung nhánh AoE, còn SpiritProjectileScript mở rộng cho các buff hệ spirit như xuyên mục tiêu hoặc nổ diện rộng.",
            "Thông số cụ thể: Bộ thông số mặc định của player projectile hiện là projectileSpeed = 30 và projectileLifetime = 10; các modifier runtime đáng chú ý là multiShotCount mặc định 1, multiShotAngle = 12, aoeRadius = 3 và aoeAtkMultiplier tương ứng với buff AoE. Ở boss, projectileSpeed nằm trong khoảng 12 đến 15 và projectileLifetime là 6 giây tùy config.",
            "Dữ liệu vào/ra: Đầu vào là owner, targetLayer, hướng bắn và các modifier chiến đấu lấy từ PlayerData hoặc EnemyData. Đầu ra là damage hợp lệ đẩy sang IDamageable, danh sách collider bị dính AoE và trạng thái despawn đúng lúc để không rò rỉ object pool.",
            "Kết quả đầu ra: Cụm lớp này bảo đảm mọi nguồn sát thương trong game đi qua một khuôn dạng chung, nhờ vậy việc mở rộng vũ khí mới hoặc projectile mới không làm rối logic nhận damage của toàn hệ thống.",
        ],
    ),
    SectionSpec(
        heading="3.5.5. Nhóm UI và Scene Interaction",
        next_heading="3.5.6. Nhóm Backend và Leaderboard",
        figure_caption="Class Diagram - Nhóm UI và Scene Interaction",
        body=[
            "Mô tả: Nhóm này gom phần giao diện trong trận và phần tương tác ngoài không gian scene. Trục UI dùng PanelBase và GameUI làm nền, còn trục scene interaction dùng NPC làm lớp gốc cho ChallengePostNPC và ChestBuffBox.",
            "Cách vận hành: PanelBase chuẩn hóa Show()/Hide() bằng CanvasGroup và DOTween, còn GameUI giữ tham chiếu tập trung đến các panel như ChallengePanel, NameInputPanel, CardSelectionPanel, PlayerStatsPanel và PauseMenuPanel. Ở nhánh scene, NPC kiểm tra playerInRange, lắng nghe phím F và gọi Interact(); khi panel liên quan đóng, OnPanelClosed() sẽ phục hồi prompt nếu người chơi vẫn đứng trong vùng tương tác.",
            "Thông số cụ thể: PanelBase hiện dùng showDuration = 0.3 giây và hideDuration = 0.1 giây; phím tương tác của NPC là F; ChallengePanel có hai event tĩnh là OnClosed và onGameStart; LeaderboardPanel cũng phát OnClosed để PauseMenuPanel xử lý quay về trạng thái trước đó.",
            "Dữ liệu vào/ra: Đầu vào của nhóm là input UI, trigger collider, trạng thái playerInRange và các event gameplay như onGameStart. Đầu ra là HUD đang mở đúng ngữ cảnh, gameplay input bị khóa hoặc trả lại đúng lúc và scene interaction luôn đồng bộ với trạng thái panel.",
            "Kết quả đầu ra: Việc đặt GameUI làm hub và NPC làm base class giúp phần UI/interaction mở rộng được thêm panel hoặc thêm đối tượng tương tác mà không làm đứt mạch điều phối input của gameplay.",
        ],
    ),
    SectionSpec(
        heading="3.5.6. Nhóm Backend và Leaderboard",
        next_heading="3.5.7. Nhóm Spawn Enemy và Wave Runtime",
        figure_caption="Class Diagram - Nhóm Backend và Leaderboard",
        body=[
            "Mô tả: Nhóm Backend và Leaderboard biểu diễn lớp kết nối duy nhất với PlayFab là PlayFabLeaderboardManager cùng hai panel giao diện bám theo nó là NameInputPanel và LeaderboardPanel.",
            "Cách vận hành: PlayFabLeaderboardManager thực hiện Login(), GetPlayerProfile(), SubmitName(), SubmitScore(), GetLeaderboardData() và GetPlayerLeaderboardData(). NameInputPanel chỉ chịu trách nhiệm nhập liệu và gọi SubmitName(), còn LeaderboardPanel nhận dữ liệu từ manager để dựng entry, đánh dấu người chơi hiện tại và xử lý hiển thị/ẩn panel trong nhiều ngữ cảnh như pause hoặc game over.",
            "Thông số cụ thể: Statistic dùng để gửi điểm là LeaderboardStatisticName = HighScore; Display Name hợp lệ nằm trong khoảng 3 đến 25 ký tự; Custom ID được lấy từ PlayerPrefs hoặc SystemInfo.deviceUniqueIdentifier nếu chưa có dữ liệu cục bộ. GetLeaderboardData() đang tải tối đa 100 dòng, còn GetLeaderboardAroundPlayer() lấy 1 dòng quanh người chơi hiện tại.",
            "Dữ liệu vào/ra: Đầu vào là PlayFabTitleId, Custom ID, CurrentDisplayName, CurrentPlayFabId và finalScore của lượt chơi. Đầu ra là hồ sơ PlayFab đã xác thực, leaderboard tổng, leaderboard quanh người chơi và trạng thái giao diện bảng xếp hạng.",
            "Kết quả đầu ra: Nhóm backend được cô lập tốt khỏi gameplay, nên thay đổi cách lưu điểm hoặc hiển thị leaderboard sẽ không làm ảnh hưởng trực tiếp tới logic combat hay wave runtime.",
        ],
    ),
    SectionSpec(
        heading="3.5.7. Nhóm Spawn Enemy và Wave Runtime",
        next_heading="CHƯƠNG 4 \nXÂY DỰNG CHƯƠNG TRÌNH",
        figure_caption="Class Diagram - Nhóm Spawn Enemy và Wave Runtime",
        body=[
            "Mô tả: Nhóm Spawn Enemy và Wave Runtime mô tả đúng cụm lớp hiện thực sequence spawn ở mục 3.2.3, gồm WaveSpawner, WaveConfig, SimpleWaveData, EnemyGroup và ObjectPool.",
            "Cách vận hành: WaveConfig cung cấp danh sách waves, mỗi wave là một SimpleWaveData gồm nhiều EnemyGroup và các cờ như isBossWave hoặc preparationTime. WaveSpawner gọi StartNextWave(), lấy wave bằng GetWave() hoặc GenerateEndlessWave(), sau đó dùng RunWave() để spawn từng group qua ObjectPool. Trong toàn bộ quá trình, spawner duy trì currentWave, activeEnemies, totalEnemiesToSpawn và phát các event OnWaveStart, OnWaveComplete, OnEnemyCountChanged.",
            "Thông số cụ thể: Dữ liệu hiện tại đang bật autoScale = true với scalePerWave = 1.1. WaveSpawner dùng spawnRandomRadius = 2, maxEnemySpawnsPerFrame = 8 và maxSpawnEffectsPerFrame = 12. Wave 1 trong WaveConfig.asset gồm 4 nhóm melee, mỗi nhóm 3 enemy, spawnDelay lần lượt là 0, 1, 2 và 3 giây, preparationTime = 0. Đây là cụm thông số thể hiện rõ nhất việc cấu hình thiết kế đã được nối trực tiếp vào runtime spawn.",
            "Dữ liệu vào/ra: Đầu vào của nhóm là waveConfig, enemyPoolType, spawnPosition, spreadRadius, spawnDelay và bossPoolTypes. Đầu ra là danh sách enemy đang hoạt động, số lượng enemy còn sống trên HUD, trạng thái hoàn tất wave và các session spawn tiếp theo của trận đấu.",
            "Kết quả đầu ra: Sơ đồ lớp này là phần bổ sung cấu trúc cho sequence spawn, chứng minh rằng luồng ChallengePanel -> WaveSpawner -> WaveConfig -> EnemyGroup -> ObjectPool trong mục 3.2.3 đã được hiện thực đầy đủ bằng các lớp dữ liệu và lớp điều phối riêng biệt.",
        ],
        table=WAVE_TABLE,
    ),
]


def rewrite_35_intro(document: Document) -> None:
    start = "3.5. Class Diagram"
    end = "3.5.1. Nhóm Player Core"
    remove_between(document, start, end)
    anchor = find_paragraph(document, start)
    intro = (
        "Phần Class Diagram được viết lại theo 7 nhóm lớp bám trực tiếp vào mã nguồn hiện tại thay vì gộp ba sơ đồ lớn như bản cũ. "
        "Mỗi nhóm giữ một sơ đồ riêng, phần mô tả runtime bằng chữ và khi cần có thêm bảng thông số để người đọc đối chiếu nhanh giữa cấu trúc lớp và dữ liệu đang dùng trong game."
    )
    para = insert_paragraph_after(anchor, intro, style="Normal")
    set_paragraph_font(para, size=13)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def rewrite_sections(document: Document) -> None:
    figure_index = 1
    for spec in SECTION_SPECS:
        anchor, caption = cleanup_section(document, spec.heading, spec.next_heading)
        if spec.figure_caption and caption is not None:
            replace_paragraph_text(
                caption,
                f"Hình 3.{figure_index}: {spec.figure_caption}.",
                size=12,
                italic=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
            figure_index += 1

        body_anchor = caption if caption is not None else anchor
        body_end = add_body_blocks(body_anchor, spec.body)
        if spec.table is not None:
            add_table(body_end, spec.table)


def validate_document(document: Document) -> None:
    headings_to_check = [
        "3.1. Use Case Diagram",
        "3.1.1. Use Case tổng quan",
        "3.1.2. Use Case chi tiết",
        "3.2. Sequence Diagram",
        "3.2.1. Luồng đăng nhập và nhập tên hiển thị",
        "3.2.2. Luồng bắt đầu trận và khởi tạo wave",
        "3.2.3. Luồng hệ thống spawn quái",
        "3.2.4. Luồng chiến đấu, nhận EXP và chọn buff",
        "3.2.5. Luồng game over, gửi điểm và leaderboard",
        "3.3. Activity Diagram",
        "3.4. Component Diagram",
        "3.5. Class Diagram",
        "3.5.1. Nhóm Player Core",
        "3.5.2. Nhóm Progression và Buff",
        "3.5.3. Nhóm Enemy Core",
        "3.5.4. Nhóm Projectile và Damage Flow",
        "3.5.5. Nhóm UI và Scene Interaction",
        "3.5.6. Nhóm Backend và Leaderboard",
        "3.5.7. Nhóm Spawn Enemy và Wave Runtime",
    ]
    for heading in headings_to_check:
        find_paragraph(document, heading)

    figure_captions = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("Hình 3.")
    ]
    table_captions = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("Bảng 3.")
    ]

    expected_figure_count = 16
    expected_table_count = 5
    if len(figure_captions) != expected_figure_count:
        raise ValueError(f"Expected {expected_figure_count} figure captions, found {len(figure_captions)}")
    if len(table_captions) != expected_table_count:
        raise ValueError(f"Expected {expected_table_count} table captions, found {len(table_captions)}")

    non_black = 0
    in_chapter3 = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == "3.1. Use Case Diagram":
            in_chapter3 = True
        if text == "CHƯƠNG 4 \nXÂY DỰNG CHƯƠNG TRÌNH":
            in_chapter3 = False
        if not in_chapter3:
            continue
        for run in paragraph.runs:
            if run.font.color.rgb not in (None, BLACK):
                non_black += 1
    if non_black:
        raise ValueError(f"Found {non_black} non-black runs in chapter 3")


def main() -> int:
    if not OUTPUT_DOC.exists():
        raise FileNotFoundError(f"Output document not found: {OUTPUT_DOC}")

    BACKUP_DOC.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(OUTPUT_DOC, BACKUP_DOC)

    document = Document(OUTPUT_DOC)
    remove_existing_chapter3_tables(document)
    rewrite_35_intro(document)
    rewrite_sections(document)
    normalize_chapter3_text_black(document)
    validate_document(document)
    document.save(OUTPUT_DOC)

    print(f"Updated: {OUTPUT_DOC}")
    print(f"Backup: {BACKUP_DOC}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
