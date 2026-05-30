from __future__ import annotations

import html
import math
import shutil
import uuid
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SNAPSHOT_DOC = ROOT / "tmp" / "docs" / "BaoCao_DATN_2_onedrive_snapshot.docx"
OUTPUT_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_ch3_restructured_full.docx"
DIAGRAM_DRAWIO_DIR = ROOT / "output" / "diagrams" / "drawio"
DIAGRAM_PNG_DIR = ROOT / "output" / "diagrams" / "png"
INDEX_PATH = DIAGRAM_DRAWIO_DIR / "INDEX.md"
ONEDRIVE_DOC = Path(r"C:\Users\haov8\OneDrive\Tài liệu\BaoCao_DATN (2).docx")

FONT_REGULAR = Path(r"C:\Windows\Fonts\times.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\timesbd.ttf")

BG = "#ffffff"
BOX_FILL = "#fff2cc"
BOX_OUTLINE = "#b45f06"
TERM_FILL = "#d9ead3"
TERM_OUTLINE = "#38761d"
NOTE_FILL = "#e8f0fe"
NOTE_OUTLINE = "#3c78d8"
TEXT = "#222222"
EDGE = "#333333"


@dataclass
class Node:
    id: str
    text: str
    x: int
    y: int
    w: int = 220
    h: int = 80
    kind: str = "box"  # box, terminator, note
    attributes: list[str] | None = None
    methods: list[str] | None = None


@dataclass
class Edge:
    source: str
    target: str
    label: str = ""
    dashed: bool = False


@dataclass
class DiagramSpec:
    filename: str
    report_section: str
    caption: str
    page_name: str
    width: int
    height: int
    nodes: list[Node]
    edges: list[Edge]


@dataclass
class ReportItem:
    heading: str
    filename: str
    caption: str
    paragraphs: list[str]


BOX_STYLE = (
    "whiteSpace=wrap;html=1;align=left;verticalAlign=top;spacing=10;"
    "fontFamily=Times New Roman;fontSize=24;strokeWidth=2;rounded=0;"
    "strokeColor=#000000;fillColor=#ffffff;fontColor=#000000;"
)
TERM_STYLE = (
    "whiteSpace=wrap;html=1;align=center;verticalAlign=middle;"
    "fontFamily=Times New Roman;fontSize=24;strokeWidth=2;rounded=0;"
    "strokeColor=#000000;fillColor=#ffffff;fontColor=#000000;"
)
NOTE_STYLE = (
    "whiteSpace=wrap;html=1;align=center;verticalAlign=middle;"
    "fontFamily=Times New Roman;fontSize=24;strokeWidth=2;rounded=0;"
    "strokeColor=#000000;fillColor=#ffffff;fontColor=#000000;"
)
TITLE_STYLE = (
    "text;html=1;strokeColor=none;fillColor=none;align=left;verticalAlign=middle;"
    "fontFamily=Times New Roman;fontSize=24;fontStyle=1;"
)
EDGE_STYLE = (
    "edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;"
    "endArrow=block;endFill=1;strokeColor=#000000;fontColor=#000000;fontSize=12;"
    "fontFamily=Times New Roman;"
)
DASHED_EDGE_STYLE = (
    "edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;"
    "endArrow=open;dashed=1;strokeColor=#000000;fontColor=#000000;fontSize=12;"
    "fontFamily=Times New Roman;"
)


def n(
    id_: str,
    text: str,
    x: int,
    y: int,
    w: int = 220,
    h: int = 80,
    kind: str = "box",
    attributes: list[str] | None = None,
    methods: list[str] | None = None,
) -> Node:
    return Node(id=id_, text=text, x=x, y=y, w=w, h=h, kind=kind, attributes=attributes, methods=methods)


def uml(
    id_: str,
    text: str,
    x: int,
    y: int,
    attributes: list[str],
    methods: list[str],
    w: int = 320,
    h: int = 240,
) -> Node:
    return Node(id=id_, text=text, x=x, y=y, w=w, h=h, kind="uml", attributes=attributes, methods=methods)


def e(src: str, tgt: str, label: str = "", dashed: bool = False) -> Edge:
    return Edge(source=src, target=tgt, label=label, dashed=dashed)


def esc(text: str) -> str:
    return html.escape(text).replace("\n", "&lt;br&gt;")


def style_for(kind: str) -> str:
    return {"box": BOX_STYLE, "uml": BOX_STYLE, "terminator": TERM_STYLE, "note": NOTE_STYLE}[kind]


def node_display_text(node: Node) -> str:
    return node.text


def uml_table_html(node: Node) -> str:
    attrs = "<br>".join(html.escape(item) for item in (node.attributes or ["-"]))
    methods = "<br>".join(html.escape(item) for item in (node.methods or ["-"]))
    title = html.escape(node.text)
    return (
        "<table style='width:100%;height:100%;border-collapse:collapse;'>"
        "<tr><td style='text-align:center;border-bottom:1px solid #000;padding:6px;'>"
        f"<b>{title}</b>"
        "</td></tr>"
        "<tr><td style='border-bottom:1px solid #000;padding:8px;'>"
        "<b>Thuộc tính</b><br>"
        f"{attrs}"
        "</td></tr>"
        "<tr><td style='padding:8px;'>"
        "<b>Phương thức</b><br>"
        f"{methods}"
        "</td></tr>"
        "</table>"
    )


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidate = FONT_BOLD if bold else FONT_REGULAR
    if candidate.exists():
        return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def build_diagram_xml(spec: DiagramSpec) -> str:
    parts = [
        '<mxfile host="app.diagrams.net" modified="2026-05-21T00:00:00.000Z" agent="Codex" version="24.7.17">',
        f'  <diagram id="{uuid.uuid4().hex[:12]}" name="{esc(spec.page_name)}">',
        (
            f'    <mxGraphModel dx="{spec.width}" dy="{spec.height}" grid="1" gridSize="10" '
            f'guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" '
            f'pageWidth="{spec.width}" pageHeight="{spec.height}" math="0" shadow="0">'
        ),
        "      <root>",
        '        <mxCell id="0" />',
        '        <mxCell id="1" parent="0" />',
        (
            f'        <mxCell id="title" value="{esc(spec.page_name)}" '
            f'style="{TITLE_STYLE}" vertex="1" parent="1">'
        ),
        f'          <mxGeometry x="80" y="40" width="{max(1200, spec.width - 160)}" height="40" as="geometry" />',
        "        </mxCell>",
    ]
    for node in spec.nodes:
        value = esc(uml_table_html(node)) if node.kind == "uml" else esc(node_display_text(node))
        parts.extend(
            [
                (
                    f'        <mxCell id="{node.id}" value="{value}" '
                    f'style="{style_for(node.kind)}" vertex="1" parent="1">'
                ),
                (
                    f'          <mxGeometry x="{node.x}" y="{node.y}" width="{node.w}" '
                    f'height="{node.h}" as="geometry" />'
                ),
                "        </mxCell>",
            ]
        )
    edge_id = 1000
    for edge in spec.edges:
        style = DASHED_EDGE_STYLE if edge.dashed else EDGE_STYLE
        parts.extend(
            [
                (
                    f'        <mxCell id="e{edge_id}" value="{esc(edge.label)}" style="{style}" '
                    f'edge="1" parent="1" source="{edge.source}" target="{edge.target}">'
                ),
                '          <mxGeometry relative="1" as="geometry" />',
                "        </mxCell>",
            ]
        )
        edge_id += 1
    parts.extend(["      </root>", "    </mxGraphModel>", "  </diagram>", "</mxfile>"])
    return "\n".join(parts)


def node_center(node: Node) -> tuple[float, float]:
    return (node.x + (node.w / 2), node.y + (node.h / 2))


def edge_points(source: Node, target: Node) -> tuple[tuple[int, int], tuple[int, int]]:
    sx, sy = node_center(source)
    tx, ty = node_center(target)
    dx = tx - sx
    dy = ty - sy
    if abs(dx) >= abs(dy):
        start = (source.x + source.w, int(sy)) if dx >= 0 else (source.x, int(sy))
        end = (target.x, int(ty)) if dx >= 0 else (target.x + target.w, int(ty))
    else:
        start = (int(sx), source.y + source.h) if dy >= 0 else (int(sx), source.y)
        end = (int(tx), target.y) if dy >= 0 else (int(tx), target.y + target.h)
    return start, end


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw_line in text.splitlines() or [""]:
        words = raw_line.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            trial = f"{current} {word}"
            width = draw.textbbox((0, 0), trial, font=font)[2]
            if width <= max_width:
                current = trial
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def draw_multiline_centered(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str = TEXT,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, max(40, x2 - x1 - 16))
    line_heights = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line or "Ag", font=font)
        line_heights.append(bbox[3] - bbox[1])
    total_h = sum(line_heights) + max(0, len(lines) - 1) * 4
    y = y1 + ((y2 - y1 - total_h) / 2)
    for line, line_h in zip(lines, line_heights):
        bbox = draw.textbbox((0, 0), line or " ", font=font)
        line_w = bbox[2] - bbox[0]
        x = x1 + ((x2 - x1 - line_w) / 2)
        draw.text((x, y), line, font=font, fill=fill)
        y += line_h + 4


def draw_multiline_left(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str = TEXT,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, max(40, x2 - x1 - 20))
    y = y1
    for line in lines:
        bbox = draw.textbbox((0, 0), line or "Ag", font=font)
        line_h = bbox[3] - bbox[1]
        if y + line_h > y2:
            break
        draw.text((x1, y), line, font=font, fill=fill)
        y += line_h + 3


def draw_uml_box(draw: ImageDraw.ImageDraw, node: Node, title_font, body_font) -> None:
    x1, y1, x2, y2 = node.x, node.y, node.x + node.w, node.y + node.h
    draw.rectangle((x1, y1, x2, y2), fill="#ffffff", outline="#000000", width=3)

    title_h = max(44, int(node.h * 0.16))
    attr_h = max(70, int(node.h * 0.42))
    method_h_start = y1 + title_h + attr_h

    draw.line((x1, y1 + title_h, x2, y1 + title_h), fill="#000000", width=2)
    draw.line((x1, method_h_start, x2, method_h_start), fill="#000000", width=2)

    draw_multiline_centered(draw, (x1 + 8, y1 + 4, x2 - 8, y1 + title_h - 4), node.text, title_font)
    draw_multiline_left(
        draw,
        (x1 + 10, y1 + title_h + 8, x2 - 10, method_h_start - 8),
        "\n".join(node.attributes or ["-"]),
        body_font,
    )
    draw_multiline_left(
        draw,
        (x1 + 10, method_h_start + 8, x2 - 10, y2 - 8),
        "\n".join(node.methods or ["-"]),
        body_font,
    )


def orthogonal_path(start: tuple[int, int], end: tuple[int, int]) -> list[tuple[int, int]]:
    if start[0] == end[0] or start[1] == end[1]:
        return [start, end]
    if abs(end[0] - start[0]) >= abs(end[1] - start[1]):
        mid_x = int((start[0] + end[0]) / 2)
        return [start, (mid_x, start[1]), (mid_x, end[1]), end]
    mid_y = int((start[1] + end[1]) / 2)
    return [start, (start[0], mid_y), (end[0], mid_y), end]


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], dashed: bool) -> None:
    path = orthogonal_path(start, end)
    if dashed:
        for p1, p2 in zip(path, path[1:]):
            segments = 12
            for i in range(segments):
                if i % 2 == 0:
                    a = (
                        int(p1[0] + (p2[0] - p1[0]) * (i / segments)),
                        int(p1[1] + (p2[1] - p1[1]) * (i / segments)),
                    )
                    b = (
                        int(p1[0] + (p2[0] - p1[0]) * ((i + 1) / segments)),
                        int(p1[1] + (p2[1] - p1[1]) * ((i + 1) / segments)),
                    )
                    draw.line([a, b], fill=EDGE, width=3)
    else:
        draw.line(path, fill=EDGE, width=3)

    tail = path[-2]
    angle = math.atan2(end[1] - tail[1], end[0] - tail[0])
    arrow_len = 14
    left = (
        int(end[0] - arrow_len * math.cos(angle - math.pi / 6)),
        int(end[1] - arrow_len * math.sin(angle - math.pi / 6)),
    )
    right = (
        int(end[0] - arrow_len * math.cos(angle + math.pi / 6)),
        int(end[1] - arrow_len * math.sin(angle + math.pi / 6)),
    )
    draw.polygon([end, left, right], fill=EDGE)


def render_diagram(spec: DiagramSpec) -> Path:
    DIAGRAM_PNG_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (spec.width, spec.height), BG)
    draw = ImageDraw.Draw(image)
    title_font = load_font(28, bold=True)
    font = load_font(24)
    uml_title_font = load_font(20, bold=True)
    uml_body_font = load_font(15)
    label_font = load_font(15)
    nodes_by_id = {node.id: node for node in spec.nodes}

    draw.text((80, 40), spec.page_name, font=title_font, fill=TEXT)

    for edge in spec.edges:
        source = nodes_by_id[edge.source]
        target = nodes_by_id[edge.target]
        start, end = edge_points(source, target)
        draw_arrow(draw, start, end, edge.dashed)
        if edge.label:
            mx = int((start[0] + end[0]) / 2)
            my = int((start[1] + end[1]) / 2)
            bbox = draw.textbbox((0, 0), edge.label, font=label_font)
            pad = 6
            label_box = (
                mx - ((bbox[2] - bbox[0]) // 2) - pad,
                my - 12,
                mx + ((bbox[2] - bbox[0]) // 2) + pad,
                my + 12,
            )
            draw.rectangle(label_box, fill="#ffffff", outline="#cccccc", width=1)
            draw.text((label_box[0] + pad, label_box[1] + 2), edge.label, font=label_font, fill=TEXT)

    for node in spec.nodes:
        box = (node.x, node.y, node.x + node.w, node.y + node.h)
        if node.kind == "uml":
            draw_uml_box(draw, node, uml_title_font, uml_body_font)
            continue
        if node.kind == "terminator":
            draw.rectangle(box, fill=TERM_FILL, outline=TERM_OUTLINE, width=3)
        elif node.kind == "note":
            draw.rectangle(box, fill=NOTE_FILL, outline=NOTE_OUTLINE, width=3)
        else:
            draw.rectangle(box, fill=BOX_FILL, outline=BOX_OUTLINE, width=3)
        draw_multiline_centered(draw, box, node.text, font)

    output_path = DIAGRAM_PNG_DIR / f"{Path(spec.filename).stem}.png"
    image.save(output_path, format="PNG")
    return output_path


def set_run_font(run, size: int = 13, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_font(paragraph, size: int = 13) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size)


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_paragraph(document: Document, text: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def remove_between(document: Document, start_text: str, end_text: str) -> None:
    removing = False
    to_remove = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == start_text:
            removing = True
            continue
        if removing and text == end_text:
            break
        if removing:
            to_remove.append(paragraph)
    for paragraph in to_remove:
        remove_paragraph(paragraph)


def desired_width(spec: DiagramSpec) -> Inches:
    ratio = spec.width / spec.height
    if ratio >= 1.7:
        return Inches(6.2)
    if ratio >= 1.3:
        return Inches(5.9)
    return Inches(5.5)


def figure_label_for_heading(heading: str) -> str:
    return heading.split(" ", 1)[0].rstrip(".")


SEQUENCE_ITEMS = [
    ReportItem(
        "3.2.1. Luồng đăng nhập và nhập tên hiển thị",
        "3_2_1_sequence_dang_nhap_nhap_ten.drawio",
        "Luồng đăng nhập và nhập tên hiển thị",
        [
            "Sequence đầu tiên mô tả chuỗi xác thực PlayFab và thiết lập Display Name cho người chơi. Từ thời điểm ứng dụng khởi động, PlayFabLeaderboardManager gửi yêu cầu LoginWithCustomID, đọc hồ sơ hiện có rồi quyết định có cần mở NameInputPanel để nhập tên hiển thị hay không.",
            "Luồng này làm rõ dữ liệu nền của hệ thống trước khi gameplay bắt đầu, bao gồm Custom ID, CurrentPlayFabId và CurrentDisplayName. Việc tách riêng sequence đăng nhập giúp báo cáo phân biệt rõ phần xác thực trực tuyến với phần vòng lặp chiến đấu diễn ra trong scene Game.",
        ],
    ),
    ReportItem(
        "3.2.2. Luồng bắt đầu trận và khởi tạo wave",
        "3_2_2_sequence_bat_dau_tran_va_khoi_tao_wave.drawio",
        "Luồng bắt đầu trận và khởi tạo wave",
        [
            "Sequence này mô tả giai đoạn chuyển từ trạng thái chờ sang một phiên chơi chủ động. Người chơi tương tác với ChallengePostNPC, ChallengePanel hiển thị hướng dẫn, phát sự kiện onGameStart và gọi WaveSpawner.StartNextWave để khởi tạo wave đầu tiên.",
            "Cách trình bày này cho thấy thao tác bắt đầu trận không chỉ là một nút giao diện, mà là điểm kích hoạt đồng thời nhiều thành phần như khóa/mở input, ẩn NPC, bắt đầu hệ thống level và chuyển gameplay sang pha runtime. Đây là cầu nối giữa nhóm UI/NPC và nhóm spawn wave trong thiết kế tổng thể.",
        ],
    ),
    ReportItem(
        "3.2.3. Luồng hệ thống spawn quái",
        "3_2_3_sequence_spawn_quai.drawio",
        "Luồng hệ thống spawn quái",
        [
            "Sequence Diagram này mô tả cách hệ thống tạo quái mỗi khi một wave mới bắt đầu. ChallengePanel chỉ đóng vai trò gửi tín hiệu StartNextWave, còn WaveSpawner là lớp điều phối trung tâm: tăng currentWave, lấy cấu hình từ WaveConfig, khởi tạo wave session mới và quyết định dùng wave có sẵn hay GenerateEndlessWave nếu đã vượt phần cấu hình định trước.",
            "Sau thời gian preparationTime, WaveSpawner cập nhật trạng thái wave, gửi số lượng quái ban đầu lên HUD và đi vào hai nhánh chính. Với wave thường, hệ thống duyệt từng EnemyGroup trong SimpleWaveData rồi gọi ObjectPool để sinh quái theo nhóm; với boss wave, hệ thống chọn bossPoolType, bossSpawnPosition và thực hiện SpawnBossNow hoặc SpawnBossAfterEffect. Về mặt cấu trúc lớp, luồng này được giải thích chi tiết hơn ở mục 3.5.7.",
        ],
    ),
    ReportItem(
        "3.2.4. Luồng chiến đấu, nhận EXP và chọn buff",
        "3_2_4_sequence_chien_dau_exp_buff.drawio",
        "Luồng chiến đấu, nhận EXP và chọn buff",
        [
            "Sequence combat/progression tập trung vào mạch chiến đấu cốt lõi sau khi enemy đã có mặt trong scene. PlayerAttack sinh projectile qua ObjectPool, Enemy nhận sát thương, ExpDropper cộng expValue cho PlayerLevelSystem và sự kiện OnLevelUp kích hoạt BuffCardManager cùng CardSelectionPanel.",
            "Luồng này cho thấy rõ cách gameplay runtime và hệ thống progression liên kết với nhau thông qua event. Khi EXP đủ ngưỡng, quá trình lên cấp không dừng ở việc đổi số liệu trong PlayerLevelSystem mà còn dẫn tới quyết định buff mới, làm thay đổi trực tiếp nhịp độ và sức mạnh của người chơi ở các wave tiếp theo.",
        ],
    ),
    ReportItem(
        "3.2.5. Luồng game over, gửi điểm và leaderboard",
        "3_2_5_sequence_game_over_gui_diem_tai_leaderboard.drawio",
        "Luồng game over, gửi điểm và leaderboard",
        [
            "Sequence cuối cùng bắt đầu tại thời điểm PlayerHealth xác nhận nhân vật đã chết. Hệ thống lấy totalExpGained từ PlayerLevelSystem, quy đổi thành điểm cuối trận, gọi SubmitScore tại PlayFabLeaderboardManager rồi tải lại leaderboard tổng và leaderboard quanh người chơi hiện tại.",
            "Nhờ có một sequence riêng, báo cáo làm rõ rằng kết thúc trận là giao điểm giữa gameplay cục bộ và backend. Điểm số ở đây bám đúng thống kê HighScore, không dùng thời gian sống hay số quái tiêu diệt, qua đó thống nhất với logic gửi dữ liệu đang được hiện thực trong lớp PlayFabLeaderboardManager.",
        ],
    ),
]


CLASS_ITEMS = [
    ReportItem(
        "3.5.1. Nhóm Player Core",
        "3_5_1_class_player_core.drawio",
        "Nhóm Player Core",
        [
            "Sơ đồ lớp Player Core tập trung vào bốn lớp trực tiếp điều khiển đối tượng người chơi trong runtime: PlayerController, PlayerAttack, PlayerHealth và PlayerData, cùng quan hệ Singleton của PlayerController. Cách tách này giúp phần thiết kế làm rõ ranh giới giữa lớp điều phối input, lớp xử lý tấn công, lớp quản lý máu và lớp lưu tham số vận hành của nhân vật.",
            "Trong hiện thực, PlayerController lấy dữ liệu từ PlayerData để xử lý di chuyển, dash và gravity; đồng thời phối hợp với PlayerHealth để chặn input khi người chơi đã chết. PlayerAttack và PlayerHealth cũng cùng phụ thuộc vào PlayerData cho các thông số chiến đấu, vì vậy PlayerData trở thành điểm quy tụ của toàn bộ trạng thái cốt lõi mà nhân vật sử dụng trong một phiên chơi.",
        ],
    ),
    ReportItem(
        "3.5.2. Nhóm Progression và Buff",
        "3_5_2_class_progression_buff.drawio",
        "Nhóm Progression và Buff",
        [
            "Sơ đồ lớp này mô tả cụm phát triển sức mạnh của người chơi, gồm PlayerLevelSystem, BuffCardManager, CardSelectionPanel và PlayerStatsPanel. PlayerLevelSystem là trung tâm tính toán level/EXP, còn BuffCardManager và CardSelectionPanel chịu trách nhiệm chuyển một sự kiện tăng cấp thành lựa chọn buff cụ thể trên giao diện.",
            "Điểm quan trọng của nhóm này là cơ chế event. PlayerLevelSystem phát OnExpChanged, OnLevelChanged và đặc biệt là OnLevelUp; từ đó PlayerStatsPanel cập nhật HUD, còn CardSelectionPanel lắng nghe sự kiện để hiển thị bộ thẻ buff mới. Nhờ tách riêng nhóm Progression và Buff, báo cáo giải thích rõ rằng việc tăng sức mạnh của player được điều khiển bởi event runtime chứ không phải các lời gọi trực tiếp rời rạc.",
        ],
    ),
    ReportItem(
        "3.5.3. Nhóm Enemy Core",
        "3_5_3_class_enemy_core.drawio",
        "Nhóm Enemy Core",
        [
            "Nhóm Enemy Core trình bày cấu trúc kế thừa từ lớp Enemy sang MeleeEnemy, RangedEnemy, FlyEnemy và BossEnemy, đồng thời giữ mối liên hệ giữa EnemyData, EnemyConfig và giao diện IDamageable. Đây là cụm lớp quyết định hành vi AI, chỉ số chiến đấu và vòng đời của mỗi địch thủ xuất hiện trong trận.",
            "Lớp Enemy đóng vai trò xương sống của toàn bộ hệ thống địch, chịu trách nhiệm cập nhật AI, di chuyển bằng CharacterController, nhận sát thương và phát các sự kiện như OnDeath hoặc OnHealthChanged. BossEnemy kế thừa từ Enemy rồi bổ sung logic phase, nhờ đó báo cáo có thể chỉ ra rằng boss không phải một hệ riêng biệt mà là một nhánh mở rộng từ cùng một nền tảng đối kháng.",
        ],
    ),
    ReportItem(
        "3.5.4. Nhóm Projectile và Damage Flow",
        "3_5_4_class_projectile_damage.drawio",
        "Nhóm Projectile và Damage Flow",
        [
            "Sơ đồ lớp Projectile và Damage Flow tập trung vào Projectile, PlayerProjectile, EnemyProjectile và SpiritProjectileScript, kèm theo mối liên hệ của chúng với PlayerAttack, Enemy và IDamageable. Tách nhóm này ra khỏi Enemy Core giúp báo cáo diễn giải riêng đường đi của sát thương thay vì gộp chung với phần AI hoặc dữ liệu enemy.",
            "Trong runtime, PlayerAttack và các lớp enemy khởi tạo projectile với bộ tham số như damage, speed, lifetime và hướng bắn. Khi projectile va chạm, giao diện IDamageable đóng vai trò lớp trừu tượng hóa đối tượng nhận sát thương. Điều này cho phép hệ thống dùng chung một mẫu phát đạn cho nhiều kiểu nguồn gây damage nhưng vẫn giữ phần tiếp nhận sát thương nhất quán.",
        ],
    ),
    ReportItem(
        "3.5.5. Nhóm UI và Scene Interaction",
        "3_5_5_class_ui_scene_interaction.drawio",
        "Nhóm UI và Scene Interaction",
        [
            "Nhóm UI và Scene Interaction mô tả PanelBase, các panel con trọng yếu, GameUI cùng nhánh NPC trong scene gồm ChallengePostNPC và ChestBuffBox. Phần tách này làm rõ rằng giao diện trong project không chỉ là tập hợp các canvas độc lập, mà được tổ chức theo một lớp gốc dùng chung hiệu ứng hiện/ẩn và một hub quản lý là GameUI.",
            "Song song với đó, NPC là lớp cha của các đối tượng tương tác trong scene, chịu trách nhiệm trigger, prompt và xử lý input F. ChallengePostNPC và ChestBuffBox là hai trường hợp sử dụng tiêu biểu, nơi gameplay runtime bắt đầu hoặc phát sinh tương tác ngoài combat. Việc gom các lớp này vào một sơ đồ riêng giúp báo cáo phân biệt rõ UI panel với interaction trong không gian 3D nhưng vẫn chỉ ra mối nối giữa chúng.",
        ],
    ),
    ReportItem(
        "3.5.6. Nhóm Backend và Leaderboard",
        "3_5_6_class_backend_leaderboard.drawio",
        "Nhóm Backend và Leaderboard",
        [
            "Sơ đồ Backend và Leaderboard lấy PlayFabLeaderboardManager làm lớp trung tâm, liên hệ với NameInputPanel và LeaderboardPanel là hai cửa sổ giao diện trực tiếp phản ánh dữ liệu xác thực và bảng xếp hạng. Nhóm này được tách riêng để nhấn mạnh đây là cụm đồng bộ trực tuyến, không thuộc mạch combat hay progression trong runtime nội bộ.",
            "PlayFabLeaderboardManager quản lý đầy đủ các bước Login, đọc profile, SubmitName, SubmitScore và tải leaderboard. Do đó, khi mô tả kết quả đồ án, báo cáo có thể chỉ ra rằng NameInputPanel và LeaderboardPanel không tự giao tiếp với dịch vụ ngoài mà thông qua một lớp trung gian duy nhất, giúp chuẩn hóa dữ liệu CurrentDisplayName, CurrentPlayFabId và HighScore.",
        ],
    ),
    ReportItem(
        "3.5.7. Nhóm Spawn Enemy và Wave Runtime",
        "3_5_7_class_spawn_enemy_wave_runtime.drawio",
        "Nhóm Spawn Enemy và Wave Runtime",
        [
            "Nhóm Spawn Enemy và Wave Runtime mô tả cụm lớp chịu trách nhiệm sinh quái trong lúc trận đấu đang diễn ra: WaveSpawner, WaveConfig, SimpleWaveData, EnemyGroup và ObjectPool, đồng thời thể hiện mối liên hệ vào trận từ ChallengePanel và đầu ra là các instance Enemy trong scene. Đây là phần cấu trúc lớp tương ứng trực tiếp với sequence 3.2.3.",
            "WaveConfig lưu danh sách waves, mỗi SimpleWaveData chứa nhiều EnemyGroup cùng các tham số preparationTime, bossPoolTypes và bossSpawnPosition. WaveSpawner đọc cấu hình này, quản lý currentWave và waveSessionId, gọi GenerateEndlessWave khi vượt phần cấu hình có sẵn, rồi dùng ObjectPool để sinh Enemy hoặc boss theo từng nhánh. Nhờ tách riêng nhóm spawn, báo cáo giải thích được cả cấu trúc dữ liệu lẫn trách nhiệm runtime đứng sau cơ chế spawn quái theo đợt.",
        ],
    ),
]


OLD_CLASS_SPECS = [
    DiagramSpec(
        "3_5_1_class_player_core.drawio",
        "3.5.1",
        "Class diagram nhóm Player Core",
        "3.5.1 Class Player Core",
        1500,
        900,
        [
            n("n1", "Singleton<T>", 90, 300),
            n("n2", "PlayerController", 380, 130),
            n("n3", "PlayerAttack", 380, 300),
            n("n4", "PlayerHealth", 380, 470),
            n("n5", "PlayerData", 760, 300),
            n("n6", "Input / Dash /\nGravity", 1100, 130, 230, 90, "note"),
            n("n7", "Attack stats /\nmoveSpeed /\ndashCooldown", 1100, 420, 240, 90, "note"),
        ],
        [
            e("n1", "n2", "singleton"),
            e("n2", "n3", "phối hợp"),
            e("n2", "n4", "kiểm tra chết"),
            e("n2", "n5", "GetComponent"),
            e("n3", "n5", "damage/range"),
            e("n4", "n5", "health params"),
            e("n2", "n6", "runtime"),
            e("n5", "n7", "tham số"),
        ],
    ),
    DiagramSpec(
        "3_5_2_class_progression_buff.drawio",
        "3.5.2",
        "Class diagram nhóm Progression và Buff",
        "3.5.2 Class Progression va Buff",
        1600,
        920,
        [
            n("n1", "Singleton<T>", 90, 280),
            n("n2", "PlayerLevelSystem", 380, 120, 240, 90),
            n("n3", "BuffCardManager", 380, 420, 240, 90),
            n("n4", "CardSelectionPanel", 820, 420, 260, 90),
            n("n5", "PlayerStatsPanel", 820, 120, 240, 90),
            n("n6", "OnLevelUp", 1120, 280, 220, 80, "terminator"),
            n("n7", "OnExpChanged /\nOnLevelChanged", 1120, 120, 240, 90, "note"),
        ],
        [
            e("n1", "n2", "singleton"),
            e("n1", "n3", "singleton"),
            e("n2", "n5", "hud update"),
            e("n2", "n4", "event listener"),
            e("n2", "n6", "phát sự kiện"),
            e("n3", "n4", "cards"),
            e("n2", "n7", "event stream"),
        ],
    ),
    DiagramSpec(
        "3_5_3_class_enemy_core.drawio",
        "3.5.3",
        "Class diagram nhóm Enemy Core",
        "3.5.3 Class Enemy Core",
        1800,
        1000,
        [
            n("n1", "IDamageable", 90, 360),
            n("n2", "Enemy", 430, 360),
            n("n3", "MeleeEnemy", 780, 120),
            n("n4", "RangedEnemy", 780, 280),
            n("n5", "FlyEnemy", 780, 440),
            n("n6", "BossEnemy", 780, 600),
            n("n7", "EnemyData", 1120, 240),
            n("n8", "EnemyConfig", 1120, 460),
            n("n9", "Boss phase /\nOnBossDied", 1460, 600, 240, 90, "note"),
        ],
        [
            e("n1", "n2", "nhận damage"),
            e("n2", "n3", "kế thừa"),
            e("n2", "n4", "kế thừa"),
            e("n2", "n5", "kế thừa"),
            e("n2", "n6", "kế thừa"),
            e("n7", "n2", "runtime state"),
            e("n8", "n7", "load config"),
            e("n6", "n9", "phase logic"),
        ],
    ),
    DiagramSpec(
        "3_5_4_class_projectile_damage.drawio",
        "3.5.4",
        "Class diagram nhóm Projectile và Damage Flow",
        "3.5.4 Class Projectile va Damage Flow",
        1800,
        920,
        [
            n("n1", "PlayerAttack", 90, 200),
            n("n2", "Enemy", 90, 460),
            n("n3", "Projectile", 470, 320),
            n("n4", "PlayerProjectile", 860, 120),
            n("n5", "EnemyProjectile", 860, 320),
            n("n6", "SpiritProjectileScript", 860, 520, 280, 90),
            n("n7", "IDamageable", 1280, 320),
            n("n8", "damage / speed /\nlifetime", 1280, 140, 240, 90, "note"),
        ],
        [
            e("n1", "n4", "spawn shot"),
            e("n2", "n5", "shoot"),
            e("n2", "n6", "special shot"),
            e("n3", "n4", "base class"),
            e("n3", "n5", "base class"),
            e("n3", "n6", "base class"),
            e("n4", "n7", "apply damage"),
            e("n5", "n7", "apply damage"),
            e("n6", "n7", "apply damage"),
            e("n3", "n8", "initialize"),
        ],
    ),
    DiagramSpec(
        "3_5_5_class_ui_scene_interaction.drawio",
        "3.5.5",
        "Class diagram nhóm UI và Scene Interaction",
        "3.5.5 Class UI va Scene Interaction",
        1900,
        980,
        [
            n("n1", "PanelBase", 100, 260),
            n("n2", "ChallengePanel", 430, 100),
            n("n3", "NameInputPanel", 430, 260),
            n("n4", "LeaderboardPanel", 430, 420),
            n("n5", "PauseMenuPanel", 430, 580),
            n("n6", "GameUI", 830, 260, 240, 90),
            n("n7", "NPC", 1180, 260),
            n("n8", "ChallengePostNPC", 1480, 120, 260, 90),
            n("n9", "ChestBuffBox", 1480, 400, 240, 90),
            n("n10", "prompt F /\nshow-hide panel", 1180, 520, 240, 90, "note"),
        ],
        [
            e("n1", "n2", "kế thừa"),
            e("n1", "n3", "kế thừa"),
            e("n1", "n4", "kế thừa"),
            e("n1", "n5", "kế thừa"),
            e("n6", "n2", "hub panel"),
            e("n6", "n3", "hub panel"),
            e("n6", "n4", "hub panel"),
            e("n6", "n5", "hub panel"),
            e("n7", "n8", "kế thừa"),
            e("n7", "n9", "kế thừa"),
            e("n7", "n10", "scene interaction"),
        ],
    ),
    DiagramSpec(
        "3_5_6_class_backend_leaderboard.drawio",
        "3.5.6",
        "Class diagram nhóm Backend và Leaderboard",
        "3.5.6 Class Backend va Leaderboard",
        1650,
        900,
        [
            n("n1", "Singleton<T>", 110, 300),
            n("n2", "PlayFabLeaderboardManager", 420, 260, 300, 100),
            n("n3", "NameInputPanel", 860, 120),
            n("n4", "LeaderboardPanel", 860, 400),
            n("n5", "Login /\nGetPlayerProfile", 1240, 120, 240, 90, "terminator"),
            n("n6", "SubmitScore /\nGetLeaderboardData", 1240, 400, 260, 90, "terminator"),
            n("n7", "CurrentPlayFabId /\nCurrentDisplayName /\nHighScore", 1240, 610, 260, 110, "note"),
        ],
        [
            e("n1", "n2", "singleton"),
            e("n2", "n3", "nhập tên"),
            e("n2", "n4", "hiển thị bảng"),
            e("n2", "n5", "xác thực"),
            e("n2", "n6", "leaderboard"),
            e("n2", "n7", "quản lý dữ liệu"),
        ],
    ),
    DiagramSpec(
        "3_5_7_class_spawn_enemy_wave_runtime.drawio",
        "3.5.7",
        "Class diagram nhóm Spawn Enemy và Wave Runtime",
        "3.5.7 Class Spawn Enemy va Wave Runtime",
        1900,
        980,
        [
            n("n1", "ChallengePanel", 90, 140),
            n("n2", "WaveSpawner", 420, 140, 240, 90),
            n("n3", "WaveConfig", 780, 140, 220, 90),
            n("n4", "SimpleWaveData", 1120, 140, 240, 90),
            n("n5", "EnemyGroup", 1120, 380, 220, 90),
            n("n6", "ObjectPool", 1480, 300, 220, 90),
            n("n7", "Enemy", 1480, 520),
            n("n8", "GenerateEndlessWave", 780, 380, 240, 90, "note"),
            n("n9", "bossPoolTypes /\nspawnPosition /\npreparationTime", 1120, 620, 280, 110, "note"),
        ],
        [
            e("n1", "n2", "StartNextWave"),
            e("n2", "n3", "đọc config"),
            e("n3", "n4", "wave"),
            e("n4", "n5", "enemyGroups"),
            e("n2", "n8", "vượt cấu hình"),
            e("n2", "n6", "spawn"),
            e("n6", "n7", "enemy instance"),
            e("n4", "n9", "wave data"),
        ],
    ),
]


NEW_CLASS_SPECS = [
    DiagramSpec(
        "3_5_1_class_player_core.drawio",
        "3.5.1",
        "Class diagram UML nhóm Player Core",
        "3.5.1 Class Diagram - Player Core",
        2100,
        1280,
        [
            uml("n1", "Singleton<T>", 120, 480, ["+ Instance: T"], ["# Awake()"], 260, 170),
            uml("n2", "PlayerController", 480, 120, ["- controller: CharacterController", "- playerHealth: PlayerHealth", "- playerData: PlayerData", "- moveInput: Vector2"], ["+ Update()", "+ SetInputActive(active: bool)", "+ IsInputActive(): bool"], 360, 250),
            uml("n3", "PlayerAttack", 480, 450, ["- attackTimer: float", "- playerHealth: PlayerHealth", "- playerData: PlayerData", "- currentTarget: Transform"], ["+ Update()", "- TryAttack()", "- SpawnProjectiles()", "+ SetPlayerData(data: PlayerData)"], 360, 260),
            uml("n4", "PlayerHealth", 480, 790, ["- playerData: PlayerData", "+ OnHealthChanged", "+ OnDeath"], ["+ TakeDamage(...)", "+ Heal(amount: float)", "+ IsDead(): bool", "+ GetCurrentHealth(): float"], 360, 240),
            uml("n5", "PlayerData", 1120, 450, ["+ moveSpeed: float", "+ dashCooldown: float", "+ attackDamage: float", "+ attackRange: float"], ["+ LoadFromConfig()", "+ ResetData()", "+ GetTotalDamage(): float", "+ GetAttackCooldown(): float"], 400, 260),
        ],
        [e("n1", "n2", "kế thừa"), e("n2", "n3", ""), e("n2", "n4", ""), e("n2", "n5", "dùng dữ liệu"), e("n3", "n5", ""), e("n4", "n5", "")],
    ),
    DiagramSpec(
        "3_5_2_class_progression_buff.drawio",
        "3.5.2",
        "Class diagram UML nhóm Progression và Buff",
        "3.5.2 Class Diagram - Progression va Buff",
        2200,
        1280,
        [
            uml("n1", "PlayerLevelSystem", 120, 140, ["- currentLevel: int", "- currentExp: float", "+ totalExpGained: float", "+ OnLevelUp: UnityEvent<int>"], ["+ AddExp(amount: float)", "+ GrantLevels(count: int)", "+ GetCurrentLevel(): int", "+ GetTotalExpGained(): float"], 390, 280),
            uml("n2", "BuffCardManager", 740, 140, ["+ Instance", "- availableCards: List<BuffCardConfig>"], ["+ SelectCards()", "+ ApplyBuff(...)", "+ GetCardLevel(...): int", "+ GetMaxLevelForBuff(...): int"], 390, 260),
            uml("n3", "CardSelectionPanel", 1360, 140, ["- queuedLevelRewards: Queue<int>", "- cardManager: BuffCardManager", "- waitingForThemeTransition: bool"], ["- OnPlayerLevelUp(newLevel: int)", "+ ShowCards(cards: List<BuffCardConfig>)", "+ Hide()", "- TryShowNextQueuedCards()"], 430, 290),
            uml("n4", "PlayerStatsPanel", 740, 610, ["- boundLevelSystem: PlayerLevelSystem", "- boundWaveSpawner: WaveSpawner", "- boundPlayerHealth: PlayerHealth"], ["+ ResetForReplay()", "- BindRuntimeReferences()", "- UpdateExp(...)", "- UpdateWave(waveNumber: int)"], 430, 270),
        ],
        [e("n1", "n3", "OnLevelUp"), e("n1", "n4", "OnExpChanged / OnLevelChanged"), e("n2", "n3", "cards"), e("n3", "n2", "")],
    ),
    DiagramSpec(
        "3_5_3_class_enemy_core.drawio",
        "3.5.3",
        "Class diagram UML nhóm Enemy Core",
        "3.5.3 Class Diagram - Enemy Core",
        2500,
        1380,
        [
            uml("n1", "IDamageable", 120, 520, ["+ IsDead(): bool"], ["+ TakeDamage(...)", "+ Heal(amount: float)"], 280, 180),
            uml("n2", "Enemy", 540, 420, ["- enemyData: EnemyData", "- currentState: EnemyState", "+ OnDeath: UnityEvent", "+ OnHealthChanged: UnityEvent<float,float>"], ["+ Update()", "# UpdateAI()", "+ TakeDamage(...)", "# Die()"], 430, 300),
            uml("n3", "MeleeEnemy", 1160, 80, ["-"], ["# UpdateMeleeAI(...)"], 300, 170),
            uml("n4", "RangedEnemy", 1160, 330, ["-"], ["# UpdateRangedAI(...)", "# ShootProjectile()"], 320, 180),
            uml("n5", "FlyEnemy", 1160, 580, ["-"], ["# UpdateAI()"], 300, 170),
            uml("n6", "BossEnemy", 1160, 830, ["- currentPhase: int", "+ OnPhaseChanged: UnityEvent<int>", "+ OnBossDied: UnityEvent"], ["# CheckPhaseTransition()", "# TransitionToPhase(...)", "# OnPhase2()", "# OnPhase3()"], 360, 260),
            uml("n7", "EnemyData", 1720, 280, ["+ maxHealth: float", "+ moveSpeed: float", "+ attackRange: float", "+ enemyType: EnemyType"], ["+ ResetHealth()", "+ GetConfig<T>()", "+ GetCurrentHealth(): float"], 360, 240),
            uml("n8", "EnemyConfig", 1720, 690, ["+ maxHealth: float", "+ moveSpeed: float", "+ expValue: int"], ["+ ApplyTo(data: EnemyData)"], 320, 190),
        ],
        [e("n1", "n2", "thực thi"), e("n2", "n3", "kế thừa"), e("n2", "n4", "kế thừa"), e("n2", "n5", "kế thừa"), e("n2", "n6", "kế thừa"), e("n7", "n2", "state / stats"), e("n8", "n7", "cấu hình")],
    ),
    DiagramSpec(
        "3_5_4_class_projectile_damage.drawio",
        "3.5.4",
        "Class diagram UML nhóm Projectile và Damage Flow",
        "3.5.4 Class Diagram - Projectile va Damage Flow",
        2250,
        1280,
        [
            uml("n1", "PlayerAttack", 120, 140, ["- attackTimer: float", "- currentTarget: Transform"], ["- PerformAttack()", "- SpawnSingleProjectile(...)", "- FindNearestEnemy(): Transform"], 360, 240),
            uml("n2", "Enemy", 120, 720, ["- firePoint: Transform", "- projectilePrefab: GameObject"], ["# ShootProjectile()", "+ TakeDamage(...)"], 340, 210),
            uml("n3", "Projectile", 760, 430, ["# damage: float", "# speed: float", "# lifetime: float", "# direction: Vector3"], ["+ Initialize(...)", "# Update()", "# OnHit(other: Collider)", "# DispawnProjectile()"], 390, 280),
            uml("n4", "PlayerProjectile", 1360, 120, ["- isAoEEnabled: bool", "- aoeRadius: float"], ["+ InitializeExtra(...)", "# OnHit(other: Collider)", "# DispawnProjectile()"], 360, 230),
            uml("n5", "EnemyProjectile", 1360, 430, ["- owner: GameObject"], ["# OnHit(other: Collider)", "# DispawnProjectile()"], 340, 200),
            uml("n6", "SpiritProjectileScript", 1360, 730, ["- ownerEnemy: Enemy"], ["+ InitializeSpirit(...)", "# OnHit(other: Collider)"], 360, 200),
            uml("n7", "IDamageable", 1820, 430, ["+ IsDead(): bool"], ["+ TakeDamage(...)"], 280, 170),
        ],
        [e("n3", "n4", "kế thừa"), e("n3", "n5", "kế thừa"), e("n3", "n6", "kế thừa"), e("n1", "n4", "spawn"), e("n2", "n5", "shoot"), e("n2", "n6", "special"), e("n4", "n7", ""), e("n5", "n7", ""), e("n6", "n7", "")],
    ),
    DiagramSpec(
        "3_5_5_class_ui_scene_interaction.drawio",
        "3.5.5",
        "Class diagram UML nhóm UI và Scene Interaction",
        "3.5.5 Class Diagram - UI va Scene Interaction",
        2500,
        1380,
        [
            uml("n1", "PanelBase", 120, 480, ["- menu: GameObject", "+ IsOpen: bool"], ["+ Show()", "+ Hide()", "+ HideImmediate()"], 320, 210),
            uml("n2", "ChallengePanel", 560, 120, ["- bg: GameObject", "- startGameButton: Button", "+ onGameStart: Action"], ["+ ShowTutorial()", "+ Dismiss()", "+ StartGame()"], 380, 240),
            uml("n3", "NameInputPanel", 560, 410, ["- inputField: TMP_InputField"], ["+ Show()", "+ SubmitName()"], 320, 180),
            uml("n4", "LeaderboardPanel", 560, 660, ["- entries: List<...>"], ["+ Show()", "+ RefreshLeaderboard()"], 320, 180),
            uml("n5", "PauseMenuPanel", 560, 910, ["- isPaused: bool"], ["+ PauseGame()", "+ ResumeGame()"], 320, 180),
            uml("n6", "GameUI", 1080, 480, ["- challengePanel: ChallengePanel", "- leaderboardPanel: LeaderboardPanel", "- pauseMenuPanel: PauseMenuPanel"], ["+ PrepareForSceneReload()", "- ResolveMissingReferences()"], 430, 220),
            uml("n7", "NPC", 1640, 480, ["- playerInRange: bool", "- interactText: string"], ["# Interact()", "# IsPanelOpen(): bool", "# OnPanelClosed()"], 360, 220),
            uml("n8", "ChallengePostNPC", 2060, 220, ["-"], ["# Interact()", "# IsPanelOpen(): bool"], 340, 180),
            uml("n9", "ChestBuffBox", 2060, 700, ["-"], ["# Interact()"], 320, 170),
        ],
        [e("n1", "n2", "kế thừa"), e("n1", "n3", "kế thừa"), e("n1", "n4", "kế thừa"), e("n1", "n5", "kế thừa"), e("n6", "n2", ""), e("n6", "n3", ""), e("n6", "n4", ""), e("n6", "n5", ""), e("n7", "n8", "kế thừa"), e("n7", "n9", "kế thừa")],
    ),
    DiagramSpec(
        "3_5_6_class_backend_leaderboard.drawio",
        "3.5.6",
        "Class diagram UML nhóm Backend và Leaderboard",
        "3.5.6 Class Diagram - Backend va Leaderboard",
        2100,
        1120,
        [
            uml("n1", "Singleton<T>", 120, 430, ["+ Instance: T"], ["# Awake()"], 260, 170),
            uml("n2", "PlayFabLeaderboardManager", 540, 290, ["+ PlayFabTitleId: string", "+ LeaderboardStatisticName: string", "+ CurrentDisplayName: string", "+ CurrentPlayFabId: string"], ["+ Login()", "+ SubmitName(...)", "+ SubmitScore(...)", "+ GetLeaderboardData()"], 450, 300),
            uml("n3", "NameInputPanel", 1200, 120, ["- inputField: TMP_InputField"], ["+ Show()", "+ SubmitName()"], 320, 180),
            uml("n4", "LeaderboardPanel", 1200, 500, ["- entries: List<...>"], ["+ Show()", "+ LoadData()"], 320, 180),
            uml("n5", "PlayFab Client API", 1640, 290, ["-"], ["+ LoginWithCustomID()", "+ UpdatePlayerStatistics()", "+ GetLeaderboard()"], 360, 220),
        ],
        [e("n1", "n2", "kế thừa"), e("n2", "n3", "nhập tên"), e("n2", "n4", "hiển thị bảng"), e("n2", "n5", "gọi API")],
    ),
    DiagramSpec(
        "3_5_7_class_spawn_enemy_wave_runtime.drawio",
        "3.5.7",
        "Class diagram UML nhóm Spawn Enemy và Wave Runtime",
        "3.5.7 Class Diagram - Spawn Enemy va Wave Runtime",
        2800,
        1320,
        [
            uml("n1", "ChallengePanel", 120, 250, ["+ onGameStart: Action"], ["+ StartGame()"], 300, 160),
            uml("n2", "WaveSpawner", 520, 170, ["- waveConfig: WaveConfig", "- currentWave: int", "- waveSessionId: int", "+ OnWaveStart: UnityEvent<int>"], ["+ StartNextWave()", "+ JumpToWave(...)", "- GenerateEndlessWave(...)", "- RunWave(...)"], 430, 300),
            uml("n3", "WaveConfig", 1080, 170, ["+ waves: List<SimpleWaveData>", "+ autoScale: bool", "+ scalePerWave: float"], ["+ GetWave(waveNumber: int)", "+ Generate30Waves()", "+ CreateDefaultBossPoolTypes()"], 390, 250),
            uml("n4", "SimpleWaveData", 1600, 170, ["+ enemyGroups: List<EnemyGroup>", "+ preparationTime: float", "+ isBossWave: bool", "+ bossSpawnPosition: Vector3"], ["-"], 400, 250),
            uml("n5", "EnemyGroup", 1600, 570, ["+ enemyPoolType: PoolType", "+ enemyCount: int", "+ spawnPosition: Vector3", "+ spawnDelay: float"], ["-"], 390, 230),
            uml("n6", "ObjectPool", 2140, 400, ["+ Instance"], ["+ Spawn(poolType, position, rotation)", "+ Despawn(obj)"], 340, 190),
            uml("n7", "Enemy", 2540, 400, ["+ OnDeath"], ["+ TakeDamage(...)", "+ IsDead(): bool"], 220, 180),
        ],
        [e("n1", "n2", "StartNextWave"), e("n2", "n3", "đọc cấu hình"), e("n3", "n4", "waves"), e("n4", "n5", "enemyGroups"), e("n2", "n6", "spawn"), e("n6", "n7", "enemy instance")],
    ),
]


def write_diagrams() -> None:
    DIAGRAM_DRAWIO_DIR.mkdir(parents=True, exist_ok=True)
    for spec in NEW_CLASS_SPECS:
        (DIAGRAM_DRAWIO_DIR / spec.filename).write_text(build_diagram_xml(spec), encoding="utf-8")
        render_diagram(spec)


def write_index() -> None:
    lines = [
        "# Draw.io Diagram Index",
        "",
        "File | Mục báo cáo | Caption đề xuất",
        "--- | --- | ---",
        "1_4_1_1_dinh_danh_nguoi_choi.drawio | 1.4.1.1 | Sơ đồ định danh người chơi qua PlayFab",
        "1_4_1_2_bat_dau_tran_dau.drawio | 1.4.1.2 | Sơ đồ bắt đầu trận đấu",
        "1_4_1_3_player_controller.drawio | 1.4.1.3 | Sơ đồ điều khiển nhân vật",
        "1_4_1_4_chien_dau_player.drawio | 1.4.1.4 | Sơ đồ chiến đấu của người chơi",
        "1_4_1_5_enemy_system.drawio | 1.4.1.5 | Sơ đồ nhóm enemy",
        "1_4_1_6_wave_va_do_kho.drawio | 1.4.1.6 | Sơ đồ quản lý wave và độ khó",
        "1_4_1_7_theme_ban_do.drawio | 1.4.1.7 | Sơ đồ đổi theme bản đồ",
        "1_4_1_8_exp_va_len_cap.drawio | 1.4.1.8 | Sơ đồ EXP và lên cấp",
        "1_4_1_9_buff_va_tang_suc_manh.drawio | 1.4.1.9 | Sơ đồ chọn buff và tăng sức mạnh",
        "1_4_1_10_ui_trong_tran.drawio | 1.4.1.10 | Sơ đồ UI trong trận",
        "1_4_1_11_pause_va_ket_thuc_tran.drawio | 1.4.1.11 | Sơ đồ pause và kết thúc trận",
        "1_4_1_12_leaderboard.drawio | 1.4.1.12 | Sơ đồ leaderboard",
        "3_1_1_use_case_nguoi_choi_truoc_tran.drawio | 3.1 Use Case | Use Case người chơi trước trận",
        "3_1_2_use_case_gameplay_core_loop.drawio | 3.1 Use Case | Use Case gameplay core loop",
        "3_1_3_use_case_ket_thuc_tran_leaderboard.drawio | 3.1 Use Case | Use Case kết thúc trận và leaderboard",
        "3_2_1_sequence_dang_nhap_nhap_ten.drawio | 3.2 Sequence | Sequence đăng nhập và nhập tên hiển thị",
        "3_2_2_sequence_bat_dau_tran_va_khoi_tao_wave.drawio | 3.2 Sequence | Sequence bắt đầu trận và khởi tạo wave",
        "3_2_3_sequence_spawn_quai.drawio | 3.2 Sequence | Sequence hệ thống spawn quái",
        "3_2_4_sequence_chien_dau_exp_buff.drawio | 3.2 Sequence | Sequence chiến đấu - nhận EXP - chọn buff",
        "3_2_5_sequence_game_over_gui_diem_tai_leaderboard.drawio | 3.2 Sequence | Sequence game over - gửi điểm - tải leaderboard",
        "3_3_activity_gameplay_core_loop.drawio | 3.3 Activity | Activity gameplay core loop",
        "3_4_1_component_gameplay_runtime.drawio | 3.4 Component | Component gameplay runtime",
        "3_4_2_component_progression_ui.drawio | 3.4 Component | Component progression và UI",
        "3_4_3_component_backend_services.drawio | 3.4 Component | Component backend và services",
        "3_5_1_class_player_core.drawio | 3.5.1 Class | UML class diagram nhóm Player Core",
        "3_5_2_class_progression_buff.drawio | 3.5.2 Class | UML class diagram nhóm Progression và Buff",
        "3_5_3_class_enemy_core.drawio | 3.5.3 Class | UML class diagram nhóm Enemy Core",
        "3_5_4_class_projectile_damage.drawio | 3.5.4 Class | UML class diagram nhóm Projectile và Damage Flow",
        "3_5_5_class_ui_scene_interaction.drawio | 3.5.5 Class | UML class diagram nhóm UI và Scene Interaction",
        "3_5_6_class_backend_leaderboard.drawio | 3.5.6 Class | UML class diagram nhóm Backend và Leaderboard",
        "3_5_7_class_spawn_enemy_wave_runtime.drawio | 3.5.7 Class | UML class diagram nhóm Spawn Enemy và Wave Runtime",
        "",
        "Ghi chú: mỗi file chỉ chứa một sơ đồ, mở trực tiếp bằng draw.io / diagrams.net.",
    ]
    INDEX_PATH.write_text("\n".join(lines), encoding="utf-8")


def add_source_line(anchor, filename: str, heading: str):
    para = insert_paragraph_after(anchor, style="Normal")
    run1 = para.add_run("Tệp sơ đồ nguồn: ")
    set_run_font(run1, bold=True)
    run2 = para.add_run(f"{filename} - {heading}.")
    set_run_font(run2)
    return para


def add_image(anchor, spec: DiagramSpec, heading: str):
    image_para = insert_paragraph_after(anchor, style="Normal")
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = image_para.add_run()
    run.add_picture(str(DIAGRAM_PNG_DIR / f"{Path(spec.filename).stem}.png"), width=desired_width(spec))

    caption_para = insert_paragraph_after(image_para, style="Normal")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_para.add_run(f"Hình {figure_label_for_heading(heading)}: {heading.split('. ', 1)[1]}.")
    set_run_font(caption_run, italic=True)
    return caption_para


def add_body_paragraph(anchor, text: str):
    para = insert_paragraph_after(anchor, style="Normal")
    run = para.add_run(text)
    set_run_font(run)
    return para


def replace_sequence_section(document: Document) -> None:
    remove_between(document, "3.2. Sequence Diagram", "3.3. Activity Diagram")
    anchor = find_paragraph(document, "3.2. Sequence Diagram")

    dummy_specs = {spec.filename: spec for spec in NEW_CLASS_SPECS}
    dummy_specs["3_2_3_sequence_spawn_quai.drawio"] = DiagramSpec(
        "3_2_3_sequence_spawn_quai.drawio",
        "3.2.3",
        "Luồng hệ thống spawn quái",
        "3.2.3 Sequence spawn quái",
        1400,
        1100,
        [],
        [],
    )
    dummy_specs["3_2_2_sequence_bat_dau_tran_va_khoi_tao_wave.drawio"] = DiagramSpec(
        "3_2_2_sequence_bat_dau_tran_va_khoi_tao_wave.drawio",
        "3.2.2",
        "Luồng bắt đầu trận và khởi tạo wave",
        "3.2.2 Sequence bắt đầu trận",
        2200,
        1100,
        [],
        [],
    )
    dummy_specs["3_2_1_sequence_dang_nhap_nhap_ten.drawio"] = DiagramSpec(
        "3_2_1_sequence_dang_nhap_nhap_ten.drawio",
        "3.2.1",
        "Luồng đăng nhập và nhập tên hiển thị",
        "3.2.1 Sequence đăng nhập",
        1800,
        1000,
        [],
        [],
    )
    dummy_specs["3_2_4_sequence_chien_dau_exp_buff.drawio"] = DiagramSpec(
        "3_2_4_sequence_chien_dau_exp_buff.drawio",
        "3.2.4",
        "Luồng chiến đấu, nhận EXP và chọn buff",
        "3.2.4 Sequence chiến đấu",
        2200,
        1200,
        [],
        [],
    )
    dummy_specs["3_2_5_sequence_game_over_gui_diem_tai_leaderboard.drawio"] = DiagramSpec(
        "3_2_5_sequence_game_over_gui_diem_tai_leaderboard.drawio",
        "3.2.5",
        "Luồng game over, gửi điểm và leaderboard",
        "3.2.5 Sequence game over",
        2100,
        1100,
        [],
        [],
    )

    for item in SEQUENCE_ITEMS:
        anchor = insert_paragraph_after(anchor, item.heading, style="Heading 3")
        anchor = add_source_line(anchor, item.filename, item.heading)
        anchor = add_image(anchor, dummy_specs[item.filename], item.heading)
        for text in item.paragraphs:
            anchor = add_body_paragraph(anchor, text)


def replace_class_section(document: Document) -> None:
    remove_between(document, "3.5. Class Diagram", "CHƯƠNG 4 \nXÂY DỰNG CHƯƠNG TRÌNH")
    anchor = find_paragraph(document, "3.5. Class Diagram")
    specs_by_name = {spec.filename: spec for spec in NEW_CLASS_SPECS}
    intro = (
        "Phần Class Diagram được tái cấu trúc theo các cụm lớp thực tế trong mã nguồn thay vì gộp ba sơ đồ lớn như trước. "
        "Cách chia này giúp người đọc theo dõi rõ hơn ranh giới giữa player runtime, progression, enemy core, projectile, "
        "UI/interaction, backend và hệ thống spawn quái."
    )
    anchor = add_body_paragraph(anchor, intro)
    for item in CLASS_ITEMS:
        anchor = insert_paragraph_after(anchor, item.heading, style="Heading 3")
        anchor = add_source_line(anchor, item.filename, item.heading)
        anchor = add_image(anchor, specs_by_name[item.filename], item.heading)
        for text in item.paragraphs:
            anchor = add_body_paragraph(anchor, text)


def rebuild_report() -> None:
    if not SNAPSHOT_DOC.exists():
        raise FileNotFoundError(f"Snapshot doc not found: {SNAPSHOT_DOC}")
    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(SNAPSHOT_DOC, OUTPUT_DOC)
    document = Document(OUTPUT_DOC)
    replace_sequence_section(document)
    replace_class_section(document)
    document.save(OUTPUT_DOC)


def try_sync_to_onedrive() -> str:
    try:
        shutil.copyfile(OUTPUT_DOC, ONEDRIVE_DOC)
        return "synced"
    except Exception as exc:  # noqa: BLE001
        return f"failed: {exc}"


def main() -> int:
    write_diagrams()
    write_index()
    rebuild_report()
    sync_result = try_sync_to_onedrive()
    import sys

    sys.stdout.reconfigure(encoding="utf-8")
    print(f"Output doc: {OUTPUT_DOC}")
    print(f"Sync result: {sync_result}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
