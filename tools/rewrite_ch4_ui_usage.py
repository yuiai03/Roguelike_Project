from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = ROOT / "tmp" / "docs" / "BaoCao_DATN_2_onedrive_live_snapshot.docx"
OUTPUT_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_ch4_ui_usage.docx"
BACKUP_DOC = ROOT / "tmp" / "docs" / "BaoCao_DATN_2_ch4_before_ui_rewrite.docx"

BLACK = RGBColor(0, 0, 0)
CHAPTER4_HEADING = "CHƯƠNG 4 \nXÂY DỰNG CHƯƠNG TRÌNH"
CHAPTER4_FIRST_SECTION = "4.1. Toàn cảnh gameplay trong scene Game"
CHAPTER4_END_HEADING = "KẾT LUẬN"


@dataclass(frozen=True)
class SectionSpec:
    heading: str
    next_heading: str
    description: str
    usage: str
    functions: str
    result: str


SECTION_SPECS: list[SectionSpec] = [
    SectionSpec(
        heading="4.1. Toàn cảnh gameplay trong scene Game",
        next_heading="4.2. Nhân vật chính của trò chơi",
        description=(
            "Hai hình trong mục này mô tả scene Game ở hai trạng thái chính: trước khi vào combat "
            "và khi wave đã được kích hoạt. Người dùng nhìn thấy đồng thời khu vực bản đồ, nhân vật, "
            "enemy, HUD và các điểm tương tác cần thiết cho một lượt chơi hoàn chỉnh."
        ),
        usage=(
            "Khi vào scene Game, người chơi dùng WASD hoặc các phím mũi tên để di chuyển nhân vật trong "
            "bản đồ. Sau khi tiếp cận NPC challenge và nhấn F để tương tác, trận đấu được bắt đầu và màn "
            "hình chuyển sang trạng thái chiến đấu liên tục theo từng wave."
        ),
        functions=(
            "Scene này kết hợp toàn bộ chức năng chính của chương trình gồm di chuyển, dash bằng chuột phải, "
            "tấn công tự động, sinh quái theo wave, nhận EXP, lên cấp, chọn buff và theo dõi kết quả qua HUD."
        ),
        result=(
            "Người chơi quan sát được đầy đủ vòng lặp Roguelike ngay trên một màn hình thống nhất, từ chuẩn bị "
            "trận, giao chiến, tăng sức mạnh cho tới khi kết thúc lượt chơi."
        ),
    ),
    SectionSpec(
        heading="4.2. Nhân vật chính của trò chơi",
        next_heading="4.3. Danh sách enemy cơ bản trong trò chơi",
        description=(
            "Hình minh họa trong mục này tập trung vào nhân vật người chơi, là đối tượng trung tâm của toàn bộ "
            "trải nghiệm. Trên giao diện, người dùng dễ nhận biết player qua model chính, hướng nhìn và các hiệu "
            "ứng chiến đấu xuất hiện quanh nhân vật."
        ),
        usage=(
            "Người chơi điều khiển nhân vật bằng WASD hoặc phím mũi tên để chạy trong không gian 3D, dùng chuột "
            "phải để dash né đòn, còn việc tấn công được game tự động xử lý khi phát hiện enemy trong tầm đánh. "
            "Khi một panel lớn được mở, input di chuyển sẽ tạm khóa để tránh thao tác chồng chéo."
        ),
        functions=(
            "Nhân vật chính đảm nhiệm di chuyển, đổi hướng theo camera, nhận sát thương, hồi phản hồi animation "
            "và là điểm gắn của các hệ combat, level, buff và HUD. Đây là nơi người dùng trực tiếp tương tác nhiều "
            "nhất trong suốt quá trình sử dụng chương trình."
        ),
        result=(
            "Người chơi có thể tập trung vào định vị, né tránh và giữ khoảng cách an toàn, trong khi phần tấn công "
            "tự động giúp nhịp chơi mượt và dễ tiếp cận hơn."
        ),
    ),
    SectionSpec(
        heading="4.3. Danh sách enemy cơ bản trong trò chơi",
        next_heading="4.4. Nhóm enemy nguyên tố LawaChurl",
        description=(
            "Ba hình trong mục này lần lượt minh họa các enemy cơ bản của game: enemy bay, enemy tấn công tầm xa "
            "và enemy cận chiến. Người dùng có thể phân biệt chúng trực tiếp qua hình dáng, vị trí xuất hiện và "
            "kiểu tiếp cận trên màn hình."
        ),
        usage=(
            "Trong quá trình chơi, người dùng quan sát loại enemy đang xuất hiện để chọn cách di chuyển phù hợp. "
            "Enemy cận chiến buộc người chơi phải giữ khoảng cách, enemy tầm xa yêu cầu đổi vị trí liên tục, còn "
            "enemy bay khiến không gian chiến đấu trở nên khó kiểm soát hơn."
        ),
        functions=(
            "Nhóm enemy cơ bản tạo áp lực chính trong đa số wave thường. Chúng đóng vai trò làm mục tiêu cho hệ "
            "tấn công tự động của player, đồng thời tạo ra các tình huống né tránh và dồn hướng di chuyển trong runtime."
        ),
        result=(
            "Người chơi học được cách ưu tiên mối đe dọa và đọc tình huống ngay từ các wave đầu, từ đó thích nghi "
            "tốt hơn khi độ khó trận đấu tăng dần."
        ),
    ),
    SectionSpec(
        heading="4.4. Nhóm enemy nguyên tố LawaChurl",
        next_heading="4.5. Các theme map và chuyển đổi môi trường",
        description=(
            "Mục này trình bày nhóm LawaChurl, là các enemy cỡ lớn dùng để đánh dấu những pha giao tranh áp lực cao. "
            "Trên giao diện, người dùng nhận ra chúng qua kích thước lớn, tạo hình nổi bật và hiệu ứng nguyên tố rõ ràng hơn nhóm quái thường."
        ),
        usage=(
            "Khi gặp LawaChurl, người chơi cần giữ khoảng cách an toàn, quan sát hiệu ứng cảnh báo và tận dụng dash đúng thời điểm. "
            "Đây là nhóm enemy khiến người dùng phải chủ động hơn trong việc di chuyển thay vì chỉ đứng giữ vị trí."
        ),
        functions=(
            "Nhóm enemy này đóng vai trò như mốc tăng độ khó của trận đấu, tạo ra các đòn đánh diện rộng và thay đổi nhịp combat. "
            "Sự xuất hiện của chúng giúp chương trình có thêm các pha boss wave hoặc combat cao trào dễ nhận biết."
        ),
        result=(
            "Người chơi cảm nhận rõ sự leo thang độ khó của game, đồng thời hiểu rằng việc chọn buff và đọc chuyển động của enemy "
            "ảnh hưởng trực tiếp đến khả năng sống sót."
        ),
    ),
    SectionSpec(
        heading="4.5. Các theme map và chuyển đổi môi trường",
        next_heading="4.6. Giao diện HUD và hỗ trợ gameplay",
        description=(
            "Ba hình trong mục này thể hiện các theme môi trường khác nhau được dùng trong scene Game. Người dùng nhìn thấy sự thay đổi "
            "ở màu sắc nền, vật liệu mặt đất, tường và hiệu ứng phụ trợ, nhờ đó phân biệt được các giai đoạn của trận đấu."
        ),
        usage=(
            "Người chơi không cần thao tác thủ công để đổi map. Hệ thống sẽ tự chuyển theme theo tiến trình wave và hiển thị transition "
            "đen khi cần đổi môi trường, giúp người dùng nhận biết rõ thời điểm scene đang chuyển trạng thái."
        ),
        functions=(
            "Các theme map làm mới trải nghiệm thị giác và đóng vai trò đánh dấu mốc tiến triển của run. Việc đổi theme còn phối hợp với "
            "loading/transition để khóa tạm thao tác trong lúc hệ thống áp dụng môi trường mới."
        ),
        result=(
            "Người chơi cảm nhận được sự thay đổi không khí qua từng chặng wave, giảm cảm giác lặp lại và dễ nhận ra trận đấu đang bước sang giai đoạn mới."
        ),
    ),
    SectionSpec(
        heading="4.6. Giao diện HUD và hỗ trợ gameplay",
        next_heading="4.7. Giao diện challenge và chọn buff",
        description=(
            "HUD là lớp giao diện xuất hiện liên tục khi người chơi đang ở trong scene Game. Các thành phần chính gồm thanh máu, thanh kinh nghiệm, "
            "cấp độ hiện tại, chỉ báo wave và các prompt hỗ trợ như nhắc người dùng nhấn ESC để mở menu."
        ),
        usage=(
            "Trong lúc chiến đấu, người dùng theo dõi máu để quyết định có nên giữ khoảng cách hay không, nhìn thanh EXP để biết thời điểm sắp lên cấp "
            "và quan sát chỉ báo wave để nắm tiến độ trận. Prompt trên HUD cũng giúp người chơi nhớ nhanh các thao tác mở menu hoặc tương tác."
        ),
        functions=(
            "HUD cung cấp dữ liệu thời gian thực để hỗ trợ ra quyết định, thay vì chỉ đóng vai trò hiển thị trang trí. Tất cả thông tin quan trọng liên quan "
            "đến trạng thái player và trận đấu đều được gom về một cụm giao diện dễ đọc."
        ),
        result=(
            "Người dùng có thể phản ứng kịp thời trước thay đổi của gameplay mà không cần rời khỏi màn hình chiến đấu, từ đó sử dụng game thuận tiện hơn."
        ),
    ),
    SectionSpec(
        heading="4.7. Giao diện challenge và chọn buff",
        next_heading="4.8. Giao diện nhập tên và leaderboard",
        description=(
            "Mục này gồm hai giao diện tác động trực tiếp tới tiến trình trận đấu: ChallengePanel trước khi bắt đầu combat và CardSelectionPanel khi người chơi lên cấp. "
            "Đây là hai màn hình mà người dùng cần đưa ra lựa chọn để tiếp tục lượt chơi."
        ),
        usage=(
            "Người chơi tiếp cận NPC challenge, nhấn F để mở panel và bấm nút bắt đầu trận đấu sau phần giới thiệu ngắn. Trong quá trình chiến đấu, mỗi lần lên cấp, "
            "CardSelectionPanel sẽ hiện 3 thẻ buff để người dùng chọn 1 nâng cấp, lúc đó game tạm dừng cho đến khi lựa chọn hoàn tất."
        ),
        functions=(
            "ChallengePanel dùng để kích hoạt run, còn CardSelectionPanel dùng để tạo nhánh phát triển sức mạnh theo từng lượt chơi. Hai giao diện này biến gameplay "
            "từ một vòng lặp chiến đấu đơn thuần thành một trải nghiệm có quyết định chiến thuật."
        ),
        result=(
            "Người chơi chủ động khởi đầu trận và định hình build nhân vật theo từng lần lên cấp, nhờ đó mỗi lượt chơi có thể tạo ra cách tiếp cận khác nhau."
        ),
    ),
    SectionSpec(
        heading="4.8. Giao diện nhập tên và leaderboard",
        next_heading="4.9. Giao diện tạm dừng, cài đặt âm thanh và loading",
        description=(
            "Mục này trình bày giao diện nhập tên người chơi và giao diện bảng xếp hạng kết nối với PlayFab. Người dùng nhìn thấy tên hiển thị của mình, danh sách xếp hạng "
            "và các nút điều hướng phục vụ phần hậu trận."
        ),
        usage=(
            "Ở lần vào game đầu tiên, nếu chưa có Display Name, hệ thống sẽ hiện form nhập tên và yêu cầu tên dài từ 3 đến 25 ký tự. Sau khi kết thúc trận hoặc mở từ menu, "
            "người dùng có thể xem leaderboard để so sánh điểm số được gửi lên statistic HighScore."
        ),
        functions=(
            "Giao diện này dùng để định danh người chơi và hiển thị kết quả cạnh tranh giữa nhiều lần chơi. Tên hiển thị được lưu cùng tài khoản PlayFab đăng nhập bằng CustomId, "
            "trong khi bảng xếp hạng phản hồi trực tiếp dữ liệu score đã submit."
        ),
        result=(
            "Người dùng có một danh tính ổn định trong game và có thể theo dõi thành tích của mình trên bảng xếp hạng ngay sau khi hoàn thành một lượt chơi."
        ),
    ),
    SectionSpec(
        heading="4.9. Giao diện tạm dừng, cài đặt âm thanh và loading",
        next_heading="KẾT LUẬN",
        description=(
            "Mục cuối của Chương 4 mô tả các giao diện phụ trợ quan trọng như menu tạm dừng, phần cài đặt âm thanh và lớp loading/black fade. Dù không trực tiếp tạo combat, "
            "đây là nhóm giao diện ảnh hưởng mạnh đến cảm giác sử dụng của chương trình."
        ),
        usage=(
            "Trong lúc chơi, người dùng nhấn ESC để mở hoặc đóng menu tạm dừng. Tại đây có thể xem leaderboard, vào phần cài đặt để chỉnh riêng Music và SFX, hoặc khởi động lại lượt chơi; "
            "ngoài ra màn hình loading đen sẽ tự xuất hiện khi đổi theme hoặc reload scene."
        ),
        functions=(
            "PauseMenuPanel giúp dừng thời gian, khóa input gameplay và cung cấp các thao tác điều hướng an toàn. Loading/black fade được dùng để che quá trình chuyển trạng thái, "
            "giúp trải nghiệm mượt hơn khi restart hoặc chuyển môi trường."
        ),
        result=(
            "Người chơi có thể kiểm soát phiên chơi tốt hơn, tạm dừng đúng lúc, chỉnh âm thanh nhanh và vẫn giữ được cảm giác chuyển cảnh ổn định khi hệ thống thay đổi trạng thái."
        ),
    ),
]


FIGURE_CAPTIONS = [
    "Hình 4.1: Toàn cảnh scene Game trước khi bắt đầu trận đấu.",
    "Hình 4.2: Toàn cảnh gameplay khi wave chiến đấu đang diễn ra.",
    "Hình 4.3: Nhân vật chính trong không gian chiến đấu của trò chơi.",
    "Hình 4.4: Enemy hệ bay trong gameplay.",
    "Hình 4.5: Enemy tấn công tầm xa.",
    "Hình 4.6: Enemy cận chiến áp sát người chơi.",
    "Hình 4.7: Nhóm enemy nguyên tố LawaChurl.",
    "Hình 4.8: Theme bản đồ thứ nhất.",
    "Hình 4.9: Theme bản đồ thứ hai.",
    "Hình 4.10: Theme bản đồ thứ ba.",
    "Hình 4.11: Prompt mở menu bằng phím ESC trên HUD.",
    "Hình 4.12: Chỉ báo wave hiện tại trên HUD.",
    "Hình 4.13: Thanh máu, thanh kinh nghiệm và cấp độ của người chơi.",
    "Hình 4.14: ChallengePanel trước khi bắt đầu trận đấu.",
    "Hình 4.15: CardSelectionPanel khi người chơi lên cấp.",
    "Hình 4.16: Giao diện hiển thị tên người chơi trong game.",
    "Hình 4.17: Giao diện bảng xếp hạng theo điểm số.",
    "Hình 4.18: Giao diện nhập tên ở lần vào game đầu tiên.",
    "Hình 4.19: Giao diện cài đặt âm thanh trong menu tạm dừng.",
    "Hình 4.20: Giao diện tạm dừng trận đấu.",
]


def set_run_font(run, size: int = 13, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size: int = 13, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def replace_paragraph_text(
    paragraph,
    text: str,
    *,
    size: int = 13,
    bold: bool | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if alignment is not None:
        paragraph.alignment = alignment


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def has_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//*[local-name()='drawing']"))


def is_figure_caption(text: str) -> bool:
    return bool(re.match(r"^Hình\s+\d", text))


def find_paragraph(document: Document, text: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def paragraphs_between(document: Document, start_text: str, end_text: str):
    collecting = False
    items = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == start_text:
            collecting = True
            continue
        if collecting and text == end_text:
            break
        if collecting:
            items.append(paragraph)
    return items


def add_labeled_paragraph(anchor, label: str, text: str):
    paragraph = insert_paragraph_after(anchor, style="Normal")
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=13, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=13, bold=False)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def rewrite_chapter_intro(document: Document) -> None:
    chapter_heading = find_paragraph(document, CHAPTER4_HEADING)
    for paragraph in list(paragraphs_between(document, CHAPTER4_HEADING, CHAPTER4_FIRST_SECTION)):
        remove_paragraph(paragraph)

    intro = insert_paragraph_after(
        chapter_heading,
        (
            "Chương này trình bày các giao diện chính của chương trình và cách sử dụng những chức năng quan trọng "
            "trong quá trình chơi. Các hình minh họa được giữ theo đúng bố cục triển khai thực tế để người đọc dễ "
            "đối chiếu giữa thao tác của người dùng và phản hồi hiển thị trên màn hình."
        ),
        style="Normal",
    )
    set_paragraph_font(intro, size=13)
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def cleanup_section(document: Document, heading: str, next_heading: str):
    for paragraph in list(paragraphs_between(document, heading, next_heading)):
        text = paragraph.text.strip()
        if has_drawing(paragraph) or is_figure_caption(text):
            continue
        remove_paragraph(paragraph)


def rewrite_sections(document: Document) -> None:
    for spec in SECTION_SPECS:
        cleanup_section(document, spec.heading, spec.next_heading)
        heading_paragraph = find_paragraph(document, spec.heading)
        anchor = heading_paragraph
        anchor = add_labeled_paragraph(anchor, "Mô tả giao diện", spec.description)
        anchor = add_labeled_paragraph(anchor, "Cách sử dụng", spec.usage)
        anchor = add_labeled_paragraph(anchor, "Chức năng chính", spec.functions)
        add_labeled_paragraph(anchor, "Kết quả sử dụng", spec.result)


def rewrite_figure_captions(document: Document) -> None:
    chapter4_captions = []
    in_chapter4 = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == CHAPTER4_HEADING:
            in_chapter4 = True
            continue
        if in_chapter4 and text == CHAPTER4_END_HEADING:
            break
        if in_chapter4 and is_figure_caption(text):
            chapter4_captions.append(paragraph)

    if len(chapter4_captions) != len(FIGURE_CAPTIONS):
        raise ValueError(
            f"Expected {len(FIGURE_CAPTIONS)} captions in Chapter 4, found {len(chapter4_captions)}"
        )

    for paragraph, caption in zip(chapter4_captions, FIGURE_CAPTIONS, strict=True):
        replace_paragraph_text(
            paragraph,
            caption,
            size=12,
            bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )


def normalize_chapter4_text_black(document: Document) -> int:
    in_chapter4 = False
    non_black = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == CHAPTER4_HEADING:
            in_chapter4 = True
        if in_chapter4:
            for run in paragraph.runs:
                if run.font.color.rgb not in (None, BLACK):
                    non_black += 1
                run.font.color.rgb = BLACK
        if in_chapter4 and text == CHAPTER4_END_HEADING:
            in_chapter4 = False
    return non_black


def verify_structure(document: Document) -> None:
    for spec in SECTION_SPECS:
        find_paragraph(document, spec.heading)

    find_paragraph(document, CHAPTER4_END_HEADING)

    labels = {
        "Mô tả giao diện:",
        "Cách sử dụng:",
        "Chức năng chính:",
        "Kết quả sử dụng:",
    }
    found_labels = 0
    in_chapter4 = False
    figure_count = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == CHAPTER4_HEADING:
            in_chapter4 = True
            continue
        if in_chapter4 and text == CHAPTER4_END_HEADING:
            break
        if not in_chapter4:
            continue
        if any(text.startswith(label) for label in labels):
            found_labels += 1
        if is_figure_caption(text) and text.startswith("Hình 4."):
            figure_count += 1

    expected_labels = len(SECTION_SPECS) * 4
    if found_labels != expected_labels:
        raise ValueError(f"Expected {expected_labels} labeled paragraphs, found {found_labels}")
    if figure_count != len(FIGURE_CAPTIONS):
        raise ValueError(f"Expected {len(FIGURE_CAPTIONS)} Chapter 4 figure captions, found {figure_count}")


def main() -> None:
    if not SOURCE_DOC.exists():
        raise FileNotFoundError(SOURCE_DOC)

    BACKUP_DOC.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)

    shutil.copyfile(SOURCE_DOC, BACKUP_DOC)
    shutil.copyfile(SOURCE_DOC, OUTPUT_DOC)

    document = Document(OUTPUT_DOC)
    rewrite_chapter_intro(document)
    rewrite_sections(document)
    rewrite_figure_captions(document)
    non_black_before_fix = normalize_chapter4_text_black(document)
    verify_structure(document)
    document.save(OUTPUT_DOC)

    print(f"output={OUTPUT_DOC}")
    print(f"backup={BACKUP_DOC}")
    print(f"chapter4_non_black_before_fix={non_black_before_fix}")
    print(f"chapter4_figures={len(FIGURE_CAPTIONS)}")
    print(f"chapter4_sections={len(SECTION_SPECS)}")


if __name__ == "__main__":
    main()
