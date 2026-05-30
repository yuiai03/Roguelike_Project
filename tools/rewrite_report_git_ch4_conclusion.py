from __future__ import annotations

import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = ROOT / "tmp" / "docs" / "BaoCao_DATN_2_current_for_git_ch4_conclusion.docx"
OUTPUT_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_git_ch4_conclusion_updated.docx"
BACKUP_DOC = ROOT / "tmp" / "docs" / "BaoCao_DATN_2_before_git_ch4_conclusion.docx"

BLACK = RGBColor(0, 0, 0)


@dataclass(frozen=True)
class TableSpec:
    caption: str
    headers: list[str]
    rows: list[list[str]]
    widths_cm: list[float]
    font_size: int = 11


@dataclass(frozen=True)
class DetailSection:
    prefix: str
    next_prefix: str
    blocks: list[tuple[str, str]]


def set_run_font(run, size: int = 13, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size: int = 13, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def replace_paragraph_text(
    paragraph,
    text: str,
    *,
    size: int = 13,
    bold: bool | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if alignment is not None:
        paragraph.alignment = alignment


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def insert_table_after(paragraph, rows: int, cols: int, width_cm: float = 16.0):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Cm(width_cm))
    paragraph._p.addnext(table._tbl)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    return table


def insert_paragraph_after_table(table, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    new_para = table._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def has_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//*[local-name()='drawing']"))


def is_caption(paragraph) -> bool:
    style_name = paragraph.style.name if paragraph.style else ""
    text = paragraph.text.strip()
    return style_name == "Caption" or text.startswith("Hình ") or text.startswith("Bảng ")


def find_paragraph_by_prefix(document: Document, prefix: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Could not find paragraph with prefix: {prefix}")


def find_paragraph_by_prefix_after(document: Document, prefix: str, after_paragraph):
    seen = False
    for paragraph in document.paragraphs:
        if paragraph._p is after_paragraph._p:
            seen = True
            continue
        if seen and paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Could not find paragraph with prefix {prefix} after anchor.")


def find_next_heading1_after(document: Document, after_paragraph):
    seen = False
    for paragraph in document.paragraphs:
        if paragraph._p is after_paragraph._p:
            seen = True
            continue
        if seen and paragraph.style and paragraph.style.name == "Heading 1":
            return paragraph
    raise ValueError("Could not find next Heading 1.")


def paragraphs_between(document: Document, start_paragraph, end_paragraph):
    collecting = False
    collected = []
    for paragraph in document.paragraphs:
        if paragraph._p is start_paragraph._p:
            collecting = True
            continue
        if collecting and paragraph._p is end_paragraph._p:
            break
        if collecting:
            collected.append(paragraph)
    return collected


def remove_non_visual_content(document: Document, start_prefix: str, next_prefix: str) -> None:
    start = find_paragraph_by_prefix(document, start_prefix)
    end = find_paragraph_by_prefix(document, next_prefix)
    for paragraph in list(paragraphs_between(document, start, end)):
        if has_drawing(paragraph) or is_caption(paragraph):
            continue
        remove_paragraph(paragraph)


def add_labeled_blocks(anchor, blocks: list[tuple[str, str]]):
    current = anchor
    for label, text in blocks:
        para = insert_paragraph_after(current, style="Normal")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        label_run = para.add_run(f"{label}: ")
        set_run_font(label_run, size=13, bold=True)
        text_run = para.add_run(text)
        set_run_font(text_run, size=13)
        current = para
    return current


def add_table(anchor, spec: TableSpec):
    caption = insert_paragraph_after(anchor, spec.caption, style="Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(caption, size=12, bold=True)

    table = insert_table_after(caption, rows=1 + len(spec.rows), cols=len(spec.headers), width_cm=sum(spec.widths_cm))
    for idx, header in enumerate(spec.headers):
        table.rows[0].cells[idx].text = header

    for row_idx, row_values in enumerate(spec.rows, start=1):
        for col_idx, value in enumerate(row_values):
            table.rows[row_idx].cells[col_idx].text = value

    for row in table.rows:
        for idx, width in enumerate(spec.widths_cm):
            row.cells[idx].width = Cm(width)

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    set_run_font(run, size=spec.font_size, bold=(row_idx == 0))

    return table


GIT_TABLE = TableSpec(
    caption="Bảng 2.1: Quy trình quản lý mã nguồn bằng Git trong dự án.",
    headers=["Bước", "Cách nhóm thực hiện", "Ý nghĩa quản lý project"],
    rows=[
        ["Cập nhật local", "Pull hoặc lấy bản mới nhất trước khi sửa code.", "Giảm nguy cơ làm việc trên phiên bản cũ."],
        ["Phát triển chức năng", "Chỉnh script C#, prefab, ScriptableObject hoặc tài liệu theo từng nhiệm vụ.", "Tách thay đổi theo phạm vi rõ ràng, dễ kiểm soát."],
        ["Commit", "Lưu mốc sau mỗi nhóm chức năng như spawn tool, wave config, effect hoặc map delay.", "Giữ lịch sử sửa đổi để xem lại và khôi phục khi cần."],
        ["Đối chiếu lịch sử", "Dùng log để so sánh các mốc chỉnh sửa gần nhau.", "Hỗ trợ truy vết lỗi và đánh giá tiến độ thực hiện."],
    ],
    widths_cm=[3.0, 6.5, 6.5],
)

GITHUB_TABLE = TableSpec(
    caption="Bảng 2.2: Vai trò của Git local và GitHub remote trong quản lý project.",
    headers=["Thành phần", "Dữ liệu thực tế của dự án", "Vai trò sử dụng"],
    rows=[
        ["Git local", "Nhánh chính hiện tại là main.", "Quản lý chỉnh sửa trên máy, commit và kiểm tra lịch sử trước khi đồng bộ."],
        ["GitHub remote", "origin -> https://github.com/yuiai03/Roguelike_Project.git", "Làm kho trung tâm để chia sẻ mã nguồn và đồng bộ giữa các máy."],
        ["Commit history", "Ví dụ: update effect, update spawn tool, update hiển thị wave và tool config.", "Phản ánh các mốc triển khai thật của gameplay và công cụ nội bộ."],
        ["Contributor", "Local history hiện thấy lseanl03 và YuiAI.", "Cho phép theo dõi người tham gia và mức độ đóng góp vào project."],
    ],
    widths_cm=[3.0, 6.5, 6.5],
)


SECTION_4_SPECS = [
    DetailSection(
        prefix="4.1.",
        next_prefix="4.2.",
        blocks=[
            ("Mô tả giao diện", "Hai hình ở mục này trình bày scene Game ở hai trạng thái quan trọng nhất của một lượt chơi: trước khi vào combat và khi wave đã được kích hoạt. Đây là lớp giao diện tổng hợp toàn bộ không gian chiến đấu, các đối tượng gameplay và vùng thông tin mà người chơi phải quan sát liên tục."),
            ("Đối tượng hiển thị trong hình", "Hình 4.1 cho thấy nhân vật người chơi, NPC challenge dùng để bắt đầu trận, mặt bằng chiến đấu, tường giới hạn map và HUD ở trạng thái chờ. Hình 4.2 bổ sung thêm enemy đang hoạt động, projectile đang bay, chỉ báo wave và các hiệu ứng combat, qua đó thể hiện rõ sự chuyển trạng thái từ chuẩn bị sang chiến đấu runtime."),
            ("Cách sử dụng", "Người chơi dùng WASD hoặc phím mũi tên để di chuyển trong khu vực map, tiếp cận NPC challenge và nhấn F để khởi động trận. Sau khi run bắt đầu, thao tác chính chuyển sang giữ vị trí, né đòn bằng chuột phải để dash, đồng thời theo dõi HUD để phản ứng với tình trạng của nhân vật."),
            ("Chức năng chính", "Scene Game là nơi tích hợp trọn vẹn các hệ gameplay như điều khiển nhân vật, spawn enemy theo wave, tấn công tự động, nhận EXP, đổi theme map, chọn buff và cập nhật leaderboard sau trận. Vì vậy đây là màn hình phản ánh đầy đủ nhất kiến trúc runtime của chương trình."),
            ("Thông số liên quan", "Run hiện dùng wave tăng tiến theo WaveSpawner với scalePerWave = 1.1; player khởi đầu từ PlayerConfig.asset với maxHealth = 1000, attackDamage = 50, projectileSpeed = 30 và projectileLifetime = 10; wave 1 trong WaveConfig.asset gồm 4 cụm enemy melee, mỗi cụm 3 enemy, spawnDelay từ 0 đến 3 giây."),
            ("Kết quả sử dụng", "Người đọc có thể thấy rõ toàn bộ vòng lặp của trò chơi được gom vào một scene duy nhất: bắt đầu từ không gian chờ, chuyển sang combat theo wave, tích lũy tài nguyên và liên tục thay đổi trạng thái cho đến khi kết thúc trận."),
        ],
    ),
    DetailSection(
        prefix="4.2.",
        next_prefix="4.3.",
        blocks=[
            ("Mô tả giao diện", "Mục này tập trung vào đối tượng player, là trung tâm của mọi thao tác điều khiển và là đầu mối liên kết giữa combat, level-up, buff và HUD."),
            ("Đối tượng hiển thị trong hình", "Trong hình có thể quan sát model nhân vật chính, hướng quay theo trục di chuyển, hiệu ứng bắn tự động khi phát hiện enemy và không gian trống xung quanh dùng để né tránh. Nếu ảnh chụp trong lúc chiến đấu, projectile hoặc muzzle effect cũng là thành phần quan trọng để nhận diện cơ chế auto attack."),
            ("Cách sử dụng", "Người chơi di chuyển bằng WASD hoặc phím mũi tên, dùng chuột phải để dash qua vùng nguy hiểm. Phần tấn công không cần bấm thủ công; PlayerAttack tự tìm mục tiêu gần nhất trong tầm và bắn projectile, nên thao tác của người chơi tập trung vào vị trí, khoảng cách và thời điểm né đòn."),
            ("Chức năng chính", "Player vừa là đối tượng di chuyển trong scene, vừa là nơi gắn toàn bộ dữ liệu chiến đấu như máu, sát thương, tốc độ bắn, tốc độ đạn và các chỉ số được buff sửa đổi theo thời gian thực."),
            ("Thông số liên quan", "Theo PlayerConfig.asset, moveSpeed = 7, dashSpeed = 20, dashDuration = 0.5 giây, dashCooldown = 1 giây, maxHealth = 1000, attackDamage = 50, projectileSpeed = 30 và projectileLifetime = 10. Trong PlayerData, multiShotAngle mặc định là 12 độ và các bonus như moveSpeedBonus hay damageBonus sẽ cộng trực tiếp vào bộ chỉ số gốc này."),
            ("Kết quả sử dụng", "Người chơi không phải thao tác tấn công phức tạp mà có thể tập trung tối đa vào di chuyển chiến thuật, trong khi hệ chỉ số vẫn đủ chi tiết để tạo khác biệt rõ ràng sau mỗi lần nâng cấp buff."),
        ],
    ),
    DetailSection(
        prefix="4.3.",
        next_prefix="4.4.",
        blocks=[
            ("Mô tả giao diện", "Ba hình trong mục này được dùng để giới thiệu ba nhóm enemy cơ bản của gameplay: quái bay, quái đánh xa và quái cận chiến. Đây là các đối tượng mà người chơi gặp thường xuyên nhất trong các wave chuẩn."),
            ("Đối tượng hiển thị trong hình", "Hình 4.4 là FlyEnemy với vị trí bay lơ lửng và quỹ đạo di chuyển linh hoạt. Hình 4.5 là RangedEnemy, thường giữ khoảng cách và gây áp lực bằng projectile. Hình 4.6 là MeleeEnemy, chủ yếu áp sát trực diện để ép người chơi phải đổi vị trí liên tục."),
            ("Cách sử dụng", "Người chơi đọc loại enemy qua hình dáng và cách tiếp cận của chúng để đổi hướng di chuyển. Với MeleeEnemy cần mở khoảng cách để tránh contact damage; với RangedEnemy cần thay đổi vị trí đều để né đạn; với FlyEnemy phải quan sát thêm hướng xuất hiện vì chúng tấn công từ không gian ít bị chắn hơn."),
            ("Chức năng chính", "Ba enemy cơ bản tạo nền cho hầu hết các wave thường, đóng vai trò cung cấp EXP, tạo áp lực nhịp độ và buộc người chơi phải phối hợp né đòn với việc chọn buff hợp lý."),
            ("Thông số liên quan", "MeleeEnemy_.asset có maxHealth = 100, moveSpeed = 6, attackRange = 2, contactDamage = 10 và expValue = 20. RangedEnemy_.asset có maxHealth = 100, moveSpeed = 8, attackRange = 20, projectileDamage = 15, projectileSpeed = 30, projectileLifetime = 10 và expValue = 50. FlyEnemy_.asset có maxHealth = 100, moveSpeed = 8, attackRange = 50, projectileDamage = 10, projectileSpeed = 30, projectileLifetime = 5, burstCount = 6 và expValue = 50."),
            ("Kết quả sử dụng", "Ngay từ các wave đầu, người chơi đã phải phân loại mối đe dọa theo hành vi thật chứ không chỉ theo ngoại hình, từ đó tạo nền cho việc lựa chọn buff và kiểm soát không gian chiến đấu ở các wave khó hơn."),
        ],
    ),
    DetailSection(
        prefix="4.4.",
        next_prefix="4.5.",
        blocks=[
            ("Mô tả giao diện", "Mục này trình bày nhóm LawaChurl, là các enemy cỡ lớn dùng như boss hoặc elite trong các nhịp combat cao trào. Chúng khác nhóm quái cơ bản ở kích thước, animation, hiệu ứng nguyên tố và khả năng đổi phase."),
            ("Đối tượng hiển thị trong hình", "Trong ảnh có ba biến thể LawaChurl gồm Geo, Pyro và Electro. Điểm dễ nhận biết là thân hình lớn, hiệu ứng nguyên tố quanh mô hình, vùng tấn công rộng và sự hiện diện nổi bật hơn hẳn so với enemy thường trong cùng scene."),
            ("Cách sử dụng", "Khi gặp LawaChurl, người chơi không thể đứng giữ vị trí lâu mà phải vừa đọc chuyển động vừa canh chuột phải để dash. Các boss này nên được ưu tiên quan sát trước, vì mỗi phase mới sẽ làm nhịp di chuyển và vùng nguy hiểm thay đổi đáng kể."),
            ("Chức năng chính", "LawaChurl là mốc tăng độ khó, giúp game tạo ra boss wave rõ ràng và kiểm tra mức độ tối ưu build của người chơi sau nhiều lần cộng buff. Chúng cũng là đối tượng làm rõ nhất cơ chế OnPhaseChanged trong hệ boss."),
            ("Thông số liên quan", "Ba boss config hiện đều dùng maxHealth = 1000, moveSpeed = 5, phase2Threshold = 0.6, phase3Threshold = 0.3, phase2SpeedMult = 1.3 và phase3SpeedMult = 1.6. Boss projectileDamage = 25, bossProjectileLifetime = 6; tốc độ đạn của Geo là 12, còn Pyro và Electro là 15. Mỗi boss chết sẽ kích hoạt buffCardDropCount = 2."),
            ("Kết quả sử dụng", "Sự xuất hiện của nhóm LawaChurl làm cho combat không chỉ tăng số lượng enemy mà còn tăng chiều sâu về nhịp đọc tình huống, giúp game có điểm nhấn rõ ràng giữa các giai đoạn của lượt chơi."),
        ],
    ),
    DetailSection(
        prefix="4.5.",
        next_prefix="4.6.",
        blocks=[
            ("Mô tả giao diện", "Ba hình trong mục này thể hiện ba theme môi trường đang được áp dụng cho scene Game. Việc thay đổi theme không chỉ là đổi màu nền mà còn thay đổi vật liệu mặt đất, vật liệu tường và hệ effect đi kèm."),
            ("Đối tượng hiển thị trong hình", "Người đọc có thể quan sát khác biệt giữa nền đất, tường bao, ánh sáng tổng thể và hiệu ứng trang trí ở từng ảnh. Các thành phần này đều thuộc về cùng một map nhưng được thay vật liệu và effectRoot để tạo cảm giác chuyển giai đoạn gameplay."),
            ("Cách sử dụng", "Người chơi không cần thao tác thủ công để đổi map. MapThemeManager tự động xác định theme theo wave, phát black transition qua LoadingUIManager rồi áp dụng material và effect mới trong lúc gameplay bị khóa ngắn hạn."),
            ("Chức năng chính", "Nhóm theme map giúp phân đoạn lượt chơi theo mốc wave, làm mới cảm giác thị giác và hỗ trợ người chơi nhận ra rằng run đã bước sang giai đoạn mới, đặc biệt ở những lần chuyển sau nhiều wave liên tiếp."),
            ("Thông số liên quan", "Scene hiện có 3 theme được cấu hình là Night, AfterNight và Evening. MapThemeManager dùng fadeInDuration = 0.4, holdDuration = 0.2 và fadeOutDuration = 0.4; theme index được tính theo công thức ((wave - 1) / 10) % themes.Length, nên về thực tế mỗi block 10 wave sẽ tương ứng với một theme khác nhau."),
            ("Kết quả sử dụng", "Người chơi cảm nhận được sự tiến triển của run không chỉ qua độ khó mà còn qua thay đổi môi trường, giúp trận đấu bớt lặp và tăng cảm giác hoàn thiện cho sản phẩm."),
        ],
    ),
    DetailSection(
        prefix="4.6.",
        next_prefix="4.7.",
        blocks=[
            ("Mô tả giao diện", "HUD là lớp giao diện xuất hiện xuyên suốt trong scene Game để người chơi đọc trạng thái của nhân vật và tiến độ trận đấu mà không phải mở thêm panel phụ."),
            ("Đối tượng hiển thị trong hình", "Hình 4.11 thể hiện prompt mở menu bằng ESC, giúp người chơi nhận biết thao tác tạm dừng. Hình 4.12 tập trung vào chỉ báo wave hiện tại. Hình 4.13 mô tả cụm thông tin sống còn gồm thanh máu, thanh kinh nghiệm và cấp độ hiện tại của player."),
            ("Cách sử dụng", "Người chơi cần nhìn HUD liên tục để quyết định có nên lùi ra giữ khoảng cách, có sắp lên cấp hay chưa và trận đấu đang ở wave nào. Trong những lúc áp lực cao, chỉ riêng việc đọc đúng ba cụm HP, EXP và wave đã đủ ảnh hưởng trực tiếp tới cách ra quyết định."),
            ("Chức năng chính", "HUD gom dữ liệu runtime từ PlayerStatsPanel, WaveSpawner và GameUI để hiển thị theo thời gian thực. Đây là lớp giao diện phản hồi nhanh nhất mọi thay đổi trong combat, từ máu giảm đến level tăng hoặc wave hoàn tất."),
            ("Thông số liên quan", "Máu gốc của người chơi bắt đầu ở 1000 và được hiển thị dưới dạng currentHealth/maxHealth. Hệ level khởi tạo expToNextLevel = 100 và tăng theo expScalingFactor = 1.1. Chỉ báo wave bám theo biến currentWave của WaveSpawner, còn prompt ESC được dùng để mở PauseMenuPanel trong mọi thời điểm cho phép."),
            ("Kết quả sử dụng", "HUD giúp người chơi đọc đúng trạng thái hiện thời mà không bị gián đoạn nhịp combat, qua đó làm trải nghiệm game rõ ràng và dễ kiểm soát hơn đáng kể."),
        ],
    ),
    DetailSection(
        prefix="4.7.",
        next_prefix="4.8.",
        blocks=[
            ("Mô tả giao diện", "Mục này gồm hai giao diện có ảnh hưởng trực tiếp tới tiến trình run: ChallengePanel để khởi động trận và CardSelectionPanel để chọn buff khi lên cấp."),
            ("Đối tượng hiển thị trong hình", "Hình 4.14 thường hiển thị NPC challenge, nội dung giới thiệu ngắn và nút bắt đầu trận. Hình 4.15 hiển thị ba thẻ buff cùng tên, mô tả, rarity và thông số mà người chơi sẽ nhận nếu chọn thẻ đó."),
            ("Cách sử dụng", "Người chơi tiếp cận NPC challenge, nhấn F rồi xác nhận bắt đầu để hệ thống vào run. Trong quá trình chiến đấu, mỗi lần OnLevelUp được phát, CardSelectionPanel mở ra với 3 lựa chọn; game tạm dừng đến khi người chơi chọn xong 1 thẻ."),
            ("Chức năng chính", "ChallengePanel kiểm soát điểm bắt đầu của gameplay, còn CardSelectionPanel quyết định nhánh phát triển sức mạnh của player. Hai giao diện này biến mỗi lượt chơi thành một run có lựa chọn chiến thuật thay vì chỉ là chiến đấu lặp lại."),
            ("Thông số liên quan", "BuffType hiện có 15 loại trong enum và 13 asset thẻ đã cấu hình sẵn trong thư mục BuffCard. RarityType gồm 4 mức là Common, Rare, Epic và Legendary. Một số thẻ tiêu biểu: MultiShot dùng shotCount để cộng thêm đạn và người chơi mặc định có multiShotAngle = 12 độ; AoEExplosion bật nổ diện rộng với aoeRadius mặc định 2; các spirit buff dùng attackDamageMultiplier để tính theo phần trăm ATK."),
            ("Kết quả sử dụng", "Người chơi được quyền điều chỉnh build theo tình huống thực tế của run, ví dụ ưu tiên damage, tốc độ đạn, tốc độ di chuyển hay kỹ năng phụ, nhờ đó mỗi lần chơi có thể cho ra cách tiếp cận khác nhau."),
        ],
    ),
    DetailSection(
        prefix="4.8.",
        next_prefix="4.9.",
        blocks=[
            ("Mô tả giao diện", "Mục này trình bày giao diện định danh người chơi và giao diện xếp hạng trực tuyến, là phần kết nối trực tiếp giữa client gameplay với backend PlayFab."),
            ("Đối tượng hiển thị trong hình", "Hình 4.16 thể hiện tên hiển thị của người chơi sau khi đăng nhập. Hình 4.17 là danh sách xếp hạng theo điểm số. Hình 4.18 là form nhập tên trong lần đầu vào game khi hệ thống chưa có Display Name hợp lệ."),
            ("Cách sử dụng", "Nếu người chơi chưa có tên, NameInputPanel sẽ yêu cầu nhập trước khi gameplay hoàn chỉnh. Sau khi kết thúc trận hoặc mở từ menu, LeaderboardPanel sẽ đọc dữ liệu top và hiển thị thứ hạng dựa trên điểm đã gửi lên PlayFab."),
            ("Chức năng chính", "Hai panel này giải quyết hai nhu cầu: định danh ổn định cho từng người chơi và hiển thị kết quả cạnh tranh sau mỗi lượt chơi. Đây là cầu nối quan trọng để sản phẩm không dừng ở mức chơi cục bộ đơn lẻ."),
            ("Thông số liên quan", "LeaderboardStatisticName hiện cấu hình là HighScore. Tên hợp lệ sau khi trim phải dài từ 3 đến 25 ký tự. CustomId được lấy từ PlayerPrefs hoặc sinh từ SystemInfo.deviceUniqueIdentifier nếu chưa có dữ liệu cũ; điểm gửi lên backend lấy từ finalScore/currentRunScore của lượt chơi."),
            ("Kết quả sử dụng", "Người chơi có một danh tính nhất quán trong hệ thống và có thể theo dõi thành tích của mình qua bảng xếp hạng ngay sau khi kết thúc trận, làm tăng cảm giác hoàn chỉnh của trò chơi."),
        ],
    ),
    DetailSection(
        prefix="4.9.",
        next_prefix="__CONCLUSION__",
        blocks=[
            ("Mô tả giao diện", "Đây là nhóm giao diện phụ trợ nhưng ảnh hưởng mạnh đến trải nghiệm sử dụng, gồm menu tạm dừng, phần chỉnh âm thanh và lớp loading/black fade dùng khi chuyển trạng thái."),
            ("Đối tượng hiển thị trong hình", "Trong hình cài đặt, các thành phần quan trọng là slider âm lượng Music và SFX, nút điều hướng quay lại và khung nền che gameplay. Trong hình pause, người chơi nhìn thấy các lựa chọn tiếp tục, mở leaderboard, mở settings hoặc khởi động lại lượt chơi."),
            ("Cách sử dụng", "Người chơi nhấn ESC để mở hoặc đóng PauseMenuPanel. Từ đây có thể tạm dừng trận, chỉnh âm lượng, xem bảng xếp hạng hoặc thực hiện thao tác liên quan tới phiên chơi hiện tại. Khi game đổi theme hoặc reload scene, lớp loading đen sẽ tự xuất hiện để che quá trình chuyển cảnh."),
            ("Chức năng chính", "PauseMenuPanel khóa gameplay và đưa quyền điều hướng sang UI, còn LoadingPanel chịu trách nhiệm làm mượt các pha chuyển trạng thái. Nhờ đó game không bị cảm giác giật hoặc thay cảnh đột ngột khi hệ thống đổi theme hay restart."),
            ("Thông số liên quan", "MapThemeManager phối hợp với LoadingUIManager bằng bộ thời lượng fadeIn = 0.4 giây, hold = 0.2 giây và fadeOut = 0.4 giây. Việc mở menu gắn với phím ESC, còn phần âm thanh tách riêng Music và SFX để người chơi có thể điều chỉnh từng lớp âm thanh theo nhu cầu."),
            ("Kết quả sử dụng", "Người chơi kiểm soát tốt hơn nhịp chơi của mình, có thể dừng đúng lúc, chỉnh âm thanh nhanh và vẫn giữ được trải nghiệm chuyển cảnh ổn định trong những lần đổi map hoặc khởi động lại trận."),
        ],
    ),
]


def insert_git_sections(document: Document) -> None:
    start = find_paragraph_by_prefix(document, "2.2.3.")
    end = find_paragraph_by_prefix(document, "2.3.")
    between = paragraphs_between(document, start, end)
    anchor = between[-1] if between else start

    heading_224 = insert_paragraph_after(anchor, "2.2.4. Git trong quản lý phiên bản mã nguồn", style="Heading 3")
    set_paragraph_font(heading_224, size=13, bold=True)
    anchor = add_labeled_blocks(
        heading_224,
        [
            ("Mô tả", "Git là công cụ quản lý phiên bản được dùng để lưu vết thay đổi của toàn bộ mã nguồn, prefab, ScriptableObject, scene và tài liệu đi kèm trong suốt quá trình phát triển đề tài. Với một dự án Unity có nhiều asset và script phụ thuộc lẫn nhau, việc quản lý lịch sử bằng Git giúp giảm rủi ro ghi đè hoặc mất phiên bản ổn định."),
            ("Vai trò trong dự án", "Trong project này, Git được dùng để chia nhỏ quá trình phát triển thành các mốc kỹ thuật cụ thể như cập nhật hiệu ứng, điều chỉnh thời gian đổi map, bổ sung công cụ spawn hoặc hiển thị wave. Cách làm này giúp nhóm dễ kiểm tra phần nào đã sửa, phần nào đang phát sinh lỗi và phần nào có thể quay lui an toàn."),
            ("Quy trình sử dụng thực tế", "Nhóm thực hiện cập nhật mã nguồn theo chu trình: lấy bản mới nhất, chỉnh sửa trong Unity hoặc Visual Studio, kiểm tra logic cục bộ, sau đó commit theo từng nhóm chức năng. Những commit gần đây của chính repo cho thấy cách chia mốc triển khai khá rõ, ví dụ: update effect, update sửa thời gian delay đổi map, update hiển thị wave và tool config, update spawn tool."),
            ("Thông tin cụ thể", "Tại thời điểm rà soát, nhánh làm việc chính của repo là main. Lịch sử commit được lưu trực tiếp trên máy local để hỗ trợ đối chiếu lại thay đổi theo từng ngày làm việc, đồng thời là cơ sở để phục hồi khi một chỉnh sửa mới gây ảnh hưởng xấu đến gameplay hoặc giao diện."),
            ("Kết quả đạt được", "Việc dùng Git giúp quá trình phát triển có tính kiểm soát hơn, đặc biệt khi project liên quan đồng thời đến gameplay, UI, backend PlayFab và tài liệu báo cáo. Nhờ đó nhóm có thể sửa đổi liên tục mà vẫn giữ được khả năng truy vết và khôi phục phiên bản ổn định."),
        ],
    )
    first_table = add_table(anchor, GIT_TABLE)

    heading_225 = insert_paragraph_after_table(first_table, "2.2.5. GitHub trong quản lý và phối hợp project", style="Heading 3")
    set_paragraph_font(heading_225, size=13, bold=True)
    anchor = add_labeled_blocks(
        heading_225,
        [
            ("Mô tả", "GitHub là nơi lưu trữ từ xa và đóng vai trò như kho mã nguồn trung tâm của dự án. Nếu Git hỗ trợ quản lý phiên bản cục bộ thì GitHub giúp đồng bộ, chia sẻ và theo dõi tiến độ làm việc nhóm trên cùng một dự án."),
            ("Vai trò trong dự án", "Đối với đề tài này, GitHub được dùng để lưu project Roguelike dưới dạng một repository thống nhất, giúp các thành viên có thể làm việc trên nhiều máy mà vẫn bám cùng một nguồn mã. Cách tổ chức này đặc biệt quan trọng với Unity vì cấu trúc dự án gồm nhiều thư mục asset, prefab, scene và script có liên hệ chặt chẽ."),
            ("Quy trình sử dụng thực tế", "Sau khi hoàn thành một nhóm chức năng trên máy local, thành viên sẽ push commit lên GitHub để tạo mốc đồng bộ chung. Các thành viên khác có thể pull về, tiếp tục chỉnh sửa hoặc đối chiếu lịch sử thay đổi trước khi tích hợp phần việc tiếp theo. Quy trình này giúp tránh ghi đè lẫn nhau và hỗ trợ kiểm soát tiến độ theo từng mốc commit."),
            ("Thông tin cụ thể", "Repository hiện tại của dự án là https://github.com/yuiai03/Roguelike_Project.git với remote tên origin. Nhánh chính đang dùng là main và lịch sử local hiện ghi nhận contributor lseanl03 cùng YuiAI, phản ánh đúng thực tế làm việc có nhiều người tham gia chỉnh sửa project."),
            ("Kết quả đạt được", "Nhờ GitHub, project không chỉ được lưu trữ an toàn hơn mà còn có môi trường phối hợp rõ ràng giữa các thành viên. Đây là yếu tố giúp nhóm quản lý tốt vòng đời mã nguồn từ phát triển tính năng, kiểm tra thay đổi cho đến hoàn thiện báo cáo và bàn giao sản phẩm."),
        ],
    )
    add_table(anchor, GITHUB_TABLE)


def cleanup_misc_chapter2(document: Document) -> None:
    start = find_paragraph_by_prefix(document, "2.3.3.")
    end = find_next_heading1_after(document, start)
    for paragraph in list(paragraphs_between(document, start, end)):
        text = paragraph.text.strip().lower()
        if "git" in text or "package manager" in text or "asset store" in text:
            remove_paragraph(paragraph)


def rewrite_section4(document: Document) -> None:
    for spec in SECTION_4_SPECS:
        if spec.next_prefix == "__CONCLUSION__":
            heading = find_paragraph_by_prefix(document, spec.prefix)
            next_paragraph = find_next_heading1_after(document, heading)
            for paragraph in list(paragraphs_between(document, heading, next_paragraph)):
                if has_drawing(paragraph) or is_caption(paragraph):
                    continue
                remove_paragraph(paragraph)
        else:
            next_paragraph = find_paragraph_by_prefix(document, spec.next_prefix)
            remove_non_visual_content(document, spec.prefix, spec.next_prefix)
        heading = find_paragraph_by_prefix(document, spec.prefix)
        add_labeled_blocks(heading, spec.blocks)


def rewrite_conclusion(document: Document) -> None:
    start = None
    end = None
    chapter4_heading = find_paragraph_by_prefix(document, "4.9.")
    passed_49 = False
    for paragraph in document.paragraphs:
        if paragraph._p is chapter4_heading._p:
            passed_49 = True
        if passed_49 and paragraph.style and paragraph.style.name == "Heading 1" and not paragraph.text.strip().startswith("CH"):
            if start is None:
                start = paragraph
                continue
            end = paragraph
            break
    if start is None or end is None:
        raise ValueError("Could not locate conclusion block.")

    for paragraph in list(paragraphs_between(document, start, end)):
        remove_paragraph(paragraph)

    intro = insert_paragraph_after(
        start,
        "Phần kết luận tổng hợp các kết quả chính mà đề tài đã đạt được, chỉ ra các giới hạn còn tồn tại và đề xuất hướng mở rộng tiếp theo để hoàn thiện sản phẩm theo đúng định hướng một game Roguelike 3D có khả năng phát triển dài hạn.",
        style="Normal",
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_font(intro, size=13)

    blocks = [
        (
            "Các vấn đề làm được",
            "Đề tài đã xây dựng được một khung game Roguelike 3D hoàn chỉnh ở mức đồ án với các chức năng runtime quan trọng: điều khiển nhân vật trong không gian 3D, dash né đòn, tấn công tự động, spawn enemy theo wave, tăng độ khó bằng scalePerWave, lên cấp và chọn buff, thay đổi theme map theo tiến trình wave, hệ HUD hiển thị máu - EXP - level - wave, menu tạm dừng và cài đặt âm thanh, cùng với hệ leaderboard kết nối PlayFab để lưu và hiển thị điểm HighScore. Về mặt kiến trúc, project cũng đã tách được các cụm lớp rõ ràng như player, enemy, projectile, wave, UI và backend, tạo nền tương đối tốt cho việc mở rộng sau này.",
        ),
        (
            "Các vấn đề chưa làm được",
            "Bên cạnh kết quả đạt được, project vẫn còn một số hạn chế. Nội dung gameplay hiện chưa thật sự phong phú vì số lượng enemy, boss pattern, map variation và nhánh phát triển build vẫn còn ít so với một sản phẩm hoàn chỉnh. Cân bằng độ khó giữa các wave và giữa các loại buff vẫn cần tiếp tục tinh chỉnh để tránh trường hợp một số lựa chọn quá mạnh hoặc quá yếu. Một số phần giao diện và tài nguyên trình bày tuy đã hoạt động đúng nhưng chưa đồng đều về mặt mỹ thuật và trải nghiệm người dùng. Ngoài ra, game hiện mới tập trung vào một vòng chơi đơn lẻ, chưa có các cơ chế metaprogression, save tiến trình dài hạn hoặc nhiều chế độ chơi khác nhau.",
        ),
        (
            "Định hướng phát triển của đề tài",
            "Trong giai đoạn tiếp theo, hướng phát triển phù hợp là hoàn thiện thêm MainMenu và luồng điều hướng tổng thể, mở rộng danh sách buff, enemy và boss để tăng độ đa dạng chiến thuật, bổ sung thêm theme map hoặc biến thể môi trường, đồng thời tối ưu hiệu năng ở các wave có mật độ spawn lớn. Về backend, có thể phát triển thêm các thống kê ngoài HighScore như thời gian sống sót, số enemy tiêu diệt hoặc mốc tiến trình. Về gameplay, nên bổ sung các lớp nâng cấp dài hạn, nhiều nhánh build hơn, thêm cơ chế progression giữa nhiều lần chơi và cải thiện sâu hơn trải nghiệm UI/UX để sản phẩm tiến gần hơn tới một game Roguelike thương mại ở mức hoàn thiện cao hơn.",
        ),
    ]

    add_labeled_blocks(intro, blocks)


def normalize_new_ranges_black(document: Document) -> None:
    in_chapter2 = False
    in_chapter4 = False
    in_conclusion = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("2.2.4.") or text.startswith("2.2.5."):
            in_chapter2 = True
        if text.startswith("2.3."):
            in_chapter2 = False
        if text.startswith("4.1."):
            in_chapter4 = True
        if paragraph.style and paragraph.style.name == "Heading 1" and not text.startswith("CH"):
            if not in_conclusion:
                in_chapter4 = False
                in_conclusion = True
            else:
                in_conclusion = False
        if paragraph.style and paragraph.style.name == "Heading 1" and text.startswith("CH"):
            in_conclusion = False

        if in_chapter2 or in_chapter4 or in_conclusion or text.startswith("Bảng 2.1:") or text.startswith("Bảng 2.2:"):
            for run in paragraph.runs:
                run.font.color.rgb = BLACK

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = BLACK


def main() -> None:
    shutil.copy2(SOURCE_DOC, BACKUP_DOC)
    shutil.copy2(SOURCE_DOC, OUTPUT_DOC)

    document = Document(OUTPUT_DOC)

    insert_git_sections(document)
    cleanup_misc_chapter2(document)
    rewrite_section4(document)
    rewrite_conclusion(document)
    normalize_new_ranges_black(document)

    document.save(OUTPUT_DOC)
    print(f"UPDATED_DOC={OUTPUT_DOC}")


if __name__ == "__main__":
    main()
