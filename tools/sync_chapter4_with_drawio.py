from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_drawio_aligned.docx"
TARGET_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_fully_synced.docx"


CH4_MAPPINGS = [
    (
        "4.1. Toàn cảnh gameplay trong scene Game",
        [
            "1_4_1_2_bat_dau_tran_dau.drawio",
            "1_4_1_3_player_controller.drawio",
            "1_4_1_4_chien_dau_player.drawio",
            "1_4_1_6_wave_va_do_kho.drawio",
            "3_3_activity_gameplay_core_loop.drawio",
        ],
        [
            "Phần toàn cảnh gameplay cần được hiểu như điểm giao giữa các sơ đồ nguồn về bắt đầu trận, điều khiển player, chiến đấu và wave. Khi người đọc nhìn vào hai ảnh tổng quan của scene Game, họ thực chất đang quan sát kết quả vận hành đồng thời của các cụm chức năng đã được tách riêng trong bộ draw.io.",
            "Việc đồng bộ phần mô tả với activity gameplay core loop giúp Chương 4 không chỉ là phần chèn screenshot, mà còn là nơi chứng minh các sơ đồ phân tích trước đó đã được hiện thực thành một vòng lặp chơi hoàn chỉnh trong runtime."
        ],
    ),
    (
        "4.2. Nhân vật chính của trò chơi",
        [
            "1_4_1_3_player_controller.drawio",
            "3_5_1_class_player_progression.drawio",
        ],
        [
            "Mục nhân vật chính được đồng bộ trực tiếp với cụm sơ đồ điều khiển player và class diagram Player & Progression. Hình ảnh trong phần này không chỉ minh họa model nhân vật mà còn đại diện cho các lớp chịu trách nhiệm di chuyển, dash, animation, máu và dữ liệu chiến đấu của player.",
            "Nhờ cách liên kết đó, các thông số runtime như moveSpeed, dashSpeed, dashCooldown, attackDamage và attackRange không còn đứng tách rời khỏi hình minh họa, mà trở thành phần giải thích kỹ thuật của chính đối tượng được chụp trong scene."
        ],
    ),
    (
        "4.3. Danh sách enemy cơ bản trong trò chơi",
        [
            "1_4_1_5_enemy_system.drawio",
            "3_5_2_class_enemy_projectile.drawio",
        ],
        [
            "Nhóm enemy cơ bản trong Chương 4 được đồng bộ với sơ đồ nhóm enemy và class diagram Enemy & Projectile. Điều này giúp phần minh họa quái cận chiến, đánh xa và quái bay không chỉ dừng ở giá trị hình ảnh, mà còn gắn trực tiếp với tuyến kế thừa và pipeline spawn - tấn công - OnDeath trong code.",
            "Theo cách viết mới, mỗi hình enemy là bằng chứng triển khai cho một vai trò chiến đấu cụ thể đã được phân tích ở Chương 3, từ đó làm rõ vì sao dự án cần tách các biến thể enemy theo hành vi thay vì chỉ khác nhau ở chỉ số."
        ],
    ),
    (
        "4.4. Nhóm enemy nguyên tố LawaChurl",
        [
            "1_4_1_5_enemy_system.drawio",
            "1_4_1_6_wave_va_do_kho.drawio",
            "3_5_2_class_enemy_projectile.drawio",
        ],
        [
            "Nhóm LawaChurl được đồng bộ với cụm enemy, wave và class diagram boss để phản ánh đúng vai trò của chúng trong các mốc boss wave. Các biến thể nguyên tố không chỉ khác về mặt tạo hình mà còn là đại diện cho nhánh BossEnemy với bộ projectile, hiệu ứng cảnh báo và logic áp lực diện rộng riêng.",
            "Khi đối chiếu với sơ đồ wave và boss wave, người đọc có thể hiểu rõ vì sao nhóm LawaChurl được tách khỏi enemy cơ bản: đây là cụm thực thể đánh dấu bước tăng nhịp độ trận đấu và làm thay đổi cách người chơi xử lý không gian chiến đấu."
        ],
    ),
    (
        "4.5. Các theme map và chuyển đổi môi trường",
        [
            "1_4_1_7_theme_ban_do.drawio",
            "3_4_1_component_gameplay_runtime.drawio",
        ],
        [
            "Mục theme map được đồng bộ trực tiếp với sơ đồ đổi theme bản đồ và component gameplay runtime. Nhờ đó, phần hình ảnh các map không còn là minh họa thị giác độc lập mà trở thành kết quả của một cụm xử lý rõ ràng gồm WaveSpawner, MapThemeManager và LoadingUIManager.",
            "Cách viết này giúp Chương 4 giải thích được vì sao theme chỉ đổi theo mốc wave và vì sao quá trình đổi môi trường phải đi kèm transition khóa input, thay vì chỉ mô tả ngắn rằng map có nhiều phong cách khác nhau."
        ],
    ),
    (
        "4.6. Giao diện HUD và hỗ trợ gameplay",
        [
            "1_4_1_10_ui_trong_tran.drawio",
            "3_4_2_component_progression_ui.drawio",
            "3_5_1_class_player_progression.drawio",
        ],
        [
            "HUD và UI runtime trong Chương 4 được đồng bộ với sơ đồ UI trong trận, component progression/UI và class diagram Player & Progression. Nhờ vậy, ảnh hiển thị HP, EXP, level và wave được liên kết thẳng với PlayerStatsPanel, GameUI và các event OnHealthChanged, OnExpChanged, OnWaveStart trong runtime.",
            "Việc đồng bộ cách diễn giải theo sơ đồ mới giúp phần giao diện không bị xem như lớp hiển thị phụ, mà trở thành một hệ thống phản hồi trạng thái có liên quan trực tiếp tới quyết định của người chơi trong từng giai đoạn của vòng lặp sinh tồn."
        ],
    ),
    (
        "4.7. Giao diện challenge và chọn buff",
        [
            "1_4_1_2_bat_dau_tran_dau.drawio",
            "1_4_1_9_buff_va_tang_suc_manh.drawio",
            "3_2_2_sequence_chien_dau_exp_buff.drawio",
        ],
        [
            "ChallengePanel và CardSelectionPanel là hai điểm chạm UI được đồng bộ trực tiếp với sơ đồ bắt đầu trận và sơ đồ buff. Điều này làm rõ hai vai trò rất khác nhau của chúng: ChallengePanel khởi động phiên chơi, còn CardSelectionPanel chèn quyết định chiến thuật vào giữa vòng lặp chiến đấu.",
            "Theo sequence chiến đấu - EXP - buff, việc hiển thị card sau OnLevelUp không chỉ là một màn hình nâng cấp đơn giản mà là bước chuyển trạng thái có chủ đích, nơi game dừng thời gian, khóa input gameplay và chờ quyết định của người chơi trước khi tiếp tục."
        ],
    ),
    (
        "4.8. Giao diện nhập tên và leaderboard",
        [
            "1_4_1_1_dinh_danh_nguoi_choi.drawio",
            "1_4_1_12_leaderboard.drawio",
            "3_2_1_sequence_dang_nhap_nhap_ten.drawio",
            "3_2_3_sequence_game_over_gui_diem_tai_leaderboard.drawio",
        ],
        [
            "Nhóm NameInputPanel và LeaderboardPanel trong Chương 4 được đồng bộ với cả cụm định danh lẫn cụm leaderboard. Nhờ đó, báo cáo thể hiện rõ đây là hai đầu mút của trải nghiệm trực tuyến: một đầu thiết lập danh tính người chơi, đầu còn lại phản hồi thành tích cuối trận.",
            "Sự liên kết với hai sequence riêng giúp phần hình ảnh ở đây được đặt vào đúng bối cảnh nghiệp vụ. NameInputPanel là bước trung gian sau xác thực nếu chưa có Display Name, còn LeaderboardPanel là điểm hiển thị dữ liệu sau khi SubmitScore và GetLeaderboard hoàn tất."
        ],
    ),
    (
        "4.9. Giao diện tạm dừng, cài đặt âm thanh và loading",
        [
            "1_4_1_11_pause_va_ket_thuc_tran.drawio",
            "1_4_1_7_theme_ban_do.drawio",
            "3_4_3_component_backend_services.drawio",
        ],
        [
            "Mục giao diện phụ trợ được đồng bộ với sơ đồ pause/kết thúc trận và sơ đồ đổi theme bản đồ. Theo đó, PauseMenuPanel đại diện cho nhánh dừng gameplay có chủ đích của người chơi, còn LoadingUIManager đại diện cho nhánh chuyển trạng thái kỹ thuật của hệ thống khi đổi theme hoặc reload ngữ cảnh.",
            "Cách trình bày đồng bộ này giúp người đọc nhận ra phần pause, settings và loading không phải các chi tiết trang trí tách rời, mà là các thành phần kiểm soát luồng và cảm giác chuyển tiếp, có ảnh hưởng trực tiếp tới quyền input, nhịp trận và trải nghiệm sử dụng tổng thể."
        ],
    ),
]


def set_run_font(run, size: int = 13, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph: Paragraph, size: int = 13) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_paragraph(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def add_labeled_paragraph(anchor: Paragraph, label: str, text: str) -> Paragraph:
    para = insert_paragraph_after(anchor, style="Normal")
    run1 = para.add_run(f"{label}: ")
    set_run_font(run1, bold=True)
    run2 = para.add_run(text)
    set_run_font(run2)
    return para


def remove_existing_ch4_synced_notes(document: Document) -> None:
    to_remove: list[Paragraph] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Liên hệ sơ đồ nguồn:") or text.startswith("Đồng bộ theo sơ đồ:"):
            to_remove.append(paragraph)
    for paragraph in to_remove:
        remove_paragraph(paragraph)


def sync_chapter4(document: Document) -> None:
    remove_existing_ch4_synced_notes(document)
    for heading, files, paragraphs in reversed(CH4_MAPPINGS):
        anchor = find_paragraph(document, heading)
        anchor = add_labeled_paragraph(anchor, "Liên hệ sơ đồ nguồn", ", ".join(files))
        for text in paragraphs:
            anchor = add_labeled_paragraph(anchor, "Đồng bộ theo sơ đồ", text)


def main() -> None:
    shutil.copyfile(SOURCE_DOC, TARGET_DOC)
    doc = Document(TARGET_DOC)
    sync_chapter4(doc)
    doc.save(TARGET_DOC)
    print(TARGET_DOC)


if __name__ == "__main__":
    main()
