from __future__ import annotations

import math
import re
import sys
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

from generate_drawio_diagrams import DiagramSpec, chapter3_specs, section_14_specs


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_fully_synced.docx"
OUTPUT_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_with_inserted_diagrams.docx"
PNG_DIR = ROOT / "output" / "diagrams" / "png"

FONT_REGULAR = Path(r"C:\Windows\Fonts\times.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\timesbd.ttf")

BG = "#ffffff"
BOX_FILL = "#fff2cc"
BOX_OUTLINE = "#b45f06"
TERM_FILL = "#d9ead3"
TERM_OUTLINE = "#38761d"
NOTE_FILL = "#e8f0fe"
NOTE_OUTLINE = "#3c78d8"
TEXT = "#222222"
EDGE = "#333333"

def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidate = FONT_BOLD if bold else FONT_REGULAR
    if candidate.exists():
        return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def normalize_caption(text: str) -> str:
    text = text.strip().rstrip(".")
    text = re.sub(r"^\d+(?:\.\d+)+\.?\s*", "", text)
    if text.lower().startswith("sơ đồ "):
        return text
    return text if text else "Sơ đồ"


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw_line in text.splitlines() or [""]:
        words = raw_line.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            trial = f"{current} {word}"
            width = draw.textbbox((0, 0), trial, font=font)[2]
            if width <= max_width:
                current = trial
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def draw_multiline_centered(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str = TEXT,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, max(40, x2 - x1 - 16))
    line_heights = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line or "Ag", font=font)
        line_heights.append(bbox[3] - bbox[1])
    total_h = sum(line_heights) + max(0, len(lines) - 1) * 4
    y = y1 + ((y2 - y1 - total_h) / 2)
    for line, line_h in zip(lines, line_heights):
        bbox = draw.textbbox((0, 0), line or " ", font=font)
        line_w = bbox[2] - bbox[0]
        x = x1 + ((x2 - x1 - line_w) / 2)
        draw.text((x, y), line, font=font, fill=fill)
        y += line_h + 4


def node_center(node) -> tuple[float, float]:
    return (node.x + (node.w / 2), node.y + (node.h / 2))


def edge_points(source, target) -> tuple[tuple[int, int], tuple[int, int]]:
    sx, sy = node_center(source)
    tx, ty = node_center(target)
    dx = tx - sx
    dy = ty - sy
    if abs(dx) >= abs(dy):
        start = (source.x + source.w, int(sy)) if dx >= 0 else (source.x, int(sy))
        end = (target.x, int(ty)) if dx >= 0 else (target.x + target.w, int(ty))
    else:
        start = (int(sx), source.y + source.h) if dy >= 0 else (int(sx), source.y)
        end = (int(tx), target.y) if dy >= 0 else (int(tx), target.y + target.h)
    return start, end


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], dashed: bool) -> None:
    if dashed:
        segments = 18
        for i in range(segments):
            if i % 2 == 0:
                p1 = (
                    int(start[0] + (end[0] - start[0]) * (i / segments)),
                    int(start[1] + (end[1] - start[1]) * (i / segments)),
                )
                p2 = (
                    int(start[0] + (end[0] - start[0]) * ((i + 1) / segments)),
                    int(start[1] + (end[1] - start[1]) * ((i + 1) / segments)),
                )
                draw.line([p1, p2], fill=EDGE, width=3)
    else:
        draw.line([start, end], fill=EDGE, width=3)

    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_len = 14
    left = (
        int(end[0] - arrow_len * math.cos(angle - math.pi / 6)),
        int(end[1] - arrow_len * math.sin(angle - math.pi / 6)),
    )
    right = (
        int(end[0] - arrow_len * math.cos(angle + math.pi / 6)),
        int(end[1] - arrow_len * math.sin(angle + math.pi / 6)),
    )
    draw.polygon([end, left, right], fill=EDGE)


def draw_actor(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font: ImageFont.ImageFont) -> None:
    x1, y1, x2, y2 = box
    w = x2 - x1
    h = y2 - y1
    cx = x1 + (w / 2)
    head_r = min(w, h) * 0.12
    head_y = y1 + h * 0.18
    draw.ellipse((cx - head_r, head_y, cx + head_r, head_y + head_r * 2), outline=EDGE, width=3)
    neck_y = head_y + head_r * 2
    hip_y = y1 + h * 0.62
    arm_y = y1 + h * 0.42
    draw.line((cx, neck_y, cx, hip_y), fill=EDGE, width=3)
    draw.line((cx - w * 0.18, arm_y, cx + w * 0.18, arm_y), fill=EDGE, width=3)
    draw.line((cx, hip_y, cx - w * 0.16, y1 + h * 0.88), fill=EDGE, width=3)
    draw.line((cx, hip_y, cx + w * 0.16, y1 + h * 0.88), fill=EDGE, width=3)
    draw_multiline_centered(draw, (x1 + 4, int(y1 + h * 0.84), x2 - 4, y2), text, font)


def render_diagram(spec: DiagramSpec) -> Path:
    PNG_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (spec.width, spec.height), BG)
    draw = ImageDraw.Draw(image)
    font = load_font(28)
    label_font = load_font(22)

    nodes_by_id = {node.id: node for node in spec.nodes}

    for edge in spec.edges:
        source = nodes_by_id[edge.source]
        target = nodes_by_id[edge.target]
        start, end = edge_points(source, target)
        draw_arrow(draw, start, end, edge.dashed)
        if edge.label:
            mx = int((start[0] + end[0]) / 2)
            my = int((start[1] + end[1]) / 2)
            bbox = draw.textbbox((0, 0), edge.label, font=label_font)
            pad = 8
            label_box = (mx - ((bbox[2] - bbox[0]) // 2) - pad, my - 18, mx + ((bbox[2] - bbox[0]) // 2) + pad, my + 18)
            draw.rounded_rectangle(label_box, radius=10, fill="#ffffff", outline="#cccccc")
            draw.text((label_box[0] + pad, label_box[1] + 5), edge.label, font=label_font, fill=TEXT)

    for node in spec.nodes:
        box = (node.x, node.y, node.x + node.w, node.y + node.h)
        if node.kind == "actor":
            draw_actor(draw, box, node.text, font)
            continue
        if node.kind == "terminator":
            radius = max(18, min(node.h // 2, 40))
            draw.rounded_rectangle(box, radius=radius, fill=TERM_FILL, outline=TERM_OUTLINE, width=4)
        elif node.kind == "note":
            draw.rounded_rectangle(box, radius=18, fill=NOTE_FILL, outline=NOTE_OUTLINE, width=4)
        else:
            draw.rounded_rectangle(box, radius=18, fill=BOX_FILL, outline=BOX_OUTLINE, width=4)
        draw_multiline_centered(draw, box, node.text, font)

    output_path = PNG_DIR / f"{Path(spec.filename).stem}.png"
    image.save(output_path, format="PNG")
    return output_path


def insert_paragraph_after(paragraph, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    return new_para


def extract_filename(text: str) -> str | None:
    match = re.search(r"([0-9_]+[a-zA-Z0-9_]*\.drawio)", text)
    return match.group(1) if match else None


def paragraph_caption(text: str) -> str:
    if " - " in text:
        return text.split(" - ", 1)[1].strip()
    return text.strip()


def section_code_from_filename(filename: str) -> str:
    numeric_parts = []
    for part in Path(filename).stem.split("_"):
        if part.isdigit():
            numeric_parts.append(part)
        else:
            break
    return ".".join(numeric_parts)


def desired_width(spec: DiagramSpec) -> Inches:
    ratio = spec.width / spec.height
    if ratio >= 1.8:
        return Inches(6.3)
    if ratio >= 1.4:
        return Inches(6.0)
    return Inches(5.6)


def insert_images(doc: Document, specs_by_filename: dict[str, DiagramSpec], pngs_by_filename: dict[str, Path]) -> int:
    targets = []
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if text.startswith("Tệp sơ đồ nguồn:"):
            filename = extract_filename(text)
            if filename and filename in specs_by_filename:
                targets.append((para, filename, text))

    inserted = 0
    for para, filename, source_line in reversed(targets):
        spec = specs_by_filename[filename]
        image_para = insert_paragraph_after(para)
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = image_para.add_run()
        run.add_picture(str(pngs_by_filename[filename]), width=desired_width(spec))

        section_code = section_code_from_filename(filename)
        caption = f"Hình {section_code}: {normalize_caption(paragraph_caption(source_line))}."
        caption_para = insert_paragraph_after(image_para)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_para.add_run(caption)
        try:
            caption_para.style = "Caption"
        except KeyError:
            pass
        caption_run.italic = True
        inserted += 1
    return inserted


def main() -> int:
    specs = section_14_specs() + chapter3_specs()
    specs_by_filename = {spec.filename: spec for spec in specs}
    pngs_by_filename = {spec.filename: render_diagram(spec) for spec in specs}

    doc = Document(SOURCE_DOC)
    inserted = insert_images(doc, specs_by_filename, pngs_by_filename)
    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOC)

    sys.stdout.reconfigure(encoding="utf-8")
    print(f"Rendered PNGs: {len(pngs_by_filename)}")
    print(f"Inserted diagrams: {inserted}")
    print(f"Output doc: {OUTPUT_DOC}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
