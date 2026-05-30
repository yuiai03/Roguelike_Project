from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_only_14_updated.docx"
TARGET_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_14_rewritten.docx"


SECTION_141_CONTENT: dict[str, list[tuple[str, str]]] = {
    "1.4.1.1. Chức năng định danh người chơi": [
        ("Mô tả chức năng", "Chức năng định danh người chơi bảo đảm mỗi phiên chơi đều gắn với một tài khoản PlayFab hợp lệ trước khi tham gia vòng lặp gameplay và trước khi sử dụng bảng xếp hạng."),
        ("Luồng thực hiện", "1) Khi vào game, PlayFabLeaderboardManager gọi LoginWithCustomID để đăng nhập. 2) Hệ thống lấy hồ sơ người chơi bằng GetPlayerProfile và kiểm tra Display Name hiện có. 3) Nếu người chơi chưa có tên hiển thị, NameInputPanel được mở để nhập tên mới. 4) Khi người chơi xác nhận, hệ thống kiểm tra tính hợp lệ rồi cập nhật lại hồ sơ và đồng bộ CurrentDisplayName cùng CurrentPlayFabId cho toàn bộ phiên chơi."),
        ("Thông số cụ thể", "Custom ID được lấy từ dữ liệu cục bộ của thiết bị; Display Name hợp lệ trong khoảng 3 đến 25 ký tự; dữ liệu đồng bộ chính gồm CurrentDisplayName và CurrentPlayFabId."),
        ("Dữ liệu vào/ra", "Đầu vào gồm Custom ID, hồ sơ người chơi từ PlayFab và chuỗi tên nhập từ UI. Đầu ra là một danh tính đã xác thực, sẵn sàng dùng cho gửi điểm, truy xuất thứ hạng và hiển thị tên trong trận."),
        ("Kết quả đầu ra", "Người chơi có tài khoản hợp lệ và tên hiển thị rõ ràng trước khi bước vào trận đấu."),
    ],
    "1.4.1.2. Chức năng bắt đầu trận đấu": [
        ("Mô tả chức năng", "Chức năng này chuyển game từ trạng thái chờ trong scene sang trạng thái chiến đấu và khởi tạo wave đầu tiên."),
        ("Luồng thực hiện", "1) Người chơi tiến vào vùng tương tác của ChallengePostNPC. 2) InteractPanel hiển thị gợi ý thao tác. 3) Sau khi người chơi nhấn phím tương tác, ChallengePanel được mở để hiển thị hướng dẫn và nút bắt đầu. 4) Khi người chơi xác nhận Start Game, sự kiện onGameStart được phát, HUD được bind dữ liệu và WaveSpawner.StartNextWave được gọi để khởi tạo trận đấu."),
        ("Thông số cụ thể", "Dữ liệu điều kiện gồm playerInRange, phím tương tác và trạng thái hiển thị của ChallengePanel; đầu ra điều phối gồm currentWave, trạng thái HUD và trạng thái hoạt động của NPC tương tác."),
        ("Dữ liệu vào/ra", "Đầu vào là thao tác của player với NPC và lệnh Start Game trên UI. Đầu ra là tín hiệu bắt đầu trận, HUD được bật và wave đầu tiên đi vào trạng thái chuẩn bị."),
        ("Kết quả đầu ra", "Hệ thống rời pha chờ và đi vào vòng lặp gameplay chính."),
    ],
    "1.4.1.3. Chức năng điều khiển nhân vật": [
        ("Mô tả chức năng", "Chức năng điều khiển nhân vật quản lý toàn bộ thao tác di chuyển, định hướng, gravity và dash của player trong môi trường 3D."),
        ("Luồng thực hiện", "1) InputSystem_Actions gửi dữ liệu Move và thao tác chuột phải tới PlayerController. 2) Từ hướng nhìn của Camera.main, PlayerController quy đổi input thành hướng di chuyển trong không gian chơi. 3) CharacterController thực hiện Move, xử lý gravity và kiểm tra trạng thái chạm đất. 4) Khi người chơi kích hoạt dash, hệ thống chuyển sang Dash State, áp dụng tốc độ dash trong thời gian ngắn rồi quay về trạng thái di chuyển thường; đồng thời PlayerAnimationController cập nhật hoạt ảnh tương ứng."),
        ("Thông số cụ thể", "moveSpeed = 5, dashSpeed = 15, dashCooldown = 1; ngoài ra hệ thống còn theo dõi grounded state, dash trigger, hướng camera và cờ isInputActive."),
        ("Dữ liệu vào/ra", "Đầu vào là moveInput, hướng camera, trạng thái chạm đất và tín hiệu dash. Đầu ra là vị trí player, hướng quay model và trạng thái hoạt ảnh Idle, Run hoặc Dash."),
        ("Kết quả đầu ra", "Player di chuyển mượt, đổi hướng đúng theo camera và không giữ input sai ngữ cảnh khi UI khóa điều khiển."),
    ],
    "1.4.1.4. Chức năng chiến đấu của người chơi": [
        ("Mô tả chức năng", "Hệ thống chiến đấu của player được tổ chức theo hướng tự động tìm mục tiêu và tự động bắn, còn chiều sâu chiến thuật nằm ở việc di chuyển và lựa chọn buff."),
        ("Luồng thực hiện", "1) PlayerAttack liên tục đếm thời gian giữa các lần tấn công. 2) Khi đủ điều kiện bắn, hệ thống gọi FindNearestEnemy để tìm mục tiêu gần nhất trong phạm vi tấn công. 3) Sau khi có mục tiêu, ObjectPool sinh PlayerProjectile và nạp các giá trị khởi tạo như damage, speed và lifetime. 4) Nếu player đang có buff bổ sung như multishot hoặc AoE, projectile sẽ được gắn thêm dữ liệu mở rộng trước khi gây sát thương lên enemy thông qua IDamageable."),
        ("Thông số cụ thể", "attackRange = 20; các biến cấu hình chính gồm attackCooldown, attackDamage, projectileSpeed, projectileLifetime, multiShotCount, multiShotAngle, aoeRadius và aoeAtkMultiplier."),
        ("Dữ liệu vào/ra", "Đầu vào là dữ liệu chiến đấu trong PlayerData và danh sách enemy đang có trong phạm vi kiểm tra. Đầu ra là projectile đang hoạt động và lượng sát thương hợp lệ được truyền sang enemy."),
        ("Kết quả đầu ra", "Player luôn duy trì khả năng gây sát thương theo đúng trạng thái buff hiện tại mà không cần thao tác bắn thủ công."),
    ],
    "1.4.1.5. Chức năng quản lý enemy": [
        ("Mô tả chức năng", "Chức năng quản lý enemy xây dựng tập đối tượng địch với nhiều vai trò khác nhau và duy trì vòng đời từ spawn, truy đuổi hoặc tấn công cho tới khi chết và trả về pool."),
        ("Luồng thực hiện", "1) WaveSpawner chọn đúng loại enemy cần sinh dựa trên cấu hình wave. 2) EnemyData nhận dữ liệu chỉ số từ EnemyConfig, sau đó truyền cho đối tượng Enemy và các lớp kế thừa như MeleeEnemy, RangedEnemy, FlyEnemy và BossEnemy. 3) Trong runtime, enemy di chuyển và xử lý AI bằng CharacterController, tiếp cận người chơi hoặc giữ cự ly bắn tùy loại. 4) Khi máu về 0, sự kiện OnDeath được kích hoạt để cộng EXP, cập nhật lại trạng thái wave và dọn đối tượng trước khi despawn."),
        ("Thông số cụ thể", "Các thông số chính gồm maxHealth, moveSpeed, attackRange, contactDamage, projectileDamage, cooldown, expValue và randomVariation; dữ liệu này đi qua EnemyConfig và EnemyData trước khi áp vào runtime."),
        ("Dữ liệu vào/ra", "Đầu vào là loại enemy và cấu hình chỉ số tương ứng. Đầu ra là một đối tượng enemy hoạt động đúng vai trò cận chiến, đánh xa, bay hoặc boss và phát OnDeath khi bị hạ."),
        ("Kết quả đầu ra", "Hệ thống enemy tạo đủ áp lực chiến đấu mà vẫn giữ được luồng dữ liệu rõ ràng cho EXP, wave và despawn."),
    ],
    "1.4.1.6. Chức năng quản lý wave và độ khó": [
        ("Mô tả chức năng", "WaveSpawner là bộ điều phối nhịp trận đấu, chịu trách nhiệm chọn dữ liệu wave, sinh enemy, xác định boss wave và kiểm soát mức tăng độ khó theo thời gian."),
        ("Luồng thực hiện", "1) Khi StartNextWave được gọi, currentWave tăng lên và hệ thống lấy cấu hình tương ứng từ WaveConfig. 2) Sau thời gian chuẩn bị, WaveSpawner bắt đầu sinh quái theo các enemyGroups của wave thường hoặc sinh boss theo bossPoolTypes nếu là boss wave. 3) Trong quá trình này, activeEnemies được cập nhật liên tục thông qua OnEnemyCountChanged. 4) Khi số enemy còn lại về 0, CompleteWave được kích hoạt để chuyển sang wave kế tiếp; nếu vượt quá số wave đã cấu hình, hệ thống chuyển sang endless wave theo cơ chế tái sử dụng base wave."),
        ("Thông số cụ thể", "autoScale = 1.1; dữ liệu chính gồm waves, enemyGroups, spawnPosition, spreadRadius, spawnDelay, currentWave, bossPoolTypes, activeEnemies và trạng thái session."),
        ("Dữ liệu vào/ra", "Đầu vào là cấu hình wave và các nhóm enemy tương ứng. Đầu ra là danh sách enemy đang hoạt động, trạng thái wave hiện tại, tín hiệu complete wave và nhịp tăng độ khó của trận."),
        ("Kết quả đầu ra", "Trận đấu được tổ chức theo từng wave rõ ràng, có boss wave và có thể tiếp tục endless khi vượt mốc cấu hình ban đầu."),
    ],
    "1.4.1.7. Chức năng đổi theme bản đồ": [
        ("Mô tả chức năng", "Hệ thống đổi theme bản đồ làm mới không gian chiến đấu theo tiến trình wave mà không làm gián đoạn dữ liệu gameplay hiện hành."),
        ("Luồng thực hiện", "1) Sau khi một wave kết thúc, WaveSpawner xác định upcomingWave và chuyển dữ liệu này cho MapThemeManager. 2) MapThemeManager sử dụng ResolveThemeIndexForWave để xác định theme mới tương ứng. 3) Trước khi áp dụng thay đổi, LoadingUIManager thực hiện black transition để che chuyển cảnh. 4) Sau đó hệ thống thay vật liệu mặt đất, tường và effectRoot của scene; khi hoàn tất, OnThemeTransitionCompleted được phát để tiếp tục gameplay."),
        ("Thông số cụ thể", "Dữ liệu điều khiển gồm currentWave, upcomingWave, CurrentThemeIndex, danh sách theme, groundMaterial, wallMaterial và effectRoot."),
        ("Dữ liệu vào/ra", "Đầu vào là trạng thái wave và tập theme tương ứng. Đầu ra là theme mới của scene cùng môi trường hình ảnh đã được đồng bộ với mốc wave."),
        ("Kết quả đầu ra", "Bản đồ thay đổi theo tiến trình trận mà không làm sai trạng thái input hoặc luồng wave đang chạy."),
    ],
    "1.4.1.8. Chức năng kinh nghiệm và lên cấp": [
        ("Mô tả chức năng", "Hệ thống EXP và level tạo trục tăng tiến sức mạnh của người chơi và quyết định thời điểm kích hoạt cơ chế chọn buff."),
        ("Luồng thực hiện", "1) Khi enemy chết, sự kiện Enemy OnDeath được phát và ExpDropper đọc expValue để cộng EXP cho PlayerLevelSystem thông qua AddExp. 2) Sau đó HUD được cập nhật qua OnExpChanged. 3) Nếu currentExp đạt hoặc vượt ngưỡng expToNextLevel, hệ thống thực hiện LevelUp, tăng currentLevel và phát OnLevelUp. 4) Song song với đó, totalExpGained được tích lũy để dùng làm nguồn dữ liệu gửi HighScore."),
        ("Thông số cụ thể", "Các biến chính gồm expValue, currentExp, expToNextLevel, currentLevel, expScalingFactor và totalExpGained."),
        ("Dữ liệu vào/ra", "Đầu vào là dữ liệu EXP rơi ra từ enemy và trạng thái hiện tại của PlayerLevelSystem. Đầu ra là level mới, thanh EXP mới và tín hiệu kích hoạt chọn buff."),
        ("Kết quả đầu ra", "Người chơi tăng cấp đúng nhịp chiến đấu và toàn bộ HUD cùng các hệ tiến trình nhận chung một nguồn dữ liệu EXP."),
    ],
    "1.4.1.9. Chức năng chọn buff và tăng sức mạnh": [
        ("Mô tả chức năng", "Chức năng buff tạo khả năng cá nhân hóa từng lượt chơi, cho phép người chơi lựa chọn hướng tăng chỉ số hoặc mở rộng kỹ năng."),
        ("Luồng thực hiện", "1) Khi PlayerLevelSystem phát OnLevelUp, BuffCardManager thực hiện GetRandomCards để chọn các card phù hợp từ tập card còn khả dụng. 2) Hệ thống áp dụng logic rarity và luckBonus để quyết định xác suất xuất hiện của từng card. 3) CardSelectionPanel hiển thị các lựa chọn cho người chơi. 4) Khi một card được chọn, ApplyCard sẽ cập nhật trực tiếp PlayerData, PlayerHealth hoặc các hệ thống chiến đấu mở rộng như Spirit, OrbitingBall và các combat modifiers."),
        ("Thông số cụ thể", "cardsPerSelection = 3; ngoài ra hệ thống còn theo dõi rarity, luckBonus, maxLevel, cardLevels và trạng thái hiện tại của PlayerData."),
        ("Dữ liệu vào/ra", "Đầu vào là danh sách card hợp lệ và dữ liệu tiến trình hiện tại của người chơi. Đầu ra là chỉ số hoặc kỹ năng mới được áp dụng ngay trong runtime."),
        ("Kết quả đầu ra", "Mỗi lần lên cấp, người chơi nhận được đúng một quyết định tăng sức mạnh có tác động trực tiếp tới chiến đấu."),
    ],
    "1.4.1.10. Chức năng giao diện trong trận": [
        ("Mô tả chức năng", "Lớp giao diện trong trận vừa phản ánh trạng thái gameplay theo thời gian thực, vừa điều phối quyền nhận input giữa các panel."),
        ("Luồng thực hiện", "1) GameUI đóng vai trò trung tâm, giữ tham chiếu đến PlayerStatsPanel, InteractPanel, CardSelectionPanel, NameInputPanel, LeaderboardPanel và PauseMenuPanel. 2) Trong quá trình chơi, PlayerStatsPanel cập nhật các chỉ số HP, EXP, level và wave. 3) Những panel còn lại được bật hoặc tắt theo từng ngữ cảnh như tương tác NPC, nhập tên, chọn buff, xem bảng xếp hạng hoặc tạm dừng trận. 4) Khi có panel ưu tiên cao được mở, hệ thống SetInputActive sẽ khóa input gameplay để tránh xung đột thao tác."),
        ("Thông số cụ thể", "Dữ liệu hiển thị chính gồm HP, EXP, Lv, Wave; trạng thái điều phối chính gồm panel priority và cờ Input ownership / SetInputActive."),
        ("Dữ liệu vào/ra", "Đầu vào là các sự kiện và chỉ số runtime từ PlayerHealth, PlayerLevelSystem, WaveSpawner và PlayFabLeaderboardManager. Đầu ra là HUD, prompt tương tác và các panel phụ được mở đúng thời điểm."),
        ("Kết quả đầu ra", "Giao diện luôn bám sát trạng thái game và không để người chơi thao tác sai ngữ cảnh."),
    ],
    "1.4.1.11. Chức năng tạm dừng và kết thúc trận": [
        ("Mô tả chức năng", "Chức năng này quản lý hai trạng thái ngắt vòng lặp gameplay là tạm dừng chủ động và kết thúc trận khi người chơi thất bại."),
        ("Luồng thực hiện", "1) Khi người chơi nhấn ESC, PauseMenuPanel kiểm tra điều kiện mở menu. 2) Nếu hợp lệ, Time.timeScale được đặt về 0, input gameplay bị khóa và menu pause hiện ra với các lựa chọn Resume, Settings, Leaderboard và Quit. 3) Ở nhánh kết thúc trận, khi PlayerHealth nhận thấy hp <= 0, hệ thống kích hoạt chuỗi dọn dẹp gồm clear wave, despawn pool và ngắt các đối tượng gameplay còn lại. 4) Sau đó màn hình leaderboard sau khi chết được hiển thị để người chơi xem kết quả."),
        ("Thông số cụ thể", "ESC là tín hiệu mở pause; dữ liệu trạng thái gồm Time.timeScale, Player input off, hp <= 0, currentView và currentRunScore."),
        ("Dữ liệu vào/ra", "Đầu vào là thao tác mở pause hoặc trạng thái chết của player. Đầu ra là trạng thái dừng game, giao diện pause hoặc màn hình tổng kết sau khi trận kết thúc."),
        ("Kết quả đầu ra", "Trận đấu có thể được tạm dừng an toàn hoặc kết thúc đúng quy trình mà không để lại trạng thái dở dang."),
    ],
    "1.4.1.12. Chức năng bảng xếp hạng": [
        ("Mô tả chức năng", "Bảng xếp hạng là lớp giao tiếp giữa client và PlayFab để gửi điểm, lấy dữ liệu xếp hạng và hiển thị vị trí của người chơi."),
        ("Luồng thực hiện", "1) Sau khi kết thúc lượt chơi, PlayerHealth gọi GetCurrentRunScore để lấy điểm của phiên chơi hiện tại. 2) PlayFabLeaderboardManager dùng dữ liệu này để SubmitScore lên PlayFab với statistic HighScore. 3) Khi người chơi mở bảng xếp hạng, hệ thống gọi GetLeaderboardData để lấy top tổng quát và GetPlayerLeaderboardData để lấy vùng xếp hạng quanh tài khoản hiện tại. 4) Cuối cùng, LeaderboardPanel sinh các entry giao diện và đánh dấu đúng dòng của người chơi."),
        ("Thông số cụ thể", "Statistic dùng để gửi điểm là HighScore; MaxResultsCount = 100; dữ liệu đồng bộ chính gồm CurrentDisplayName, CurrentPlayFabId và điểm của lượt chơi hiện tại."),
        ("Dữ liệu vào/ra", "Đầu vào là điểm cuối trận, danh tính PlayFab và yêu cầu tải bảng xếp hạng. Đầu ra là top leaderboard, hạng cá nhân và danh sách entry hiển thị trên UI."),
        ("Kết quả đầu ra", "Người chơi xem được thành tích của mình và vị trí tương đối so với những người chơi khác."),
    ],
}


SECTION_142_CONTENT = [
    ("Hiệu năng", "Các đối tượng được tạo và hủy với tần suất cao như enemy, projectile, effect, damage text và vật phẩm cần ưu tiên tái sử dụng thông qua ObjectPool để giữ ổn định tốc độ khung hình."),
    ("Tổ chức mã nguồn", "Các script cần được chia tách theo nhóm Player, Enemy, Wave, UI, Buff, Systems và Utils để đường đi dữ liệu rõ ràng và thuận lợi cho bảo trì."),
    ("Khả năng mở rộng", "Enemy mới, buff mới, wave mới hoặc theme mới nên được bổ sung chủ yếu bằng prefab và ScriptableObject, hạn chế sửa trực tiếp vào logic lõi nếu không thật cần thiết."),
    ("Độ ổn định luồng chơi", "Các trạng thái nhập tên, pause, chọn buff, leaderboard và đổi theme phải khóa hoặc trả lại input đúng lúc, không gây xung đột với trạng thái wave đang chạy."),
]


SECTION_143_CONTENT = [
    ("Mô hình tổng quát", "Mô hình bài toán của game được tổ chức thành bốn khối chính. Khối thứ nhất là định danh và chuẩn bị vào trận, bao gồm PlayFabLeaderboardManager, NameInputPanel, ChallengePostNPC và ChallengePanel. Khối thứ hai là gameplay runtime, bao gồm PlayerController, PlayerAttack, Enemy, WaveSpawner và MapThemeManager. Khối thứ ba là tiến trình sức mạnh, bao gồm ExpDropper, PlayerLevelSystem, BuffCardManager và CardSelectionPanel. Khối thứ tư là giao diện và tổng kết, bao gồm GameUI, PauseMenuPanel, LeaderboardPanel và dịch vụ PlayFab."),
    ("Quan hệ giữa các khối", "Các khối trên liên kết với nhau theo chuỗi dữ liệu rõ ràng: danh tính người chơi được khởi tạo trước; sau đó hệ thống cho phép mở trận; trong trận đấu, player điều khiển nhân vật, chiến đấu với enemy và nhận EXP; khi lên cấp, hệ thống sinh lựa chọn buff; khi người chơi chết hoặc kết thúc lượt chơi, điểm được gửi lên PlayFab và kết quả được phản hồi qua leaderboard."),
    ("Ý nghĩa trình bày", "Thay vì dùng một sơ đồ tổng hợp quá lớn, mục 1.4 sử dụng bộ 12 sơ đồ thành phần từ 1_4_1_1 đến 1_4_1_12. Cách trình bày này giúp người đọc dễ theo dõi từng nhóm chức năng, dễ đối chiếu với code và dễ chỉnh sửa từng phần mà không ảnh hưởng toàn bộ sơ đồ."),
]


SECTION_144_CONTENT = [
    ("Nguyên tắc mô tả", "Luồng hoạt động của hệ thống được trình bày theo đúng thứ tự xảy ra trong một lượt chơi, bắt đầu từ định danh người chơi và kết thúc ở bước gửi điểm lên leaderboard."),
    ("Chuỗi hoạt động chính", "Hệ thống bắt đầu từ bước đăng nhập PlayFab, lấy hồ sơ hiện có và yêu cầu nhập Display Name nếu người chơi chưa có tên hợp lệ. Sau đó người chơi tương tác với ChallengePostNPC để mở ChallengePanel và bắt đầu trận đấu. WaveSpawner khởi tạo wave đầu tiên, PlayerController nhận input di chuyển và PlayerAttack tự động xử lý tìm mục tiêu, sinh projectile và gây sát thương. Enemy bị tiêu diệt sẽ phát OnDeath để cộng EXP, kích hoạt lên cấp và mở CardSelectionPanel khi đủ điều kiện."),
    ("Luồng tăng tiến và ngắt trạng thái", "Sau mỗi lần lên cấp, BuffCardManager sinh 3 lựa chọn buff và áp dụng đúng card người chơi chọn vào PlayerData hoặc các hệ chiến đấu mở rộng. Khi wave hoàn tất, hệ thống có thể chuyển boss wave hoặc đổi theme bản đồ. Trong các mốc chuyển trạng thái như đổi theme, pause, chọn buff hoặc mở leaderboard, GameUI chịu trách nhiệm điều phối panel và quyền nhận input để không làm hỏng luồng gameplay chính."),
    ("Kết thúc trận", "Khi người chơi chết, hệ thống dọn các đối tượng còn lại, lấy điểm của lượt chơi, gửi HighScore lên PlayFab và hiển thị LeaderboardPanel để tổng kết kết quả. Vì vậy, toàn bộ hoạt động của hệ thống có thể tóm tắt theo chuỗi: định danh người chơi, bắt đầu trận, điều khiển và chiến đấu, quản lý enemy và wave, cộng EXP và chọn buff, tạm dừng hoặc kết thúc trận, sau cùng là gửi điểm và hiển thị bảng xếp hạng."),
]


def set_run_font(run, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        set_run_font(run)


def find_paragraph(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def is_heading_text(text: str) -> bool:
    text = text.strip()
    return bool(re.match(r"^1\.4(\.|$)", text)) or text.startswith("2.") or text.startswith("CHƯƠNG 2")


def remove_detail_block(document: Document, start_heading: str) -> Paragraph:
    paragraphs = document.paragraphs
    heading_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == start_heading)
    caption_idx = None
    first_detail_idx = None
    end_idx = len(paragraphs)

    for i in range(heading_idx + 1, len(paragraphs)):
        txt = paragraphs[i].text.strip()
        if first_detail_idx is None and txt.startswith("Mô tả chức năng:"):
            first_detail_idx = i
        if caption_idx is None and txt.startswith("Hình 1.4.1."):
            caption_idx = i
        if i > heading_idx + 1 and is_heading_text(txt):
            end_idx = i
            break

    if caption_idx is None:
        raise ValueError(f"Caption not found after heading: {start_heading}")
    if first_detail_idx is None:
        raise ValueError(f"Detail block not found after heading: {start_heading}")

    to_remove = list(document.paragraphs[first_detail_idx:end_idx])
    for paragraph in to_remove:
        remove_paragraph(paragraph)
    return document.paragraphs[caption_idx]


def add_labeled_paragraph(anchor: Paragraph, label: str, text: str) -> Paragraph:
    para = insert_paragraph_after(anchor, style="Normal")
    run1 = para.add_run(f"{label}: ")
    set_run_font(run1, bold=True)
    run2 = para.add_run(text)
    set_run_font(run2)
    return para


def replace_subsection_blocks(document: Document) -> None:
    for heading, items in SECTION_141_CONTENT.items():
        anchor = remove_detail_block(document, heading)
        for label, text in items:
            anchor = add_labeled_paragraph(anchor, label, text)


def replace_simple_section(document: Document, heading: str, items: list[tuple[str, str]], next_heading: str) -> None:
    paragraphs = document.paragraphs
    start_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == heading)
    end_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == next_heading and i > start_idx)
    start_para = document.paragraphs[start_idx]
    to_remove = list(document.paragraphs[start_idx + 1:end_idx])
    for paragraph in to_remove:
        remove_paragraph(paragraph)

    anchor = start_para
    for label, text in items:
        anchor = add_labeled_paragraph(anchor, label, text)


def main() -> None:
    shutil.copyfile(SOURCE_DOC, TARGET_DOC)
    document = Document(TARGET_DOC)

    replace_subsection_blocks(document)
    replace_simple_section(document, "1.4.2. Yêu cầu phi chức năng", SECTION_142_CONTENT, "1.4.3. Mô hình bài toán tổng quát")
    replace_simple_section(document, "1.4.3. Mô hình bài toán tổng quát", SECTION_143_CONTENT, "1.4.4. Mô tả hoạt động của hệ thống")
    replace_simple_section(document, "1.4.4. Mô tả hoạt động của hệ thống", SECTION_144_CONTENT, "CHƯƠNG 2 \nMỘT SỐ KIẾN THỨC CƠ BẢN THỰC HIỆN ĐỀ TÀI")

    document.save(TARGET_DOC)
    print(TARGET_DOC)


if __name__ == "__main__":
    main()
