from __future__ import annotations

import re
import shutil
from pathlib import Path
import unicodedata

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = ROOT / "tmp" / "docs" / "BaoCao_DATN_2_caption_fix_snapshot.docx"
OUTPUT_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_caption_catalog_fixed.docx"

BLACK = RGBColor(0, 0, 0)


def ascii_upper(text: str) -> str:
    normalized = unicodedata.normalize("NFD", text)
    stripped = "".join(ch for ch in normalized if unicodedata.category(ch) != "Mn")
    return stripped.upper()


def set_run_font(run, size: int = 12) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK


def set_paragraph_font(paragraph, size: int = 12) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size)


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def replace_paragraph_text(paragraph, text: str, size: int = 12) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size)


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def paragraphs_between(document: Document, start_ascii: str, end_predicate):
    collecting = False
    items = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        ascii_text = ascii_upper(text)
        if ascii_text == start_ascii:
            collecting = True
            continue
        if collecting and end_predicate(ascii_text):
            break
        if collecting:
            items.append(paragraph)
    return items


def normalize_caption(text: str, chapter_number: int, figure_index: int) -> str:
    prefix = text.strip().split(maxsplit=1)[0] if text.strip() else "Hình"
    body = re.sub(r"^\S+\s+\d+(?:[-.]\d+)+\s*:\s*", "", text.strip())
    body = re.sub(r"^\S+\s+\d+(?:[-.]\d+)*\s*:\s*", "", body)
    if not body:
        body = "Minh họa"
    return f"{prefix} {chapter_number}.{figure_index}: {body}"


def is_figure_caption(paragraph) -> bool:
    text = paragraph.text.strip()
    return bool(re.match(r"^Hình\s+\d", text))


def fix_document() -> tuple[Path, int]:
    shutil.copyfile(SOURCE_DOC, OUTPUT_DOC)
    doc = Document(OUTPUT_DOC)

    chapter_number = 0
    in_figure_catalog = False
    figure_counts: dict[int, int] = {}
    figure_entries: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        ascii_text = ascii_upper(text)

        if ascii_text == "DANH MUC HINH VE":
            in_figure_catalog = True
            continue

        if in_figure_catalog and ascii_text.startswith("CHUONG 1"):
            in_figure_catalog = False

        if in_figure_catalog:
            continue

        chapter_match = re.match(r"^CHƯƠNG\s+(\d+)", text)
        if chapter_match:
            chapter_number = int(chapter_match.group(1))
            figure_counts.setdefault(chapter_number, 0)
            continue

        if not is_figure_caption(paragraph) or chapter_number <= 0:
            continue

        figure_counts[chapter_number] += 1
        new_caption = normalize_caption(text, chapter_number, figure_counts[chapter_number])
        replace_paragraph_text(paragraph, new_caption, size=12)
        try:
            paragraph.style = "Caption"
        except KeyError:
            pass
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        figure_entries.append(new_caption)

    lof_items = paragraphs_between(
        doc,
        start_ascii="DANH MUC HINH VE",
        end_predicate=lambda s: s.startswith("CHUONG 1"),
    )
    anchor = None
    for paragraph in lof_items:
        remove_paragraph(paragraph)

    for paragraph in doc.paragraphs:
        if ascii_upper(paragraph.text.strip()) == "DANH MUC HINH VE":
            anchor = paragraph
            break
    if anchor is None:
        raise ValueError("Không tìm thấy DANH MỤC HÌNH VẼ")

    current = anchor
    for entry in figure_entries:
        current = insert_paragraph_after(current, entry, style="Normal")
        set_paragraph_font(current, size=12)
        current.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(OUTPUT_DOC)
    return OUTPUT_DOC, len(figure_entries)


if __name__ == "__main__":
    output, total = fix_document()
    print(f"output={output}")
    print(f"figures={total}")
