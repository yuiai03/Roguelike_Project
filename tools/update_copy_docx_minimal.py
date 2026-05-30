from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
INPUT_DOC = ROOT / "tmp" / "docs" / "BaoCao_DATN_2_copy_live_snapshot.docx"
SOURCE_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_git_ch4_conclusion_updated.docx"
OUTPUT_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_copy_minimal_updated.docx"
BACKUP_DOC = ROOT / "tmp" / "docs" / "BaoCao_DATN_2_copy_before_minimal_update.docx"

BLACK = RGBColor(0, 0, 0)

CH4_SECTIONS = [
    "4.1. Toàn cảnh gameplay trong scene Game",
    "4.2. Nhân vật chính của trò chơi",
    "4.3. Danh sách enemy cơ bản trong trò chơi",
    "4.4. Nhóm enemy nguyên tố LawaChurl",
    "4.5. Các theme map và chuyển đổi môi trường",
    "4.6. Giao diện HUD và hỗ trợ gameplay",
    "4.7. Giao diện challenge và chọn buff",
    "4.8. Giao diện nhập tên và leaderboard",
    "4.9. Giao diện tạm dừng, cài đặt âm thanh và loading",
]

TABLE_FRONT_PAGES = ["11", "29", "30", "36", "37", "47", "49", "51", "53", "59"]

REFERENCE_ITEMS = [
    '[1] Microsoft, "C# language documentation." [Online]. Available: https://learn.microsoft.com/dotnet/csharp/.',
    '[2] Unity Technologies, "Unity Manual." [Online]. Available: https://docs.unity3d.com/Manual/index.html.',
    '[3] Microsoft, "PlayFab documentation." [Online]. Available: https://learn.microsoft.com/gaming/PlayFab/.',
    '[4] Unity Technologies, "Input System package documentation." [Online]. Available: https://docs.unity3d.com/Packages/com.unity.inputsystem.',
    '[5] Unity Technologies, "Character Controller." [Online]. Available: https://docs.unity3d.com/Manual/class-CharacterController.',
    '[6] Unity Technologies, "ScriptableObject." [Online]. Available: https://docs.unity3d.com/ScriptReference/ScriptableObject.html.',
    '[7] HoYoverse, "Genshin Impact Official channel." [Online]. Available: https://www.youtube.com/@GenshinImpact.',
]

GIT_TABLE_DATA = [
    ["Bước", "Cách nhóm thực hiện", "Ý nghĩa quản lý project"],
    ["Cập nhật local", "Pull hoặc lấy bản mới nhất trước khi sửa code.", "Giảm nguy cơ làm việc trên phiên bản cũ."],
    [
        "Phát triển chức năng",
        "Chỉnh script C#, prefab, ScriptableObject hoặc tài liệu theo từng nhiệm vụ.",
        "Tách thay đổi theo phạm vi rõ ràng, dễ kiểm soát.",
    ],
    [
        "Commit",
        "Lưu mốc sau mỗi nhóm chức năng như spawn tool, wave config, effect hoặc map delay.",
        "Giữ lịch sử sửa đổi để xem lại và khôi phục khi cần.",
    ],
    ["Đối chiếu lịch sử", "Dùng log để so sánh các mốc chỉnh sửa gần nhau.", "Hỗ trợ truy vết lỗi và đánh giá tiến độ thực hiện."],
]

GITHUB_TABLE_DATA = [
    ["Thành phần", "Dữ liệu thực tế của dự án", "Vai trò sử dụng"],
    ["Git local", "Nhánh chính hiện tại là main.", "Quản lý chỉnh sửa trên máy, commit và kiểm tra lịch sử trước khi đồng bộ."],
    [
        "GitHub remote",
        "origin -> https://github.com/yuiai03/Roguelike_Project.git",
        "Làm kho trung tâm để chia sẻ mã nguồn và đồng bộ giữa các máy.",
    ],
    [
        "Commit history",
        "Ví dụ: update effect, update spawn tool, update hiển thị wave và tool config.",
        "Phản ánh các mốc triển khai thật của gameplay và công cụ nội bộ.",
    ],
    ["Contributor", "Local history hiện thấy lseanl03 và YuiAI.", "Cho phép theo dõi người tham gia và mức độ đóng góp vào project."],
]


def has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.xpath(".//*[local-name()='drawing']"))


def is_heading(paragraph: Paragraph) -> bool:
    return bool(paragraph.style and paragraph.style.name.startswith("Heading"))


def is_caption(paragraph: Paragraph) -> bool:
    style_name = paragraph.style.name if paragraph.style else ""
    text = paragraph.text.strip()
    return style_name == "Caption" or text.startswith("Hình ") or text.startswith("Bảng ")


def set_run_font(run, *, size: int = 13, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph: Paragraph, *, size: int = 13, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def replace_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    size: int = 13,
    bold: bool | None = None,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    paragraph.alignment = alignment


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_paragraph(document: Document, prefix: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Missing paragraph prefix: {prefix}")


def paragraph_list(document: Document) -> list[Paragraph]:
    return list(document.paragraphs)


def paragraph_index(document: Document, paragraph: Paragraph) -> int:
    for index, item in enumerate(document.paragraphs):
        if item._p is paragraph._p:
            return index
    raise ValueError("Paragraph not found in document.")


def paragraphs_between(document: Document, start: Paragraph, end: Paragraph) -> list[Paragraph]:
    paragraphs = paragraph_list(document)
    start_idx = paragraph_index(document, start)
    end_idx = paragraph_index(document, end)
    return paragraphs[start_idx + 1 : end_idx]


def insert_paragraph_before(anchor: Paragraph, template: Paragraph | None = None, style: str | None = None, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    if template is not None and template._p.pPr is not None:
        new_p.append(deepcopy(template._p.pPr))
    anchor._p.addprevious(new_p)
    para = Paragraph(new_p, anchor._parent)
    if style is not None:
        para.style = style
    if text:
        para.add_run(text)
    return para


def insert_table_before(anchor: Paragraph, source_table: Table) -> Table:
    new_tbl = deepcopy(source_table._tbl)
    anchor._p.addprevious(new_tbl)
    return Table(new_tbl, anchor._parent)


def insert_paragraph_after(anchor: Paragraph, template: Paragraph | None = None, style: str | None = None, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    if template is not None and template._p.pPr is not None:
        new_p.append(deepcopy(template._p.pPr))
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    if style is not None:
        para.style = style
    if text:
        para.add_run(text)
    return para


def transform_report_text(text: str) -> str:
    text = re.sub(r"\b(Hình|Bảng)\s+(\d+)\.(\d+)", r"\1 \2-\3", text)
    text = text.replace("GetPlayerLeaderboardDt()", "GetPlayerLeaderboardData()")
    text = text.replace("EnemyGroup   -> ObjectPool", "EnemyGroup -> ObjectPool")
    return text


def extract_source_texts(source_doc: Document, heading: str, next_heading: str) -> list[str]:
    start = find_paragraph(source_doc, heading)
    end = find_paragraph(source_doc, next_heading)
    texts: list[str] = []
    for paragraph in paragraphs_between(source_doc, start, end):
        raw = paragraph.text.strip()
        if not raw or is_caption(paragraph) or has_drawing(paragraph):
            continue
        texts.append(transform_report_text(raw))
    return texts


def extract_source_captions(source_doc: Document, heading: str, next_heading: str) -> list[str]:
    start = find_paragraph(source_doc, heading)
    end = find_paragraph(source_doc, next_heading)
    captions: list[str] = []
    for paragraph in paragraphs_between(source_doc, start, end):
        raw = paragraph.text.strip()
        if raw and is_caption(paragraph):
            captions.append(transform_report_text(raw))
    return captions


def find_first_image_anchor(document: Document, heading: str, next_heading: str) -> Paragraph:
    start = find_paragraph(document, heading)
    end = find_paragraph(document, next_heading)
    candidates = paragraphs_between(document, start, end)
    for idx, paragraph in enumerate(candidates):
        if has_drawing(paragraph):
            return paragraph
        if paragraph.text.strip():
            continue
        for later in candidates[idx + 1 :]:
            if later.text.strip():
                if is_caption(later):
                    return paragraph
                break
    raise ValueError(f"Could not find image anchor for {heading}")


def replace_ch4_section(target_doc: Document, source_doc: Document, heading: str, next_heading: str) -> None:
    target_heading = find_paragraph(target_doc, heading)
    target_next = find_paragraph(target_doc, next_heading)
    source_texts = extract_source_texts(source_doc, heading, next_heading)
    source_captions = extract_source_captions(source_doc, heading, next_heading)
    image_anchor = find_first_image_anchor(target_doc, heading, next_heading)

    target_range = paragraphs_between(target_doc, target_heading, target_next)
    old_text_paragraphs = []
    for paragraph in target_range:
        if paragraph._p is image_anchor._p:
            break
        if paragraph.text.strip() and not is_heading(paragraph) and not is_caption(paragraph):
            old_text_paragraphs.append(paragraph)

    if not old_text_paragraphs:
        raise ValueError(f"No text paragraphs found for section {heading}")

    template = old_text_paragraphs[0]
    for paragraph in old_text_paragraphs:
        remove_paragraph(paragraph)

    for text in source_texts:
        new_para = insert_paragraph_before(image_anchor, template=template, style=template.style.name, text=text)
        new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_font(new_para, size=13)

    target_captions = [p for p in paragraphs_between(target_doc, target_heading, target_next) if p.text.strip() and is_caption(p)]
    if len(target_captions) != len(source_captions):
        raise ValueError(f"Caption count mismatch in {heading}: target={len(target_captions)} source={len(source_captions)}")
    for paragraph, text in zip(target_captions, source_captions):
        replace_paragraph_text(paragraph, text, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    if heading.startswith("4.4."):
        updated_range = paragraphs_between(target_doc, target_heading, target_next)
        for paragraph in updated_range:
            if paragraph.text.startswith("Đối tượng hiển thị trong hình:") and "[7]" not in paragraph.text:
                paragraph.text = (
                    paragraph.text
                    + " Về mặt cảm hứng hình ảnh và cách xây dựng biến thể nguyên tố, nhóm có tham khảo thêm nguồn gameplay chính thức của Genshin Impact [7]."
                )
                set_paragraph_font(paragraph, size=13)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                break


def remove_blocks_between_heading(target_doc: Document, start_prefix: str, end_prefix: str) -> None:
    start = find_paragraph(target_doc, start_prefix)
    end = find_paragraph(target_doc, end_prefix)
    paras = paragraphs_between(target_doc, start, end)
    for paragraph in list(paras):
        if has_drawing(paragraph):
            continue
        remove_paragraph(paragraph)


def add_table_caption_before(anchor: Paragraph, template: Paragraph, caption_text: str) -> Paragraph:
    para = insert_paragraph_before(anchor, template=template, style="Caption", text=caption_text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(para, size=12)
    return para


def set_table_font(table: Table, *, size: int = 11) -> None:
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=size, bold=(row_idx == 0))


def update_ch2_git_sections(target_doc: Document, source_doc: Document) -> None:
    anchor = find_paragraph(target_doc, "2.3.")
    playfab_heading = find_paragraph(target_doc, "2.2.3.")
    template = paragraphs_between(target_doc, playfab_heading, anchor)[0]

    if any(p.text.strip().startswith("2.2.4.") for p in target_doc.paragraphs):
        existing = find_paragraph(target_doc, "2.2.4.")
        for paragraph in list(paragraphs_between(target_doc, existing, anchor)):
            remove_paragraph(paragraph)
        remove_paragraph(existing)

    source_24 = extract_source_texts(source_doc, "2.2.4.", "2.2.5.")
    source_25 = extract_source_texts(source_doc, "2.2.5.", "2.3.")

    blocks = [
        ("Heading 3", "2.2.4. Git trong quản lý phiên bản mã nguồn"),
        *[("Normal", text) for text in source_24],
        ("Caption", "Bảng 2-1: Quy trình quản lý mã nguồn bằng Git trong dự án."),
        ("TABLE", 4),
        ("Heading 3", "2.2.5. GitHub trong quản lý và phối hợp project"),
        *[("Normal", text) for text in source_25],
        ("Caption", "Bảng 2-2: Vai trò của Git local và GitHub remote trong quản lý project."),
        ("TABLE", 5),
    ]

    for kind, payload in blocks:
        if kind == "TABLE":
            source_table = source_doc.tables[payload]
            table = insert_table_before(anchor, source_table)
            set_table_font(table, size=11)
            continue
        para = insert_paragraph_before(anchor, template=template, style=kind, text=str(payload))
        if kind == "Caption":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_font(para, size=12)
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if kind == "Normal" else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_font(para, size=13, bold=(kind.startswith("Heading")))


def replace_intro_between_headings(document: Document, heading_prefix: str, next_prefix: str, text: str) -> None:
    heading = find_paragraph(document, heading_prefix)
    next_heading = find_paragraph(document, next_prefix)
    between = paragraphs_between(document, heading, next_heading)
    if between:
        replace_paragraph_text(between[0], transform_report_text(text), size=13, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        for paragraph in between[1:]:
            if not has_drawing(paragraph):
                remove_paragraph(paragraph)
    else:
        para = insert_paragraph_before(next_heading, template=heading, style="Normal", text=transform_report_text(text))
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_font(para, size=13)


def replace_conclusion(target_doc: Document, source_doc: Document) -> None:
    conclusion = find_paragraph(target_doc, "KẾT LUẬN")
    references = find_paragraph(target_doc, "TÀI LIỆU THAM KHẢO")
    source_conclusion = find_paragraph(source_doc, "KẾT LUẬN")
    source_references = find_paragraph(source_doc, "TÀI LIỆU THAM KHẢO")
    new_texts = extract_source_texts(source_doc, "KẾT LUẬN", "TÀI LIỆU THAM KHẢO")

    old_paragraphs = [p for p in paragraphs_between(target_doc, conclusion, references) if not has_drawing(p)]
    template = old_paragraphs[0] if old_paragraphs else conclusion
    for paragraph in old_paragraphs:
        remove_paragraph(paragraph)

    anchor = references
    for text in new_texts:
        para = insert_paragraph_before(anchor, template=template, style=template.style.name, text=transform_report_text(text))
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_font(para, size=13)


def rebuild_references(document: Document) -> None:
    heading = find_paragraph(document, "TÀI LIỆU THAM KHẢO")
    trailing = paragraph_list(document)[paragraph_index(document, heading) + 1 :]
    template = trailing[0] if trailing else heading
    for paragraph in trailing:
        remove_paragraph(paragraph)

    anchor = heading
    for item in REFERENCE_ITEMS:
        para = insert_paragraph_after(anchor, template=template, style="List Paragraph", text=item)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.left_indent = Cm(0.75)
        para.paragraph_format.first_line_indent = Cm(-0.5)
        para.paragraph_format.space_after = Pt(6)
        set_paragraph_font(para, size=12)
        anchor = para


def append_citation(paragraph: Paragraph, token: str) -> None:
    text = paragraph.text.strip()
    if token in text:
        return
    if text.endswith("."):
        new_text = text[:-1] + f" {token}."
    else:
        new_text = text + f" {token}"
    replace_paragraph_text(paragraph, new_text, size=13, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def add_body_citations(document: Document) -> None:
    append_citation(find_paragraph(document, "C# là ngôn ngữ lập trình hướng đối tượng hiện đại"), "[1]")
    append_citation(find_paragraph(document, "Unity là công cụ phát triển trung tâm của đề tài"), "[2]")
    append_citation(find_paragraph(document, "Để tích hợp bảng xếp hạng trực tuyến"), "[3]")
    append_citation(find_paragraph(document, "Trong dự án, thao tác điều khiển của người chơi được tổ chức"), "[4]")
    append_citation(find_paragraph(document, "Lớp PlayerController đọc dữ liệu di chuyển dưới dạng vector 2 chiều"), "[5]")
    append_citation(find_paragraph(document, "Một đặc điểm đáng chú ý của project là dữ liệu cấu hình được tách"), "[6]")


def renumber_ch3_body(document: Document) -> None:
    ch3 = find_paragraph(document, "CHƯƠNG 3")
    ch4 = find_paragraph(document, "CHƯƠNG 4")
    for paragraph in paragraphs_between(document, ch3, ch4):
        if not paragraph.text.strip():
            continue
        new_text = paragraph.text.replace("Hình 2-", "Hình 3-").replace("Bảng 2-", "Bảng 3-")
        new_text = new_text.replace("GetPlayerLeaderboardDt()", "GetPlayerLeaderboardData()")
        new_text = new_text.replace("EnemyGroup   -> ObjectPool", "EnemyGroup -> ObjectPool")
        if new_text != paragraph.text:
            size = 12 if is_caption(paragraph) else 13
            align = WD_ALIGN_PARAGRAPH.CENTER if is_caption(paragraph) else WD_ALIGN_PARAGRAPH.JUSTIFY
            replace_paragraph_text(paragraph, new_text, size=size, alignment=align)


def normalize_ch4_text_and_captions(document: Document) -> None:
    ch4 = find_paragraph(document, "CHƯƠNG 4")
    conclusion = find_paragraph(document, "KẾT LUẬN")
    for paragraph in paragraphs_between(document, ch4, conclusion):
        if not paragraph.text.strip():
            continue
        new_text = transform_report_text(paragraph.text)
        new_text = new_text.replace("Hình 3-", "Hình 4-")
        if new_text != paragraph.text:
            size = 12 if is_caption(paragraph) else 13
            align = WD_ALIGN_PARAGRAPH.CENTER if is_caption(paragraph) else WD_ALIGN_PARAGRAPH.JUSTIFY
            replace_paragraph_text(paragraph, new_text, size=size, alignment=align)


def split_entry_page(text: str) -> tuple[str, str]:
    if "\t" in text:
        left, right = text.rsplit("\t", 1)
        return left.strip(), right.strip()
    return text.strip(), ""


def update_front_matter_lists(document: Document) -> None:
    table_heading = find_paragraph(document, "DANH MỤC BẢNG BIỂU")
    figure_heading = find_paragraph(document, "DANH MỤC HÌNH VẼ")
    chapter1 = find_paragraph(document, "CHƯƠNG 1")

    table_range = paragraphs_between(document, table_heading, figure_heading)
    table_entries = [p for p in table_range if p.style and p.style.name == "table of figures"]
    blank_slots = [p for p in table_range if not p.text.strip()]

    figure_entries = [p for p in paragraphs_between(document, figure_heading, chapter1) if p.style and p.style.name == "table of figures"]
    if len(table_entries) < 8 or len(blank_slots) < 2 or len(figure_entries) < 47:
        raise ValueError("Unexpected front matter structure; cannot safely update entry lists.")

    ch2 = find_paragraph(document, "2.2.4.")
    ch4 = find_paragraph(document, "CHƯƠNG 4")
    body_table_captions = [p.text.strip() for p in paragraphs_between(document, ch2, ch4) if p.text.strip().startswith("Bảng ")]
    if len(body_table_captions) != 7:
        raise ValueError(f"Unexpected body table caption count: {len(body_table_captions)}")

    table_labels = [split_entry_page(p.text)[0] for p in table_entries[:3]] + body_table_captions
    table_slots = table_entries + blank_slots[:2]
    for paragraph, label, page in zip(table_slots[:10], table_labels, TABLE_FRONT_PAGES):
        replace_paragraph_text(paragraph, f"{label}\t{page}", size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        paragraph.style = "table of figures"

    ch3 = find_paragraph(document, "CHƯƠNG 3")
    conclusion = find_paragraph(document, "KẾT LUẬN")
    body_figure_captions = [p.text.strip() for p in paragraphs_between(document, ch3, conclusion) if p.text.strip().startswith("Hình ")]
    if len(body_figure_captions) != 34:
        raise ValueError(f"Unexpected body figure caption count: {len(body_figure_captions)}")

    figure_pages = [split_entry_page(p.text)[1] for p in figure_entries]
    figure_labels = [split_entry_page(p.text)[0] for p in figure_entries[:13]] + body_figure_captions
    for paragraph, label, page in zip(figure_entries, figure_labels, figure_pages):
        replace_paragraph_text(paragraph, f"{label}\t{page}", size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        paragraph.style = "table of figures"


def set_all_text_black(document: Document) -> None:
    for paragraph in document.paragraphs:
        size = 12 if is_caption(paragraph) else 13
        bold = None
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            bold = True
        set_paragraph_font(paragraph, size=size, bold=bold)
    for table in document.tables:
        set_table_font(table, size=11)


def main() -> None:
    if not INPUT_DOC.exists():
        raise FileNotFoundError(f"Missing input DOCX: {INPUT_DOC}")
    if not SOURCE_DOC.exists():
        raise FileNotFoundError(f"Missing source DOCX: {SOURCE_DOC}")

    BACKUP_DOC.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(INPUT_DOC, BACKUP_DOC)
    shutil.copy2(INPUT_DOC, OUTPUT_DOC)

    target_doc = Document(OUTPUT_DOC)
    source_doc = Document(SOURCE_DOC)

    update_ch2_git_sections(target_doc, source_doc)
    replace_intro_between_headings(
        target_doc,
        "CHƯƠNG 4",
        "4.1.",
        "Chương này trình bày các giao diện chính của chương trình và cách sử dụng những chức năng quan trọng trong quá trình chơi. Các hình minh họa được giữ theo đúng bố cục triển khai thực tế của project để người đọc dễ đối chiếu với sản phẩm đã xây dựng.",
    )

    for current, next_heading in zip(CH4_SECTIONS, CH4_SECTIONS[1:] + ["KẾT LUẬN"]):
        replace_ch4_section(target_doc, source_doc, current, next_heading)

    replace_conclusion(target_doc, source_doc)
    rebuild_references(target_doc)
    add_body_citations(target_doc)
    renumber_ch3_body(target_doc)
    normalize_ch4_text_and_captions(target_doc)
    update_front_matter_lists(target_doc)
    set_all_text_black(target_doc)

    target_doc.save(OUTPUT_DOC)
    print(f"Updated report written to {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
