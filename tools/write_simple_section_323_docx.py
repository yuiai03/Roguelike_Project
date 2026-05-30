from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
BASE_DOC = ROOT / "output" / "doc" / "BaoCao_DATN_2_spawn_and_14_inserted.docx"
FULL_OUT = ROOT / "output" / "doc" / "BaoCao_DATN_2_spawn_and_14_inserted_3_2_3_don_gian.docx"
SECTION_OUT = ROOT / "output" / "doc" / "Muc_3_2_3_spawn_quai_don_gian.docx"

HEADING_TEXT = "3.2.3. Luồng hệ thống spawn quái"
NEXT_HEADING_TEXT = "3.2.4. Luồng chiến đấu, nhận EXP và chọn buff"

PARAGRAPHS = [
    "Sequence Diagram này mô tả cách hệ thống tạo quái mỗi khi một wave mới bắt đầu. Người chơi không sinh quái trực tiếp; thao tác bắt đầu trận chỉ là tín hiệu để WaveSpawner tiếp nhận và điều phối toàn bộ quá trình spawn.",
    "Đầu tiên, ChallengePanel gọi StartNextWave. WaveSpawner tăng số wave hiện tại, lấy dữ liệu wave từ WaveConfig và tạo một wave session mới để tránh lẫn dữ liệu giữa các wave. Nếu wave vượt quá phần cấu hình có sẵn, hệ thống tạo một wave mới bằng cách sao chép từ một wave mẫu gần nhất.",
    "Sau đó hệ thống chờ preparationTime, thông báo wave bắt đầu và gửi số lượng quái ban đầu lên HUD. Với wave thường, WaveSpawner đọc từng nhóm quái trong enemyGroups rồi lần lượt sinh quái theo nhóm. Mỗi quái được lấy từ ObjectPool để xuất hiện nhanh hơn và sau mỗi đợt spawn, HUD được cập nhật lại số quái đang hoạt động.",
    "Với boss wave, hệ thống chọn loại boss và vị trí xuất hiện trước khi sinh boss trực tiếp hoặc qua hiệu ứng triệu hồi. Các thông số quan trọng của luồng này gồm preparationTime, enemyCount, spawnDelay, spawnPosition, spreadRadius, bossPoolTypes và bossSpawnPosition. Sequence này dừng ở thời điểm quái đã được sinh ra và HUD đã cập nhật xong, chưa đi sang bước chiến đấu hay hoàn tất wave.",
]


def set_run_font(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text
    set_run_font(paragraph)


def insert_paragraph_before(paragraph, text: str, style_name: str | None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    if style_name:
        try:
            new_para.style = style_name
        except KeyError:
            pass
    run = new_para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    return new_para


def update_full_report() -> None:
    doc = Document(BASE_DOC)
    heading_idx = None
    next_idx = None

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == HEADING_TEXT:
            heading_idx = i
        elif heading_idx is not None and text == NEXT_HEADING_TEXT:
            next_idx = i
            break

    if heading_idx is None or next_idx is None:
        raise RuntimeError("Không tìm thấy mục 3.2.3 trong file báo cáo gốc.")

    text_start = heading_idx + 4
    existing_count = max(0, next_idx - text_start)
    style_name = doc.paragraphs[text_start].style.name if existing_count > 0 and doc.paragraphs[text_start].style else None

    overlap = min(existing_count, len(PARAGRAPHS))
    for offset in range(overlap):
        set_paragraph_text(doc.paragraphs[text_start + offset], PARAGRAPHS[offset])

    for idx in range(text_start + overlap, next_idx):
        doc.paragraphs[idx].text = ""

    if len(PARAGRAPHS) > existing_count:
        anchor = doc.paragraphs[next_idx]
        for extra in PARAGRAPHS[existing_count:]:
            insert_paragraph_before(anchor, extra, style_name)

    doc.save(FULL_OUT)


def build_section_doc() -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(13)

    heading = doc.add_paragraph()
    run = heading.add_run(HEADING_TEXT)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)

    source = doc.add_paragraph("Tệp sơ đồ nguồn: 3_2_3_sequence_spawn_quai.drawio - 3.2.3. Luồng hệ thống spawn quái.")
    set_run_font(source)

    caption = doc.add_paragraph("Hình 3.2.3: Luồng hệ thống spawn quái.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(caption)

    for text in PARAGRAPHS:
        paragraph = doc.add_paragraph(text)
        set_run_font(paragraph)

    doc.save(SECTION_OUT)


def main() -> None:
    update_full_report()
    build_section_doc()
    print(FULL_OUT)
    print(SECTION_OUT)


if __name__ == "__main__":
    main()
