from __future__ import annotations

from pathlib import Path
from typing import Iterable, List, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "tmp" / "docs" / "BaoCao_DATN_2_working.docx"
OUTPUT_DIR = ROOT / "output" / "doc"
ASSET_DIR = ROOT / "tmp" / "docs" / "generated_assets"

FONT_REGULAR = r"C:\Windows\Fonts\times.ttf"
FONT_BOLD = r"C:\Windows\Fonts\timesbd.ttf"


FUNCTION_BLOCKS = [
    {
        "title": "1.4.1.1. Chức năng định danh người chơi",
        "mo_ta": (
            "Chức năng định danh bảo đảm mỗi người chơi có một tài khoản PlayFab hợp lệ trước khi bước vào "
            "vòng lặp gameplay và trước khi sử dụng leaderboard."
        ),
        "luong": (
            "1) PlayFabLeaderboardManager khởi tạo và gọi LoginWithCustomID. 2) Hệ thống đọc Display Name từ "
            "hồ sơ PlayFab. 3) Nếu chưa có tên, NameInputPanel mở để nhận dữ liệu người chơi nhập. 4) Khi tên "
            "được chấp nhận, client đồng bộ CurrentPlayFabId và CurrentDisplayName để dùng cho toàn bộ phiên chơi."
        ),
        "du_lieu": (
            "Đầu vào gồm Custom ID lưu trong PlayerPrefs hoặc SystemInfo.deviceUniqueIdentifier, phản hồi hồ sơ "
            "từ PlayFab, và chuỗi tên người chơi nhập từ UI. Đầu ra là danh tính đã xác thực, sẵn sàng dùng cho "
            "gửi điểm và truy xuất thứ hạng."
        ),
        "specs": [
            ("Kích hoạt", "Start của PlayFabLeaderboardManager khi scene gameplay được mở."),
            ("Thành phần", "PlayFabLeaderboardManager, NameInputPanel, PlayerPrefs, PlayFab Client API."),
            ("Thông số chính", "Display Name hợp lệ từ 3 đến 25 ký tự; thống kê leaderboard dùng tên HighScore."),
            ("Trạng thái", "CurrentPlayFabId, CurrentDisplayName, OnProfileLoadedEvent, OnSubmitNameFailed."),
            ("Kết quả", "Người chơi có tài khoản và tên hiển thị hợp lệ trước khi tham gia gameplay chính."),
        ],
    },
    {
        "title": "1.4.1.2. Chức năng bắt đầu trận đấu",
        "mo_ta": (
            "Chức năng này chuyển hệ thống từ trạng thái chờ trong scene sang trạng thái chiến đấu và khởi tạo "
            "wave đầu tiên."
        ),
        "luong": (
            "1) Player tiến vào trigger của ChallengePostNPC. 2) InteractPanel gợi ý thao tác tương tác. 3) "
            "ChallengePanel mở, khóa input và chờ nhấn StartGame. 4) onGameStart được phát, HUD được bind dữ liệu "
            "và WaveSpawner.StartNextWave được gọi để bắt đầu trận."
        ),
        "du_lieu": (
            "Đầu vào gồm trạng thái playerInRange, phím tương tác, nút StartGame và sự tồn tại của ChallengePanel. "
            "Đầu ra là currentWave tăng lên, HUD bắt đầu hiển thị và NPC bị vô hiệu hóa cho phiên chơi hiện tại."
        ),
        "specs": [
            ("Kích hoạt", "Player đứng trong vùng trigger và nhấn F để tương tác với ChallengePostNPC."),
            ("Thành phần", "NPC, ChallengePostNPC, InteractPanel, ChallengePanel, PlayerController, WaveSpawner, PlayerStatsPanel."),
            ("Thông số chính", "tutorialDelayTime mặc định 1 giây trước khi nút StartGame hiện rõ."),
            ("Trạng thái", "Player input bị khóa khi panel mở và được mở lại ngay trước khi vào chiến đấu."),
            ("Kết quả", "Wave đầu tiên được chuẩn bị và HUD chính thức đi vào trạng thái theo dõi trận đấu."),
        ],
    },
    {
        "title": "1.4.1.3. Chức năng điều khiển nhân vật",
        "mo_ta": (
            "Chức năng điều khiển nhân vật xử lý di chuyển, xoay model, gravity, grounded check và dash trong "
            "môi trường 3D."
        ),
        "luong": (
            "1) PlayerController nhận Vector2 Move từ Input System. 2) Hướng di chuyển được quy đổi theo camera. "
            "3) CharacterController di chuyển nhân vật và cập nhật gravity. 4) Khi người chơi nhấn chuột phải và "
            "dash cooldown bằng 0, hệ thống chuyển sang trạng thái dash rồi quay lại di chuyển thường."
        ),
        "du_lieu": (
            "Đầu vào gồm moveInput, hướng camera, trạng thái grounded, dashPressed và cờ isInputActive. Đầu ra là "
            "vị trí nhân vật, hướng quay model và trạng thái animation Idle/Run/Dash."
        ),
        "specs": [
            ("Thành phần", "PlayerController, InputSystem_Actions, CharacterController, PlayerAnimationController, Camera.main."),
            ("Thông số mặc định", "moveSpeed=5, gravity=-20, rotationSpeed=1440, groundDistance=0.3."),
            ("Thông số dash", "dashSpeed=15, dashDuration=0.2 giây, dashCooldown=1 giây."),
            ("Trạng thái", "isGrounded, isDashing, dashTimer, dashCooldownTimer, isInputActive."),
            ("Kết quả", "Player di chuyển mượt theo camera và không giữ input cũ khi UI khóa điều khiển."),
        ],
    },
    {
        "title": "1.4.1.4. Chức năng chiến đấu của người chơi",
        "mo_ta": (
            "Hệ thống chiến đấu của người chơi được tự động hóa ở khâu tìm mục tiêu và bắn, còn chiều sâu chiến thuật "
            "đến từ buff và vị trí di chuyển."
        ),
        "luong": (
            "1) PlayerAttack đếm lùi attackTimer. 2) Khi timer về 0, hệ thống quét enemy trong attackRange để tìm "
            "mục tiêu gần nhất. 3) ObjectPool sinh projectile và khởi tạo hướng bay, sát thương, thời gian sống. "
            "4) Nếu player có multishot hoặc AoE, projectile được gắn thêm dữ liệu mở rộng trước khi bắn."
        ),
        "du_lieu": (
            "Đầu vào gồm attackRange, attackCooldown, attackDamage, projectileSpeed, projectileLifetime, multiShotCount, "
            "multiShotAngle, aoeRadius và aoeAtkMultiplier từ PlayerData. Đầu ra là các projectile đang hoạt động và "
            "lượng sát thương hợp lệ tác động lên enemy."
        ),
        "specs": [
            ("Thành phần", "PlayerAttack, PlayerProjectile, ObjectPool, PlayerData, IDamageable."),
            ("Thông số mặc định", "attackDamage=50, attackCooldown=1 giây, attackRange=20."),
            ("Thông số đạn", "projectileSpeed=30, projectileLifetime=10 giây."),
            ("Thông số buff", "multiShotCount mặc định 1, multiShotAngle mặc định 12 độ, AoE bật/tắt theo buff."),
            ("Kết quả", "Projectile được tái sử dụng qua pool và phản ánh đúng sức mạnh hiện tại của player."),
        ],
    },
    {
        "title": "1.4.1.5. Chức năng quản lý enemy",
        "mo_ta": (
            "Chức năng quản lý enemy xây dựng các nhóm đối kháng có vai trò khác nhau và duy trì vòng đời spawn - "
            "truy đuổi - tấn công - chết - trả về pool."
        ),
        "luong": (
            "1) WaveSpawner chọn PoolType và spawn enemy theo wave. 2) EnemyData nạp chỉ số từ EnemyConfig. 3) "
            "Enemy chạy AI bằng CharacterController, bám theo Player hoặc giữ cự ly bắn. 4) Khi máu về 0, OnDeath "
            "được phát để cộng EXP, cập nhật wave và dọn trạng thái trước khi despawn."
        ),
        "du_lieu": (
            "Đầu vào gồm enemyType, maxHealth, moveSpeed, attackRange, contactDamage, projectileDamage, cooldown, "
            "expValue và randomVariation. Đầu ra là enemy hoạt động đúng vai trò cận chiến, đánh xa, bay hoặc boss."
        ),
        "specs": [
            ("Thành phần", "Enemy, EnemyData, EnemyConfig, MeleeEnemy, RangedEnemy, FlyEnemy, BossEnemy."),
            ("Thông số dữ liệu", "EnemyData mặc định có maxHealth=100, moveSpeed=5, expValue=10, randomVariation=0.15."),
            ("Thông số cận chiến", "contactDamage=10, attackCooldown=1 giây, lungeSpeed=12, retreatSpeed=6."),
            ("Thông số đánh xa", "projectileDamage=15, projectileSpeed=10, shootCooldown=2 giây."),
            ("Kết quả", "Enemy giữ hành vi riêng nhưng vẫn dùng chung pipeline máu, sát thương, OnDeath và pooling."),
        ],
    },
    {
        "title": "1.4.1.6. Chức năng quản lý wave và độ khó",
        "mo_ta": (
            "WaveSpawner là bộ điều phối nhịp trận đấu, quyết định khi nào spawn enemy, khi nào hoàn tất wave, khi nào "
            "xuất hiện boss và khi nào chuyển sang endless wave."
        ),
        "luong": (
            "1) StartNextWave tăng currentWave và chọn dữ liệu wave từ WaveConfig. 2) Hệ thống chờ preparationTime, "
            "sau đó spawn enemy theo từng EnemyGroup hoặc boss riêng. 3) activeEnemies và totalEnemiesSpawned được "
            "theo dõi liên tục. 4) Khi tất cả enemy của wave bị hạ, CompleteWave kích hoạt chuyển sang wave kế tiếp."
        ),
        "du_lieu": (
            "Đầu vào gồm waves, enemyGroups, spawnPosition, spreadRadius, spawnDelay, isBossWave, bossPoolTypes, "
            "autoScale và scalePerWave. Đầu ra là danh sách enemy đang hoạt động, currentWave và các sự kiện wave."
        ),
        "specs": [
            ("Thành phần", "WaveSpawner, WaveConfig, EnemyGroup, ObjectPool, MapThemeManager."),
            ("Thông số wave", "preparationTime mặc định 3 giây; Generate30Waves đặt boss wave ở mốc 10, 20, 30 với 5 giây chuẩn bị."),
            ("Thông số hiệu năng", "maxEnemySpawnsPerFrame=8, maxSpawnEffectsPerFrame=12, spawnRandomRadius=2."),
            ("Thông số tăng độ khó", "autoScale bật theo WaveConfig và scalePerWave mặc định 1.1."),
            ("Kết quả", "Trận đấu giữ được nhịp tăng dần, có boss wave và tiếp tục endless khi vượt danh sách wave gốc."),
        ],
    },
    {
        "title": "1.4.1.7. Chức năng đổi theme bản đồ",
        "mo_ta": (
            "Hệ thống đổi theme làm mới môi trường chiến đấu theo chu kỳ wave mà không phá vỡ dữ liệu gameplay hiện hành."
        ),
        "luong": (
            "1) Sau khi kết thúc một wave, WaveSpawner dự báo upcomingWave. 2) MapThemeManager kiểm tra theme có cần đổi "
            "không bằng ResolveThemeIndexForWave. 3) LoadingUIManager che màn hình, áp vật liệu và effectRoot mới. "
            "4) Sau khi transition hoàn tất, gameplay và input được mở lại để bắt đầu wave tiếp theo."
        ),
        "du_lieu": (
            "Đầu vào gồm currentWave, upcomingWave, danh sách themes, groundMaterial, wallMaterial và effectRoot. "
            "Đầu ra là chỉ số CurrentThemeIndex mới và môi trường hình ảnh được đồng bộ với tiến trình wave."
        ),
        "specs": [
            ("Thành phần", "WaveSpawner, MapThemeManager, LoadingUIManager, MeshRenderer mặt đất và tường."),
            ("Chu kỳ", "Theme đổi theo công thức ((wave-1)/10) nên mỗi 10 wave chuyển một theme."),
            ("Thông số transition", "fadeInDuration=0.4, holdDuration=0.2, fadeOutDuration=0.4."),
            ("Trạng thái", "IsTransitioning, CurrentThemeIndex, OnThemeTransitionCompleted."),
            ("Kết quả", "Bản đồ đổi theme rõ ràng nhưng không làm mất trạng thái wave hoặc input của player."),
        ],
    },
    {
        "title": "1.4.1.8. Chức năng kinh nghiệm và lên cấp",
        "mo_ta": (
            "Chức năng EXP và lên cấp tạo trục tăng tiến sức mạnh của người chơi, đồng thời định thời điểm hiển thị buff."
        ),
        "luong": (
            "1) Enemy phát OnDeath. 2) ExpDropper đọc expValue từ EnemyData và cộng thẳng vào PlayerLevelSystem. 3) "
            "PlayerLevelSystem cập nhật currentExp và totalExpGained. 4) Nếu currentExp vượt expToNextLevel, hệ thống "
            "lặp LevelUp, nâng currentLevel và phát OnLevelUp cho UI chọn buff."
        ),
        "du_lieu": (
            "Đầu vào gồm expValue, currentExp, expToNextLevel, currentLevel và expScalingFactor. Đầu ra là level hiện tại, "
            "thanh EXP mới và tín hiệu kích hoạt chọn buff."
        ),
        "specs": [
            ("Thành phần", "Enemy, ExpDropper, PlayerLevelSystem, PlayerStatsPanel."),
            ("Thông số mặc định", "currentLevel=0, currentExp=0, expToNextLevel=100, expScalingFactor=1.1."),
            ("Thông số dữ liệu", "totalExpGained được dùng trực tiếp làm điểm leaderboard cuối trận."),
            ("Sự kiện", "OnExpChanged, OnLevelChanged, OnLevelUp."),
            ("Kết quả", "Người chơi tăng cấp theo tiến trình chiến đấu và mọi màn hình HUD cập nhật cùng một nguồn dữ liệu."),
        ],
    },
    {
        "title": "1.4.1.9. Chức năng chọn buff và tăng sức mạnh",
        "mo_ta": (
            "Buff là lớp cá nhân hóa chiến thuật của từng lượt chơi, cho phép người chơi tăng chỉ số hoặc mở rộng kỹ năng."
        ),
        "luong": (
            "1) CardSelectionPanel nhận OnLevelUp từ PlayerLevelSystem. 2) BuffCardManager lọc các card chưa đạt maxLevel "
            "và chọn ngẫu nhiên theo rarity cùng luckBonus. 3) UI sinh đúng số card cho phép chọn. 4) Khi người chơi "
            "nhấn một card, BuffCardManager.ApplyCard cập nhật PlayerData, PlayerHealth hoặc manager kỹ năng tương ứng."
        ),
        "du_lieu": (
            "Đầu vào gồm danh sách allCards, cardsPerSelection, rarity, maxLevel, cardLevels và trạng thái hiện tại của "
            "PlayerData. Đầu ra là bộ chỉ số hoặc kỹ năng mới đang hoạt động ngay trong runtime."
        ),
        "specs": [
            ("Thành phần", "BuffCardManager, BuffCardConfig, CardSelectionPanel, PlayerData, PlayerHealth."),
            ("Thông số chính", "cardsPerSelection=3 cho mỗi lần hiển thị."),
            ("Loại buff", "Buff chỉ số: damage, move speed, attack speed, max health, EXP bonus; buff kỹ năng: multishot, AoE, spirit, orbiting ball."),
            ("Điều kiện dừng game", "Time.timeScale=0 khi panel mở và chỉ mở lại khi người chơi đã chọn xong."),
            ("Kết quả", "Người chơi nhận được đúng một lựa chọn nâng cấp cho mỗi lần hiển thị card hợp lệ."),
        ],
    },
    {
        "title": "1.4.1.10. Chức năng giao diện trong trận",
        "mo_ta": (
            "Lớp giao diện trong trận phản ánh trạng thái gameplay theo thời gian thực và điều phối quyền sở hữu input giữa "
            "các panel."
        ),
        "luong": (
            "1) GameUI gom toàn bộ tham chiếu panel. 2) PlayerStatsPanel bind OnHealthChanged, OnExpChanged, OnLevelChanged "
            "và OnWaveStart để cập nhật HUD. 3) InteractPanel, NotiPanel, CardSelectionPanel, LeaderboardPanel và "
            "NameInputPanel được bật hoặc tắt theo sự kiện tương ứng. 4) Khi có panel ưu tiên cao mở, input gameplay bị khóa."
        ),
        "du_lieu": (
            "Đầu vào là dữ liệu từ PlayerHealth, PlayerLevelSystem, WaveSpawner và PlayFabLeaderboardManager. Đầu ra là HUD, "
            "prompt tương tác, panel thông báo và các màn hình phụ đồng bộ với đúng trạng thái trận đấu."
        ),
        "specs": [
            ("Thành phần", "GameUI, PlayerStatsPanel, InteractPanel, NotiPanel, CardSelectionPanel, NameInputPanel, LeaderboardPanel, PauseMenuPanel."),
            ("HUD chính", "Máu, EXP, level và wave được cập nhật từ event runtime thay vì polling thủ công."),
            ("Thông số hiển thị", "Wave dùng FormatWaveLabel; EXP hiển thị totalExpGained/totalExpRequiredForNextLevel."),
            ("Ràng buộc", "Các panel không được chồng sai quyền sở hữu input hoặc che lẫn nhau ở thời điểm quan trọng."),
            ("Kết quả", "Người chơi luôn nhìn thấy đúng thông tin cần thiết để ra quyết định trong trận."),
        ],
    },
    {
        "title": "1.4.1.11. Chức năng tạm dừng và kết thúc trận",
        "mo_ta": (
            "Chức năng này bao phủ hai trạng thái ngắt vòng lặp gameplay: tạm dừng để điều hướng/cài đặt và kết thúc trận "
            "khi người chơi chết hoặc chủ động rời lượt chơi."
        ),
        "luong": (
            "1) Người chơi nhấn ESC và PauseMenuPanel kiểm tra điều kiện mở. 2) Nếu hợp lệ, Time.timeScale bị dừng, input "
            "gameplay bị khóa và overlay pause hiện ra. 3) Người chơi có thể Resume, mở settings, xem leaderboard hoặc quit. "
            "4) Nếu máu về 0, PlayerHealth chạy death sequence, dọn gameplay object và hiển thị leaderboard sau khi gửi điểm."
        ),
        "du_lieu": (
            "Đầu vào gồm trạng thái panel đang mở, currentView của PauseMenuPanel, cờ isDead và currentRunScore. Đầu ra là "
            "trạng thái timeScale, input, kết thúc lượt chơi và giao diện tổng kết."
        ),
        "specs": [
            ("Thành phần", "PauseMenuPanel, LeaderboardPanel, PlayerController, PlayerHealth, PlayFabLeaderboardManager."),
            ("Điều kiện mở", "ESC chỉ có hiệu lực khi các panel khóa cứng khác không giữ quyền điều khiển."),
            ("Thông số kết thúc trận", "PlayerHealth chờ QuitScoreSubmitTimeoutSeconds=1.5, DeathCleanupDelaySeconds=1, DeathLeaderboardDelaySeconds=0.5."),
            ("Điểm số", "currentRunScore = Floor(totalExpGained) trước khi gửi lên PlayFab."),
            ("Kết quả", "Người chơi có thể dừng trận an toàn hoặc kết thúc lượt chơi với chuỗi xử lý sạch và có lưu điểm."),
        ],
    },
    {
        "title": "1.4.1.12. Chức năng bảng xếp hạng",
        "mo_ta": (
            "Bảng xếp hạng là lớp trình bày của hệ backend, dùng để đồng bộ thành tích trực tuyến và phản hồi lại vị trí "
            "của người chơi trong hệ thống."
        ),
        "luong": (
            "1) PlayerHealth hoặc LeaderboardPanel yêu cầu dữ liệu xếp hạng. 2) PlayFabLeaderboardManager gửi SubmitScore "
            "hoặc GetLeaderboardData/GetLeaderboardAroundPlayer. 3) PlayFab trả về danh sách top và vị trí quanh người chơi. "
            "4) LeaderboardPanel sinh entry UI, đánh dấu dòng của tài khoản hiện tại và hiển thị thứ hạng."
        ),
        "du_lieu": (
            "Đầu vào là totalExpGained của PlayerLevelSystem, CurrentDisplayName, CurrentPlayFabId và statistic HighScore. "
            "Đầu ra là top leaderboard, hạng cá nhân và màn hình tổng kết của lượt chơi."
        ),
        "specs": [
            ("Thành phần", "PlayerHealth, PlayerLevelSystem, PlayFabLeaderboardManager, LeaderboardPanel, LeaderboardEntryUI."),
            ("Thông số truy xuất", "GetLeaderboardData lấy tối đa 100 kết quả; GetLeaderboardAroundPlayer lấy 1 kết quả quanh người chơi."),
            ("Thông số điểm", "Score gửi đi là số nguyên lấy từ Floor(totalExpGained)."),
            ("Dữ liệu hiển thị", "Tên hiển thị, hạng, PlayFabId và điểm số của từng entry."),
            ("Kết quả", "Người chơi thấy được thành tích của mình và bối cảnh cạnh tranh trực tuyến sau mỗi lượt chơi."),
        ],
    },
]


CH3_INSERTS = {
    "3.1.1. Use Case tổng quan": (
        "Luồng và dữ liệu chính: actor Người chơi kích hoạt ba nhóm hành vi lớn là định danh, bắt đầu trận và tham gia vòng "
        "chiến đấu; hệ thống phản hồi bằng HUD, lựa chọn buff và đồng bộ leaderboard."
    ),
    "3.1.2. Use Case chi tiết": (
        "Thông số cần nhấn mạnh trong Use Case chi tiết là các điều kiện chuyển bước như có/không có Display Name, còn/hết "
        "máu, đủ/chưa đủ EXP và hoàn thành/chưa hoàn thành wave."
    ),
    "3.2.1. Luồng đăng nhập và nhập tên hiển thị": (
        "Dữ liệu vào/ra của sequence này gồm Custom ID cục bộ, hồ sơ PlayerProfile và chuỗi Display Name sau khi được xác "
        "thực với ràng buộc độ dài 3 đến 25 ký tự."
    ),
    "3.2.2. Luồng chiến đấu, nhận EXP và chọn buff": (
        "Thông số chính của sequence này là attackCooldown, attackRange, expValue, expToNextLevel và cardsPerSelection=3, "
        "vì đây là các tham số quyết định nhịp chiến đấu và nhịp tăng trưởng."
    ),
    "3.2.3. Luồng game over, gửi điểm và leaderboard": (
        "Điểm nhấn của sequence kết thúc trận là điểm số không dựa vào thời gian sống mà lấy trực tiếp từ totalExpGained, "
        "sau đó được gửi dưới statistic HighScore lên PlayFab."
    ),
    "3.3. Activity Diagram": (
        "Activity Diagram nên được đọc theo các nút điều kiện: có Display Name hay chưa, enemy đã chết hay chưa, đủ EXP để "
        "lên cấp hay chưa, wave đã kết thúc hay chưa và player còn sống hay không."
    ),
    "3.4. Component Diagram": (
        "Ở góc nhìn component, dữ liệu đi qua ba trục chính: input và trạng thái local, dữ liệu chiến đấu runtime, và dữ "
        "liệu dịch vụ trực tuyến phục vụ định danh cùng leaderboard."
    ),
    "3.5.1. Nhóm Player và Progression": (
        "Thông số trọng tâm của nhóm lớp này gồm moveSpeed, dashCooldown, attackCooldown, currentLevel, expToNextLevel và "
        "toàn bộ các bonus đang được BuffCardManager áp dụng lên PlayerData."
    ),
    "3.5.2. Nhóm Enemy và Projectile": (
        "Với nhóm lớp Enemy và Projectile, luồng dữ liệu xoay quanh chỉ số từ EnemyConfig/PlayerData, trạng thái va chạm, "
        "và sự kiện OnDeath hoặc despawn để quay lại ObjectPool."
    ),
    "3.5.3. Nhóm UI, NPC và Backend": (
        "Nhóm lớp này cần được hiểu như tầng điều hướng và đồng bộ: NPC mở luồng gameplay, UI chiếm quyền input theo trạng "
        "thái, còn PlayFabLeaderboardManager đảm nhận chuỗi xác thực - gửi điểm - tải xếp hạng."
    ),
}


CH4_INSERTS = {
    "4.1. Toàn cảnh gameplay trong scene Game": (
        "Scene Game liên kết trực tiếp PlayerController, PlayerAttack, WaveSpawner, PlayerStatsPanel "
        "và ObjectPool; đây là nơi thể hiện rõ nhất vòng lặp vào trận - spawn quái - chiến đấu - lên cấp - tiếp tục wave."
    ),
    "4.2. Nhân vật chính của trò chơi": (
        "Thông số nên đi cùng hình minh họa player là moveSpeed=5, dashSpeed=15, dashCooldown=1 giây, attackDamage=50 và "
        "attackRange=20 ở cấu hình mặc định để người đọc gắn được hình ảnh với tham số runtime."
    ),
    "4.3. Danh sách enemy cơ bản trong trò chơi": (
        "Luồng xử lý kỹ thuật của nhóm này gồm spawn từ pool, nạp chỉ số từ EnemyConfig, tiếp cận hoặc giữ cự ly theo "
        "enemyType, và phát OnDeath để cộng EXP khi bị hạ."
    ),
    "4.4. Nhóm enemy nguyên tố LawaChurl": (
        "Thông số cần nhấn mạnh ở nhóm boss là bossPoolTypes của wave, projectileDamage/projectileSpeed riêng của boss "
        "và các hiệu ứng cảnh báo hoặc diện rộng khiến boss wave khác biệt với wave thường."
    ),
    "4.5. Các theme map và chuyển đổi môi trường": (
        "Chu kỳ đổi theme được điều khiển theo mốc 10 wave/lần, với transition có fadeInDuration=0.4, holdDuration=0.2 "
        "và fadeOutDuration=0.4 để tránh chuyển cảnh đột ngột."
    ),
    "4.6. Giao diện HUD và hỗ trợ gameplay": (
        "Các thông số HUD quan trọng gồm máu hiện tại/tối đa, totalExpGained trên thanh EXP, level hiện tại và nhãn wave "
        "được tạo bởi FormatWaveLabel của WaveSpawner."
    ),
    "4.7. Giao diện challenge và chọn buff": (
        "Luồng giao diện ở đây có hai pha rõ ràng: ChallengePanel mở trước trận và CardSelectionPanel mở khi OnLevelUp xảy "
        "ra; cả hai đều khóa input gameplay để tránh xung đột thao tác."
    ),
    "4.8. Giao diện nhập tên và leaderboard": (
        "Thông số dịch vụ gắn với nhóm giao diện này là Display Name dài 3 đến 25 ký tự, statistic HighScore và top 100 "
        "entry leaderboard được tải về để hiển thị."
    ),
    "4.9. Giao diện tạm dừng, cài đặt âm thanh và loading": (
        "Phần giao diện phụ trợ cần mô tả rõ điều kiện mở bằng phím ESC, quan hệ giữa pause và quyền input, cũng như vai "
        "trò của loading transition khi game đổi theme hoặc reload màn chơi."
    ),
}


def ensure_dirs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph: Paragraph, size: float = 13) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size)


def set_table_font(table: Table, size: float = 12.5) -> None:
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_paragraph_font(paragraph, size=size)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_picture_after(paragraph: Paragraph, image_path: Path, width_inches: float) -> Paragraph:
    new_para = insert_paragraph_after(paragraph, style="Normal")
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    set_paragraph_font(new_para)
    return new_para


def insert_labeled_paragraph_after(paragraph: Paragraph, label: str, text: str) -> Paragraph:
    new_para = insert_paragraph_after(paragraph, style="Normal")
    first = new_para.add_run(f"{label}: ")
    set_run_font(first, size=13, bold=True)
    second = new_para.add_run(text)
    set_run_font(second, size=13)
    new_para.paragraph_format.space_after = Pt(6)
    return new_para


def insert_table_after(document: Document, paragraph: Paragraph, rows: Sequence[Sequence[str]]) -> Table:
    table = document.add_table(rows=0, cols=2)
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
        row[0].width = Inches(1.8)
        row[1].width = Inches(4.9)
        for p in row[0].paragraphs:
            if p.runs:
                p.runs[0].bold = True
    set_table_font(table)
    paragraph._p.addnext(table._tbl)
    return table


def insert_caption_after(paragraph: Paragraph, text: str) -> Paragraph:
    caption = insert_paragraph_after(paragraph, text=text, style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(caption, size=12)
    return caption


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_paragraph(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def remove_between(document: Document, start_text: str, end_text: str) -> None:
    deleting = False
    to_remove: List[Paragraph] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == start_text:
            deleting = True
            to_remove.append(paragraph)
            continue
        if deleting and text == end_text:
            break
        if deleting:
            to_remove.append(paragraph)
    for paragraph in to_remove:
        remove_paragraph(paragraph)


def renumber_existing_captions(document: Document) -> None:
    replacements = {
        "Hình 1-1: Use Case tổng quan của hệ thống": "Hình 3-1: Use Case tổng quan của hệ thống",
        "Hình 1-2: Use Case chi tiết của vòng lặp gameplay": "Hình 3-2: Use Case chi tiết của vòng lặp gameplay",
        "Hình 1-3: Sequence Diagram của luồng đăng nhập và nhập tên hiển thị": "Hình 3-3: Sequence Diagram của luồng đăng nhập và nhập tên hiển thị",
        "Hình 1-4: Sequence Diagram của luồng chiến đấu, nhận EXP và chọn buff": "Hình 3-4: Sequence Diagram của luồng chiến đấu, nhận EXP và chọn buff",
        "Hình 1-5: Sequence Diagram của luồng game over, gửi điểm và leaderboard": "Hình 3-5: Sequence Diagram của luồng game over, gửi điểm và leaderboard",
        "Hình 1-6: Activity Diagram của vòng lặp gameplay": "Hình 3-6: Activity Diagram của vòng lặp gameplay",
        "Hình 1-7: Component Diagram của hệ thống": "Hình 3-7: Component Diagram của hệ thống",
        "Hình 1-8: Class Diagram của nhóm Player và Progression": "Hình 3-8: Class Diagram của nhóm Player và Progression",
        "Hình 1-9: Class Diagram của nhóm Enemy và Projectile": "Hình 3-9: Class Diagram của nhóm Enemy và Projectile",
        "Hình 1-10: Class Diagram của nhóm UI, NPC và Backend": "Hình 3-10: Class Diagram của nhóm UI, NPC và Backend",
        "Hình 2-1: Giao diện khi vào game": "Hình 4-1: Giao diện khi vào game",
        "Hình 2-2: Giao diện khi chiến đấu": "Hình 4-2: Giao diện khi chiến đấu",
        "Hình 2-3: Nhân vật player điều khiển": "Hình 4-3: Nhân vật player điều khiển",
        "Hình 2-4: Quái bay": "Hình 4-4: Quái bay",
        "Hình 2-5: Kẻ thù đánh xa": "Hình 4-5: Kẻ thù đánh xa",
        "Hình 2-6: Kẻ thù đánh gần": "Hình 4-6: Kẻ thù đánh gần",
        "Hình 2-7: Ba quái vật lớn": "Hình 4-7: Ba quái vật lớn",
        "Hình 2-8: Map 1": "Hình 4-8: Map 1",
        "Hình 2-9: Map 2": "Hình 4-9: Map 2",
        "Hình 2-10: Map 3": "Hình 4-10: Map 3",
        "Hình 2-11: Hiển thị ấn ESC để mở menu": "Hình 4-11: Hiển thị ấn ESC để mở menu",
        "Hình 2-12: Hiển thị wave hiện tại": "Hình 4-12: Hiển thị wave hiện tại",
        "Hình 2-13: Máu và kinh nghiệm hiện tại của player": "Hình 4-13: Máu và kinh nghiệm hiện tại của player",
        "Hình 2-14: Giao diện hiển thị trước khi bắt đầu vào chiến đấu": "Hình 4-14: Giao diện hiển thị trước khi bắt đầu vào chiến đấu",
        "Hình 2-15: Giao diện chọn buff mỗi khi lên cấp": "Hình 4-15: Giao diện chọn buff mỗi khi lên cấp",
        "Hình 2-16: Giao diện hiển thị tên người chơi khi vào game": "Hình 4-16: Giao diện hiển thị tên người chơi khi vào game",
        "Hình 2-17: Giao diện bảng xếp hạng theo kinh nghiệm": "Hình 4-17: Giao diện bảng xếp hạng theo kinh nghiệm",
        "Hình 2-18: Giao diện nhập tên khi vào game lần đầu": "Hình 4-18: Giao diện nhập tên khi vào game lần đầu",
        "Hình 2-19:Giao diện cài đặt": "Hình 4-19: Giao diện cài đặt",
        "Hình 2-20: Giao diện Pause Game": "Hình 4-20: Giao diện Pause Game",
    }
    for paragraph in document.paragraphs:
        txt = paragraph.text.strip()
        if txt in replacements:
            paragraph.text = replacements[txt]
            set_paragraph_font(paragraph, size=12)


def add_flow_notes_after_heading(document: Document, mapping: dict[str, str]) -> None:
    for heading_text, note in mapping.items():
        anchor = find_paragraph(document, heading_text)
        next_text = None
        for paragraph in document.paragraphs:
            if next_text is None and paragraph == anchor:
                next_text = ""
                continue
            if next_text == "":
                if paragraph.text.strip():
                    insert_labeled_paragraph_after(paragraph.previous_paragraph if hasattr(paragraph, "previous_paragraph") else anchor, "Luồng và thông số chính", note)
                    break


def draw_box(draw: ImageDraw.ImageDraw, xy, text: str, font, fill: str = "#f4f4f4") -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=22, outline="black", width=4, fill=fill)
    max_width = x2 - x1 - 24
    lines = wrap_text(text, font, max_width)
    total_height = sum(font.getbbox(line or "A")[3] - font.getbbox(line or "A")[1] for line in lines) + 8 * (len(lines) - 1)
    y = y1 + (y2 - y1 - total_height) / 2
    for line in lines:
        bbox = font.getbbox(line or "A")
        w = bbox[2] - bbox[0]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=font, fill="black")
        y += (bbox[3] - bbox[1]) + 8


def wrap_text(text: str, font, max_width: int) -> List[str]:
    words = text.split()
    if not words:
        return [""]
    lines = [words[0]]
    for word in words[1:]:
        candidate = f"{lines[-1]} {word}"
        if font.getlength(candidate) <= max_width:
            lines[-1] = candidate
        else:
            lines.append(word)
    return lines


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill="black", width=4)
    x1, y1 = start
    x2, y2 = end
    if x1 == x2 and y2 > y1:
        pts = [(x2, y2), (x2 - 10, y2 - 18), (x2 + 10, y2 - 18)]
    elif x2 > x1:
        pts = [(x2, y2), (x2 - 18, y2 - 10), (x2 - 18, y2 + 10)]
    else:
        pts = [(x2, y2), (x2 + 18, y2 - 10), (x2 + 18, y2 + 10)]
    draw.polygon(pts, fill="black")


def create_diagrams() -> tuple[Path, Path]:
    regular = ImageFont.truetype(FONT_REGULAR, 28)
    bold = ImageFont.truetype(FONT_BOLD, 28)
    title_font = ImageFont.truetype(FONT_BOLD, 36)

    model_path = ASSET_DIR / "model_overview.png"
    img = Image.new("RGB", (1600, 980), "white")
    draw = ImageDraw.Draw(img)
    draw.text((450, 24), "Sơ đồ mô hình bài toán tổng quát", font=title_font, fill="black")

    boxes = {
        "player": (620, 120, 980, 220, "Player\nPlayerController, PlayerAttack,\nPlayerHealth, PlayerData"),
        "wave": (110, 320, 470, 420, "Wave System\nWaveSpawner, WaveConfig,\nBoss wave, Endless wave"),
        "enemy": (620, 320, 980, 420, "Enemy System\nEnemy, EnemyData,\nMelee/Ranged/Fly/Boss"),
        "buff": (1110, 320, 1490, 420, "Buff & Progression\nPlayerLevelSystem,\nBuffCardManager"),
        "pool": (110, 570, 470, 670, "ObjectPool\nProjectile, enemy,\neffect, spirit, chest"),
        "ui": (620, 570, 980, 670, "UI Layer\nHUD, ChallengePanel,\nNameInputPanel, Leaderboard"),
        "playfab": (1110, 570, 1490, 670, "PlayFab Backend\nLogin, Display Name,\nHighScore leaderboard"),
        "theme": (620, 790, 980, 890, "Map Theme\nMapThemeManager,\nLoading transition"),
    }
    for key, box in boxes.items():
        draw_box(draw, box[:4], box[4], bold if key == "player" else regular)

    arrow(draw, (620, 170), (470, 370))
    arrow(draw, (800, 220), (800, 320))
    arrow(draw, (980, 170), (1110, 370))
    arrow(draw, (470, 370), (620, 370))
    arrow(draw, (980, 370), (1110, 370))
    arrow(draw, (290, 420), (290, 570))
    arrow(draw, (800, 420), (800, 570))
    arrow(draw, (1300, 420), (1300, 570))
    arrow(draw, (980, 620), (1110, 620))
    arrow(draw, (800, 670), (800, 790))
    img.save(model_path)

    flow_path = ASSET_DIR / "system_flow.png"
    img = Image.new("RGB", (1600, 1500), "white")
    draw = ImageDraw.Draw(img)
    draw.text((390, 24), "Sơ đồ luồng hoạt động tổng quát của hệ thống", font=title_font, fill="black")
    flow_boxes = [
        (500, 110, 1100, 210, "1. Khởi động game và đăng nhập PlayFab"),
        (500, 250, 1100, 350, "2. Kiểm tra hồ sơ và yêu cầu nhập Display Name nếu cần"),
        (500, 390, 1100, 490, "3. Người chơi tương tác ChallengePostNPC và nhấn StartGame"),
        (500, 530, 1100, 630, "4. WaveSpawner chuẩn bị wave và spawn enemy"),
        (500, 670, 1100, 770, "5. Player di chuyển, dash và tự động tấn công"),
        (500, 810, 1100, 910, "6. Enemy chết -> ExpDropper cộng EXP vào PlayerLevelSystem"),
        (300, 950, 760, 1050, "7a. Đủ EXP -> hiện 3 buff,\náp dụng nâng cấp rồi tiếp tục"),
        (840, 950, 1300, 1050, "7b. Hết enemy -> hoàn tất wave,\ncó thể đổi theme hoặc vào boss wave"),
        (500, 1090, 1100, 1190, "8. Player chết -> lấy totalExpGained làm điểm cuối trận"),
        (500, 1230, 1100, 1330, "9. Gửi HighScore lên PlayFab và tải Leaderboard"),
    ]
    for box in flow_boxes:
        draw_box(draw, box[:4], box[4], regular, fill="#fbfbfb")
    for y1, y2 in [(210, 250), (350, 390), (490, 530), (630, 670), (770, 810), (910, 950)]:
        arrow(draw, (800, y1), (800, y2))
    arrow(draw, (760, 1000), (800, 1090))
    arrow(draw, (840, 1000), (800, 1090))
    arrow(draw, (800, 1190), (800, 1230))
    img.save(flow_path)

    return model_path, flow_path


def fill_table_document_anchor(document: Document, anchor: Paragraph, rows: Sequence[Sequence[str]]) -> Paragraph:
    table = insert_table_after(document, anchor, rows)
    table_anchor = insert_paragraph_after(anchor, style="Normal")
    table._tbl.addnext(table_anchor._p)
    return table_anchor


def rebuild_section_14(document: Document, model_path: Path, flow_path: Path) -> None:
    remove_between(document, "1.4.1. Yêu cầu chức năng", "CHƯƠNG 2 \nMỘT SỐ KIẾN THỨC CƠ BẢN THỰC HIỆN ĐỀ TÀI")
    anchor = find_paragraph(document, "1.4. Mô tả yêu cầu và mô hình bài toán")

    anchor = insert_paragraph_after(anchor, "1.4.1. Yêu cầu chức năng", style="Heading 3")
    for block in FUNCTION_BLOCKS:
        anchor = insert_paragraph_after(anchor, block["title"], style="Heading 4")
        anchor = insert_labeled_paragraph_after(anchor, "Mô tả chức năng", block["mo_ta"])
        anchor = insert_labeled_paragraph_after(anchor, "Luồng thực hiện", block["luong"])
        anchor = insert_labeled_paragraph_after(anchor, "Dữ liệu vào/ra", block["du_lieu"])
        anchor = fill_table_document_anchor(document, anchor, block["specs"])

    anchor = insert_paragraph_after(anchor, "1.4.2. Yêu cầu phi chức năng", style="Heading 3")
    anchor = insert_labeled_paragraph_after(anchor, "Hiệu năng", "Các đối tượng runtime có tần suất tạo hủy cao như enemy, projectile, effect, damage text, spirit và chest phải ưu tiên tái sử dụng qua ObjectPool để giữ ổn định tốc độ khung hình.")
    anchor = insert_labeled_paragraph_after(anchor, "Tổ chức mã nguồn", "Mã nguồn cần được chia rõ theo các cụm Player, Enemy, Wave, UI, Buff, Systems và Utils để hạn chế phụ thuộc chéo và giữ đường đi dữ liệu dễ theo dõi.")
    anchor = insert_labeled_paragraph_after(anchor, "Khả năng mở rộng", "Enemy mới, buff mới và wave mới nên được thêm bằng prefab hoặc ScriptableObject trước, chỉ sửa logic lõi khi thật cần thiết.")
    anchor = insert_labeled_paragraph_after(anchor, "Ổn định luồng chơi", "Các trạng thái nhập tên, pause, chọn buff, leaderboard và đổi theme không được để lại input cũ hoặc phá vỡ trạng thái wave hiện hành.")

    anchor = insert_paragraph_after(anchor, "1.4.3. Mô hình bài toán tổng quát", style="Heading 3")
    anchor = insert_labeled_paragraph_after(anchor, "Mục tiêu mô hình", "Mô hình bài toán tổng quát gom toàn bộ gameplay về ba trục chính: điều khiển và chiến đấu của Player, điều phối áp lực từ Wave/Enemy, và đồng bộ trạng thái qua UI cùng PlayFab.")
    anchor = insert_picture_after(anchor, model_path, 6.2)
    anchor = insert_caption_after(anchor, "Hình 1-1: Sơ đồ mô hình bài toán tổng quát của hệ thống")
    anchor = insert_labeled_paragraph_after(anchor, "Diễn giải", "Player là khối trung tâm phát sinh hầu hết dữ liệu runtime. WaveSpawner điều phối enemy và nhịp độ trận, BuffCardManager cùng PlayerLevelSystem điều chỉnh sức mạnh người chơi, còn UI và PlayFab nhận dữ liệu đã xử lý để phản hồi cho người dùng.")
    anchor = insert_caption_after(anchor, "Bảng 1-2: Các khối chức năng, dữ liệu và thông số điều phối chính")
    anchor = fill_table_document_anchor(
        document,
        anchor,
        [
            ("Khối Player", "Nhận input, xử lý di chuyển, dash, tấn công, máu và dữ liệu chiến đấu hiện thời."),
            ("Khối Wave/Enemy", "Sinh quái theo WaveConfig, quản lý boss wave, activeEnemies và cơ chế tăng độ khó."),
            ("Khối Progression", "Cộng EXP, tăng level, chọn buff và cập nhật chỉ số hoặc kỹ năng cho player."),
            ("Khối UI", "Hiển thị HUD, panel tương tác, pause, chọn buff, nhập tên và leaderboard."),
            ("Khối Backend", "Xác thực Custom ID, đồng bộ Display Name, gửi HighScore và tải dữ liệu xếp hạng."),
        ],
    )

    anchor = insert_paragraph_after(anchor, "1.4.4. Mô tả hoạt động của hệ thống", style="Heading 3")
    anchor = insert_labeled_paragraph_after(anchor, "Tổng quan luồng", "Hoạt động của hệ thống đi theo chuỗi khởi động - định danh - bắt đầu trận - chiến đấu - tăng tiến sức mạnh - kết thúc trận - đồng bộ leaderboard.")
    anchor = insert_picture_after(anchor, flow_path, 6.1)
    anchor = insert_caption_after(anchor, "Hình 1-2: Sơ đồ luồng hoạt động tổng quát của hệ thống")
    for idx, text in enumerate(
        [
            "Khởi động game: PlayFabLeaderboardManager đăng nhập bằng Custom ID và tải hồ sơ người chơi.",
            "Định danh: nếu hồ sơ chưa có Display Name, NameInputPanel mở để người chơi nhập tên hợp lệ rồi mới tiếp tục.",
            "Bắt đầu trận: người chơi tương tác ChallengePostNPC, ChallengePanel phát onGameStart và WaveSpawner chuẩn bị wave đầu tiên.",
            "Chiến đấu: PlayerController nhận input di chuyển và dash; PlayerAttack tự động quét mục tiêu, sinh projectile từ ObjectPool và gây sát thương lên enemy.",
            "Nhận EXP: khi enemy chết, ExpDropper cộng expValue vào PlayerLevelSystem, HUD đồng thời cập nhật level, EXP và tổng điểm tích lũy.",
            "Lên cấp và chọn buff: nếu vượt ngưỡng expToNextLevel, CardSelectionPanel dừng gameplay tạm thời, hiển thị 3 buff hợp lệ và áp dụng lựa chọn vào PlayerData hoặc các manager kỹ năng.",
            "Chuyển wave và đổi theme: khi toàn bộ enemy của wave bị hạ, WaveSpawner hoàn tất wave, có thể kích hoạt boss wave hoặc chuyển theme nếu bước sang chu kỳ 10 wave mới.",
            "Kết thúc trận và đồng bộ online: khi player chết, PlayerHealth lấy totalExpGained làm điểm cuối trận, gửi HighScore lên PlayFab và mở LeaderboardPanel để hiển thị kết quả.",
        ],
        start=1,
    ):
        anchor = insert_labeled_paragraph_after(anchor, f"Bước {idx}", text)


def add_contextual_notes(document: Document) -> None:
    for heading_text, note in CH3_INSERTS.items():
        anchor = find_paragraph(document, heading_text)
        insert_labeled_paragraph_after(anchor, "Luồng và thông số chính", note)

    for heading_text, note in CH4_INSERTS.items():
        anchor = find_paragraph(document, heading_text)
        insert_labeled_paragraph_after(anchor, "Luồng và thông số kỹ thuật", note)


def save_document(document: Document) -> Path:
    output_path = OUTPUT_DIR / "BaoCao_DATN_2_restructured.docx"
    document.save(output_path)
    document.save(DOC_PATH)
    return output_path


def main() -> None:
    ensure_dirs()
    model_path, flow_path = create_diagrams()
    document = Document(DOC_PATH)
    renumber_existing_captions(document)
    rebuild_section_14(document, model_path, flow_path)
    add_contextual_notes(document)
    output_path = save_document(document)
    print(output_path)


if __name__ == "__main__":
    main()
